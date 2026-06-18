#!/usr/bin/env python3
from __future__ import annotations
"""
US-China Relations Tracker Scraper

Usage:
    python scraper.py                  # run all sources
    python scraper.py --source fmprc_conf
    python scraper.py --source fmprc_remarks
    python scraper.py --source mofcom
    python scraper.py --source state
    python scraper.py --source whitehouse
    python scraper.py --source treasury
    python scraper.py --source ustr
    python scraper.py --source wardept

Environment variables (put in .env or export directly):
    GEMINI_API_KEY   — from aistudio.google.com
    GROQ_API_KEY     — fallback, from console.groq.com
"""

import argparse
import logging
import os
import re
import sqlite3
import time
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin

from dotenv import load_dotenv
load_dotenv()

import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from google import genai
from google.genai import types
from groq import Groq
from typing import Optional
from pydantic import BaseModel

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-7s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)

# ── Constants ─────────────────────────────────────────────────────────────────

DB_PATH  = "tracker.db"
DOC_PATH = "tracker_output.docx"

# 4 s between Gemini calls keeps us under the 15 RPM free-tier cap.
GEMINI_SLEEP  = 4
# 2 s between HTTP requests — polite crawl rate for every source.
REQUEST_SLEEP = 2

BROWSER_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/125.0.0.0 Safari/537.36"
)

FMPRC_SPOKESPERSONS = {
    "Lin Jian", "Mao Ning", "Wang Wenbin", "Zhao Lijian",
    "Hua Chunying", "Liu Pengyu", "Guo Jiakun",
}

# Keywords that mark a Q or A as US-China relevant.
# A Q+A pair is kept if ANY exchange in the pair matches.
RELEVANCE_KEYWORDS = re.compile(
    r"\b(?:"
    # US references
    r"U\.S\b|United States|America[n]?|Washington"
    # US officials (update as administrations change)
    r"|Trump|Biden|Rubio|Bessent|Lutnick|Navarro|Sullivan|Blinken|Yellen"
    # Trade / sanctions / tech
    r"|tariff|trade war|sanction|export control|import duty|reciprocal"
    r"|semiconductor|chip|AI|artificial intelligence|technology transfer"
    r"|Huawei|TikTok|CATL|BYD|COSCO|SMIC"
    # Territorial / political flashpoints
    r"|Taiwan|Hong Kong|Xinjiang|Tibet|South China Sea|East China Sea"
    # Finance / currency
    r"|yuan|RMB|currency|manipulation|trade deficit|trade surplus"
    # Other recurring topics
    r"|fentanyl|espionage|intellectual property|forced transfer"
    r"|NATO|G7|G20|QUAD|AUKUS|decoupling|de-risk"
    r")",
    re.IGNORECASE,
)

# ── Database ──────────────────────────────────────────────────────────────────

def init_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        "CREATE TABLE IF NOT EXISTS seen_urls "
        "(url TEXT PRIMARY KEY, date_seen TEXT)"
    )
    conn.commit()
    return conn


def is_seen(conn: sqlite3.Connection, url: str) -> bool:
    return conn.execute(
        "SELECT 1 FROM seen_urls WHERE url = ?", (url,)
    ).fetchone() is not None


def mark_seen(conn: sqlite3.Connection, url: str) -> None:
    conn.execute(
        "INSERT OR IGNORE INTO seen_urls (url, date_seen) VALUES (?, ?)",
        (url, datetime.utcnow().isoformat()),
    )
    conn.commit()


# ── LLM ───────────────────────────────────────────────────────────────────────

GEMINI_MODEL = "gemini-2.5-flash"


def init_llm() -> genai.Client:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY is not set")
    return genai.Client(api_key=api_key)


def _call_groq(prompt: str) -> str:
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("Gemini rate-limited and GROQ_API_KEY is not set")
    if len(prompt) > 9000:
        prompt = prompt[:9000] + "\n\n[content truncated]"
    client = Groq(api_key=api_key)
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=600,
    )
    return completion.choices[0].message.content.strip()


def call_llm(client: genai.Client, prompt: str, retries: int = 2) -> str:
    """Call Gemini with sleep; auto-fall back to Groq on 429."""
    time.sleep(GEMINI_SLEEP)
    for attempt in range(retries + 1):
        try:
            resp = client.models.generate_content(
                model=GEMINI_MODEL, contents=prompt
            )
            return resp.text.strip()
        except Exception as exc:
            err = str(exc).lower()
            if "429" in err or "quota" in err or "resource_exhausted" in err:
                log.warning("Gemini rate-limited — falling back to Groq")
                return _call_groq(prompt)
            if attempt < retries:
                time.sleep(8 * (attempt + 1))
                continue
            raise


def classify_relevance(model: genai.Client, text: str) -> tuple[bool, str]:
    snippet = text[:3000]
    result = call_llm(
        model,
        "Is this press release relevant to US-China relations, trade, or Taiwan? "
        "Reply with YES or NO and one sentence reason.\n\n" + snippet,
    )
    return result.upper().startswith("YES"), result


def translate_to_english(model: genai.Client, text: str) -> str:
    """Translate Chinese text to English, preserving Q&A structure."""
    parts = []
    for i in range(0, len(text), 5500):
        chunk = text[i : i + 5500]
        translated = call_llm(
            model,
            "Translate the following Chinese text to English exactly. "
            "Preserve the Q&A structure. Do not summarize or paraphrase. "
            "Romanize speaker names (e.g. '林剑：' → 'Lin Jian:').\n\n" + chunk,
        )
        parts.append(translated)
    return "\n\n".join(parts)


def generate_summary(model: genai.Client, text: str, source_name: str) -> str:
    """
    Write a tracker-style summary paragraph matching the existing doc style.
    e.g. "Foreign Ministry Spokesperson Lin Jian answered questions about U.S.
    tariffs on Chinese goods, Taiwan arms sales, and critical minerals."
    """
    snippet = text[:3500]
    return call_llm(
        model,
        "Write a 1-2 sentence summary for a US-China policy tracker. "
        "State WHO (with their full name and title) said or did WHAT, and list the "
        "specific topics covered. Be concrete. Use past tense. No fluff. "
        "Do not start with 'This', 'The entry', or 'I'. "
        "Match this style: "
        "'Foreign Ministry Spokesperson Lin Jian answered questions about U.S. tariffs "
        "on Chinese goods, Taiwan arms sales, and critical mineral export controls at "
        "the regular daily press conference.' "
        "Or: 'Treasury Secretary Scott Bessent testified before the House Financial "
        "Services Committee, discussing Chinese IPOs, rare earth leverage, and digital "
        "currency competition with China.'\n\n"
        f"Source: {source_name}\n\nContent:\n{snippet}",
    )


def extract_key_paragraphs(model: genai.Client, text: str, n: int = 4) -> list[str]:
    """
    Extract the most China-relevant verbatim paragraphs from a press release.
    Returns a list of paragraph strings.
    """
    snippet = text[:5000]
    result = call_llm(
        model,
        f"Extract the {n} most important verbatim paragraphs about China, US-China "
        f"relations, trade, or Taiwan from this text. "
        f"Return each paragraph separated by the delimiter '|||'. "
        f"Do not summarize or paraphrase — use exact text.\n\n{snippet}",
    )
    paras = [p.strip() for p in result.split("|||") if p.strip()]
    return paras[:n]


# ── HTTP client ───────────────────────────────────────────────────────────────

def make_client(verify_ssl: bool = True) -> httpx.Client:
    return httpx.Client(
        headers={"User-Agent": BROWSER_UA, "Accept-Language": "en-US,en;q=0.9"},
        follow_redirects=True,
        timeout=30.0,
        verify=verify_ssl,
    )


def fetch(client: httpx.Client, url: str, retries: int = 3) -> httpx.Response | None:
    for attempt in range(retries):
        try:
            time.sleep(REQUEST_SLEEP)
            resp = client.get(url)
            if resp.status_code == 429:
                wait = int(resp.headers.get("Retry-After", 30))
                log.warning(f"429 on {url} — waiting {wait}s")
                time.sleep(wait)
                continue
            resp.raise_for_status()
            return resp
        except httpx.HTTPStatusError as exc:
            log.warning(f"HTTP {exc.response.status_code}: {url}")
            return None
        except Exception as exc:
            log.warning(f"Request failed ({attempt+1}/{retries}): {url} — {exc}")
            if attempt < retries - 1:
                time.sleep(6 * (attempt + 1))
    return None


# ── Document writer ───────────────────────────────────────────────────────────

def _set_doc_defaults(doc: Document) -> None:
    """
    Apply document-wide defaults to the Normal style:
      - Times New Roman 12 pt
      - 1.15 line spacing
      - 8 pt space after paragraph
    All paragraphs added afterwards inherit these automatically.
    """
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after  = Pt(8)


def _run(paragraph, text: str, bold: bool = False, italic: bool = False):
    """Add a run with explicit TNR 12 so theme fonts can't override."""
    r = paragraph.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.name   = "Times New Roman"
    r.font.size   = Pt(12)
    return r


def _indented_para(doc: Document) -> object:
    """
    Return a new paragraph with a full paragraph-level left indent of 0.5 inches.
    Every line — not just the first — is indented because we set left_indent,
    NOT first_line_indent.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent       = Inches(0.5)
    pf.first_line_indent = Pt(0)          # no extra first-line indent on top
    return p


def get_or_create_doc(path: str) -> Document:
    p = Path(path)
    if p.exists():
        doc = Document(str(p))
    else:
        doc = Document()
        for para in list(doc.paragraphs):
            if not para.text.strip():
                para._element.getparent().remove(para._element)
    _set_doc_defaults(doc)
    return doc


def _fmt_date(date: datetime) -> str:
    """Format as 'Monday, June 17, 2026' — cross-platform."""
    return date.strftime("%A, %B ") + str(date.day) + ", " + str(date.year)


def add_date_heading(doc: Document, date: datetime) -> None:
    """Bold normal paragraph for the date, matching Part 3 of the existing tracker."""
    p = doc.add_paragraph()
    _run(p, _fmt_date(date), bold=True)


def add_qa_entry(
    doc: Document,
    date: datetime,
    summary: str,
    exchanges: list[dict],
) -> None:
    """
    Write a Q&A entry matching the exact format of the existing tracker docs.

    exchanges: list of dicts with keys:
        type    — "Q" | "A" | "CONT"
        speaker — outlet/spokesperson name (None for CONT)
        text    — paragraph text

    Q paragraphs:   outlet name bold+italic, ": text" italic only — full paragraph indented
    A paragraphs:   "Speaker:" bold, " text" plain              — full paragraph indented
    CONT paragraphs: plain continuation                         — full paragraph indented
    All lines in every exchange paragraph are indented (left_indent, not first_line_indent).
    """
    add_date_heading(doc, date)
    _run(doc.add_paragraph(), summary)          # summary — plain, no indent

    for ex in exchanges:
        p = _indented_para(doc)

        if ex["type"] == "Q":
            _run(p, ex["speaker"], bold=True, italic=True)
            _run(p, ": " + ex["text"], italic=True)

        elif ex["type"] == "A":
            _run(p, ex["speaker"] + ":", bold=True)
            _run(p, " " + ex["text"])

        else:  # CONT — continuation paragraph, no speaker label
            _run(p, ex["text"])

    # Minimal blank line between entries — space_after on last exchange already
    # creates visual gap; this adds a clean paragraph break.
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after  = Pt(0)
    sep.paragraph_format.space_before = Pt(0)


def add_release_entry(
    doc: Document,
    date: datetime,
    summary: str,
    body_paragraphs: list[str],
) -> None:
    """
    Write a press release / statement entry.
    body_paragraphs: verbatim paragraphs indented as a block.
    If a paragraph starts with 'Speaker: text', the speaker name is bolded.
    Every line in each block paragraph is indented (paragraph-level, not first-line).
    """
    add_date_heading(doc, date)
    _run(doc.add_paragraph(), summary)

    speaker_re = re.compile(r"^([A-Z][A-Za-z0-9 \-'\.\:]{1,50}):\s+(.+)$", re.DOTALL)
    for para_text in body_paragraphs:
        p = _indented_para(doc)
        m = speaker_re.match(para_text)
        if m:
            _run(p, m.group(1) + ":", bold=True)
            _run(p, " " + m.group(2))
        else:
            _run(p, para_text)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_after  = Pt(0)
    sep.paragraph_format.space_before = Pt(0)


# ── Q&A parser ────────────────────────────────────────────────────────────────

_QA_RE = re.compile(r"^([A-Z][A-Za-z0-9 \-'\.]{1,40}):\s+(.+)$", re.DOTALL)


def _build_exchanges(paragraphs: list[str], spokespersons: set[str]) -> list[dict]:
    """
    Convert a list of text paragraphs into exchange dicts.
    Continuation paragraphs (no 'Name: text' pattern) become CONT entries
    so they render as separate indented paragraphs — matching the existing docs.
    """
    exchanges: list[dict] = []
    in_qa = False

    for para in paragraphs:
        para = para.strip()
        if not para or len(para) < 8:
            continue
        m = _QA_RE.match(para)
        if m:
            speaker = m.group(1).strip()
            text    = m.group(2).strip()
            is_sp   = any(sp.lower() in speaker.lower() for sp in spokespersons)
            exchanges.append({
                "type":    "A" if is_sp else "Q",
                "speaker": speaker,
                "text":    text,
            })
            in_qa = True
        elif in_qa:
            exchanges.append({"type": "CONT", "speaker": None, "text": para})

    return exchanges


def filter_relevant_exchanges(exchanges: list[dict]) -> list[dict]:
    """
    Group exchanges into Q→A blocks, then keep only blocks where any exchange
    text matches RELEVANCE_KEYWORDS.

    This ensures that if a journalist asks about Trump/tariffs but the answer
    only mentions China's position (or vice versa), both Q and A are included.
    A block is: one Q + all following A/CONT until the next Q.
    A-only blocks (no preceding Q) are treated as their own block.
    """
    if not exchanges:
        return []

    blocks: list[list[dict]] = []
    current: list[dict] = []

    for ex in exchanges:
        if ex["type"] == "Q":
            if current:
                blocks.append(current)
            current = [ex]
        else:                          # A or CONT
            if current:
                current.append(ex)
            else:
                current = [ex]        # A with no preceding Q

    if current:
        blocks.append(current)

    result = []
    for block in blocks:
        combined = " ".join(ex.get("text", "") for ex in block)
        if RELEVANCE_KEYWORDS.search(combined):
            result.extend(block)

    return result


def parse_qa(html: str, spokespersons: set[str]) -> list[dict]:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()

    content = (
        soup.find("div", class_=re.compile(r"content|article|detail|main", re.I))
        or soup.find("div", id=re.compile(r"content|article|detail|main", re.I))
    )
    if not content:
        divs = soup.find_all("div")
        content = max(divs, key=lambda d: len(d.get_text()), default=soup)

    paragraphs = [p.get_text(separator=" ").strip() for p in content.find_all("p")]
    if not paragraphs:
        paragraphs = [p.get_text(separator=" ").strip() for p in soup.find_all("p")]

    return _build_exchanges(paragraphs, spokespersons)


# ── Q&A parsing schemas and helpers ──────────────────────────────────────────

class QAExchange(BaseModel):
    type: str  # "Q" | "A" | "CONT"
    speaker: Optional[str] = None
    text: str


class QAResponse(BaseModel):
    exchanges: list[QAExchange]


def call_llm_json(
    client: genai.Client,
    prompt: str,
    schema: type[BaseModel],
    retries: int = 2,
) -> dict:
    """Call Gemini to get a structured JSON response using a Pydantic schema."""
    time.sleep(GEMINI_SLEEP)
    for attempt in range(retries + 1):
        try:
            resp = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=schema,
                ),
            )
            import json
            return json.loads(resp.text)
        except Exception as exc:
            err = str(exc).lower()
            log.warning(f"Gemini JSON API call failed (attempt {attempt+1}/{retries+1}): {exc}")
            if "429" in err or "quota" in err or "resource_exhausted" in err:
                log.warning("Gemini rate-limited or quota exceeded.")
            if attempt < retries:
                time.sleep(8 * (attempt + 1))
                continue
            raise


def split_single_paragraph(text: str) -> str:
    """
    If the text has 1 or 2 lines but contains multiple speaker colon patterns,
    insert newlines before the speaker labels so that each turn is on its own line.
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if len(lines) <= 2:
        # Match word boundaries for uppercase speaker names (like 'Lin Jian:', 'Q:', 'REPORTER:')
        # which are not preceded by 'http' or 'https'.
        pattern = re.compile(
            r"(?<!https)(?<!http)\b([A-Z][A-Za-z0-9 \-'\.]{1,40}):\s+",
            re.MULTILINE
        )
        matches = list(pattern.finditer(text))
        if len(matches) >= 2:
            parts = []
            last_idx = 0
            for m in matches:
                start = m.start()
                if start > last_idx:
                    segment = text[last_idx:start].strip()
                    if segment:
                        parts.append(segment)
                last_idx = start
            parts.append(text[last_idx:].strip())
            return "\n\n".join(parts)
    return text


def parse_qa_with_llm(model: genai.Client, text: str, spokespersons: set[str]) -> list[dict]:
    """
    Use Gemini to parse and classify raw text into Q&A exchanges.
    Returns a list of dicts with keys: 'type', 'speaker', 'text'.
    """
    sp_list = ", ".join(sorted(spokespersons)) if spokespersons else "None specified"
    
    prompt = f"""You are an expert transcript parser for a foreign policy tracker.
Your task is to parse and segment the following transcript text into a structured JSON list of dialogue segments.

The known official spokespersons / government officials for this source are: [{sp_list}].

Each item in the output JSON list MUST have exactly these three keys:
- "type": "Q" or "A" or "CONT"
- "speaker": The name of the speaker (a string, or null if type is "CONT")
- "text": The exact verbatim text spoken (string)

Rules for "type":
- "Q" (Question): Spoken by a journalist, reporter, media outlet, or other questioning party. Usually contains a question, and Q always precedes A.
- "A" (Answer): Spoken by a government official, spokesperson, or official speaker. If there are no questions in the text at all (e.g., a white paper, unilateral statement, or reading of a document), everything should be classified as "A".
- "CONT" (Continuation): Use this for any subsequent/continuation paragraphs spoken by the SAME speaker as the previous paragraph, when there is a paragraph break but no new speaker label in the raw text. For CONT, "speaker" must be null.

Important structural guidelines:
- "A" always follows "Q" (or "CONT" of "Q"). "A" can also follow "A" (or "CONT" of "A") if another official speaks.
- Do NOT translate, summarize, or paraphrase. The "text" field MUST contain the exact verbatim English text from the transcript.
- If the speaker name is provided, extract it exactly (e.g. "Lin Jian", "Reporter", "Journalist", "Wall Street Journal").

Transcript text to parse:
\"\"\"
{text}
\"\"\"
"""
    result = call_llm_json(model, prompt, QAResponse)
    
    exchanges = []
    for item in result.get("exchanges", []):
        ex_type = item.get("type", "CONT")
        if ex_type not in ("Q", "A", "CONT"):
            ex_type = "CONT"
        
        exchanges.append({
            "type": ex_type,
            "speaker": item.get("speaker"),
            "text": item.get("text", "").strip(),
        })
    return exchanges


def parse_qa_from_plaintext(
    text: str,
    spokespersons: set[str],
    model: genai.Client | None = None,
) -> list[dict]:
    # 1. Preprocess text to split single massive paragraphs if needed
    text = split_single_paragraph(text)
    
    # 2. If model is provided, try LLM-driven classification
    if model:
        try:
            return parse_qa_with_llm(model, text, spokespersons)
        except Exception as exc:
            log.warning(f"LLM-driven Q&A classification failed, falling back to regex: {exc}")
    
    # 3. Fallback to classic regex-based parser
    paragraphs = text.split("\n")
    return _build_exchanges(paragraphs, spokespersons)


# ── RSS helpers ───────────────────────────────────────────────────────────────

def _rss_text(tag) -> str:
    return tag.get_text(strip=True) if tag else ""


def parse_rss(xml_text: str, limit: int = 20) -> list[dict]:
    soup = BeautifulSoup(xml_text, "xml")
    items = []
    for item in soup.find_all("item")[:limit]:
        content_tag = item.find("content:encoded") or item.find("description")
        link = _rss_text(item.find("link")) or _rss_text(item.find("guid"))
        items.append({
            "_type":       "rss",
            "title":       _rss_text(item.find("title")),
            "link":        link,
            "pubDate":     _rss_text(item.find("pubDate")),
            "content_raw": str(content_tag) if content_tag else "",
        })
    return items


def item_url(item: dict) -> str:
    return item.get("link", item.get("guid", ""))


def item_title(item: dict) -> str:
    if item.get("_type") == "rss":
        return item.get("title", "")
    raw = item.get("title", {})
    return BeautifulSoup(
        raw.get("rendered", "") if isinstance(raw, dict) else str(raw),
        "html.parser",
    ).get_text()


def item_content(item: dict) -> str:
    if item.get("_type") == "rss":
        return item.get("content_raw", "")
    raw = item.get("content", {})
    return raw.get("rendered", "") if isinstance(raw, dict) else str(raw)


def item_date(item: dict) -> datetime:
    if item.get("_type") == "rss":
        try:
            from email.utils import parsedate_to_datetime
            return parsedate_to_datetime(item["pubDate"]).replace(tzinfo=None)
        except Exception:
            return datetime.utcnow()
    try:
        return datetime.fromisoformat(item.get("date", ""))
    except Exception:
        return datetime.utcnow()


# ── Source: FMPRC ─────────────────────────────────────────────────────────────

def scrape_fmprc(
    list_url: str,
    label: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
    client: httpx.Client,
) -> None:
    log.info(f"[fmprc/{label}] {list_url}")
    resp = fetch(client, list_url)
    if not resp:
        log.error(f"[fmprc/{label}] Failed to fetch list")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    raw_links = []
    for a in soup.find_all("a", href=re.compile(r"/\d{6}/t\d+_\d+\.html")):
        raw_links.append((urljoin(list_url, a["href"]), a.get_text(strip=True)))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:10]
    log.info(f"[fmprc/{label}] {len(new_links)} new items")

    for url, title in new_links:
        try:
            resp = fetch(client, url)
            if not resp:
                continue

            raw_html = resp.text
            cjk = sum(1 for c in raw_html if "一" <= c <= "鿿")

            if cjk > 100:
                log.info(f"[fmprc/{label}] Translating: {url}")
                plaintext  = BeautifulSoup(raw_html, "html.parser").get_text(separator="\n")
                translated = translate_to_english(model, plaintext[:7000])
                exchanges  = parse_qa_from_plaintext(translated, FMPRC_SPOKESPERSONS, model)
                work_text  = translated
            else:
                exchanges = parse_qa(raw_html, FMPRC_SPOKESPERSONS)
                work_text = BeautifulSoup(raw_html, "html.parser").get_text()

            if not exchanges:
                log.info(f"[fmprc/{label}] No Q&A found — skipping: {url}")
                mark_seen(conn, url)
                continue

            exchanges = filter_relevant_exchanges(exchanges)
            if not exchanges:
                log.info(f"[fmprc/{label}] No relevant exchanges — skipping: {url}")
                mark_seen(conn, url)
                continue

            date_m = re.search(r"t(\d{8})_", url)
            date   = (
                datetime.strptime(date_m.group(1), "%Y%m%d") if date_m
                else datetime.utcnow()
            )

            summary = generate_summary(model, work_text, f"MFA {label}")
            log.info(f"[fmprc/{label}] Writing: {title}")
            add_qa_entry(doc, date, summary, exchanges)
            doc.save(DOC_PATH)
            mark_seen(conn, url)

        except Exception as exc:
            log.error(f"[fmprc/{label}] Error on {url}: {exc}")


# ── Source: MOFCOM ────────────────────────────────────────────────────────────

def scrape_mofcom(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    list_url = "https://english.mofcom.gov.cn/News/PressConference/index.html"
    log.info(f"[mofcom] {list_url}")
    client = make_client(verify_ssl=False)  # MOFCOM cert untrusted by Python CA bundle

    resp = fetch(client, list_url)
    if not resp:
        log.error("[mofcom] Failed to fetch list")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    raw_links = []
    for a in soup.find_all("a", href=re.compile(r"/News/PressConference/")):
        if "index" in a.get("href", ""):
            continue
        href  = urljoin(list_url, a["href"])
        title = a.get_text(strip=True)
        if title and len(title) > 10:
            raw_links.append((href, title))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:10]
    log.info(f"[mofcom] {len(new_links)} new items")

    for url, title in new_links:
        try:
            resp = fetch(client, url)
            if not resp:
                continue

            plain  = BeautifulSoup(resp.text, "html.parser").get_text()
            is_rel, _ = classify_relevance(model, f"{title}\n\n{plain[:2500]}")
            if not is_rel:
                mark_seen(conn, url)
                continue

            exchanges = parse_qa(resp.text, {"Spokesperson", "Minister", "Deputy"})
            exchanges = filter_relevant_exchanges(exchanges)
            if not exchanges:
                log.info(f"[mofcom] No relevant exchanges — skipping: {url}")
                mark_seen(conn, url)
                continue

            date_m    = re.search(r"(\d{4}-\d{2}-\d{2})", resp.text)
            date      = (
                datetime.strptime(date_m.group(1), "%Y-%m-%d") if date_m
                else datetime.utcnow()
            )

            summary = generate_summary(model, plain, "MOFCOM press conference")
            log.info(f"[mofcom] Writing: {title}")
            add_qa_entry(doc, date, summary, exchanges)
            doc.save(DOC_PATH)
            mark_seen(conn, url)

        except Exception as exc:
            log.error(f"[mofcom] Error on {url}: {exc}")


# ── Source: State Dept ────────────────────────────────────────────────────────

def scrape_state(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    api_url = (
        "https://www.state.gov/wp-json/wp/v2/press_releases"
        "?per_page=20&orderby=date&_fields=id,date,title,link,content,excerpt"
    )
    rss_url = "https://www.state.gov/rss-feed/press-releases/feed/"
    client  = make_client()

    log.info("[state] Fetching WP API")
    items: list[dict] = []
    resp = fetch(client, api_url)
    if resp:
        try:
            items = resp.json()
        except Exception:
            items = []

    if not items:
        log.warning("[state] WP API failed — trying RSS")
        resp_rss = fetch(client, rss_url)
        if resp_rss:
            items = parse_rss(resp_rss.text)

    if not items:
        log.error("[state] Both endpoints failed")
        return

    new_items = [it for it in items if not is_seen(conn, item_url(it))]
    log.info(f"[state] {len(new_items)} new items")

    for it in new_items:
        url = item_url(it)
        try:
            title = item_title(it)
            plain = BeautifulSoup(item_content(it), "html.parser").get_text()
            is_rel, _ = classify_relevance(model, f"{title}\n\n{plain[:2500]}")
            if not is_rel:
                mark_seen(conn, url)
                continue

            paras   = extract_key_paragraphs(model, plain)
            summary = generate_summary(model, plain, "State Department")
            log.info(f"[state] Writing: {title}")
            add_release_entry(doc, item_date(it), summary, paras)
            doc.save(DOC_PATH)
            mark_seen(conn, url)

        except Exception as exc:
            log.error(f"[state] Error on {url}: {exc}")


# ── Source: White House ───────────────────────────────────────────────────────

def scrape_whitehouse(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    rss_url = "https://www.whitehouse.gov/news/feed/"
    client  = make_client()

    log.info(f"[whitehouse] {rss_url}")
    resp = fetch(client, rss_url)
    if not resp:
        log.error("[whitehouse] Failed to fetch RSS")
        return

    items     = parse_rss(resp.text)
    new_items = [it for it in items if not is_seen(conn, item_url(it))]
    log.info(f"[whitehouse] {len(new_items)} new items")

    for it in new_items:
        url = item_url(it)
        try:
            title = item_title(it)
            plain = BeautifulSoup(item_content(it), "html.parser").get_text()
            is_rel, _ = classify_relevance(model, f"{title}\n\n{plain[:2500]}")
            if not is_rel:
                mark_seen(conn, url)
                continue

            paras   = extract_key_paragraphs(model, plain)
            summary = generate_summary(model, plain, "White House")
            log.info(f"[whitehouse] Writing: {title}")
            add_release_entry(doc, item_date(it), summary, paras)
            doc.save(DOC_PATH)
            mark_seen(conn, url)

        except Exception as exc:
            log.error(f"[whitehouse] Error on {url}: {exc}")


# ── Source: Treasury ──────────────────────────────────────────────────────────

def scrape_treasury(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    list_url = "https://home.treasury.gov/news/press-releases"
    client   = make_client()

    log.info(f"[treasury] {list_url}")
    resp = fetch(client, list_url, retries=4)
    if not resp:
        log.error("[treasury] Failed after 4 retries — skipping")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    seen_hrefs: set[str] = set()
    raw_links: list[tuple[str, str]] = []
    for a in soup.find_all("a", href=re.compile(r"/news/press-releases/")):
        href  = urljoin(list_url, a["href"])
        title = a.get_text(strip=True)
        if title and href != list_url and href not in seen_hrefs:
            seen_hrefs.add(href)
            raw_links.append((href, title))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:10]
    log.info(f"[treasury] {len(new_links)} new items")

    for url, title in new_links:
        try:
            resp = fetch(client, url, retries=4)
            if not resp:
                continue

            plain  = BeautifulSoup(resp.text, "html.parser").get_text()
            is_rel, _ = classify_relevance(model, f"{title}\n\n{plain[:2500]}")
            if not is_rel:
                mark_seen(conn, url)
                continue

            date_m = re.search(
                r"(January|February|March|April|May|June|July|August|"
                r"September|October|November|December)\s+\d{1,2},\s+\d{4}",
                plain,
            )
            date  = (
                datetime.strptime(date_m.group(0), "%B %d, %Y") if date_m
                else datetime.utcnow()
            )
            paras   = extract_key_paragraphs(model, plain)
            summary = generate_summary(model, plain, "Treasury Department")
            log.info(f"[treasury] Writing: {title}")
            add_release_entry(doc, date, summary, paras)
            doc.save(DOC_PATH)
            mark_seen(conn, url)

        except Exception as exc:
            log.error(f"[treasury] Error on {url}: {exc}")


# ── Source: USTR ──────────────────────────────────────────────────────────────

def scrape_ustr(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    list_url = "https://ustr.gov/about-us/policy-offices/press-office/press-releases"
    client   = make_client()

    log.info(f"[ustr] {list_url}")
    resp = fetch(client, list_url)
    if not resp:
        log.error("[ustr] Failed to fetch list")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    seen_hrefs: set[str] = set()
    raw_links: list[tuple[str, str]] = []
    for a in soup.find_all(
        "a",
        href=re.compile(r"/about/policy-offices/press-office/press-releases/\d{4}/"),
    ):
        href  = urljoin("https://ustr.gov", a["href"])
        title = a.get_text(strip=True)
        if title and href not in seen_hrefs:
            seen_hrefs.add(href)
            raw_links.append((href, title))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:10]
    log.info(f"[ustr] {len(new_links)} new items")

    for url, title in new_links:
        try:
            resp = fetch(client, url)
            if not resp:
                continue

            plain  = BeautifulSoup(resp.text, "html.parser").get_text()
            is_rel, _ = classify_relevance(model, f"{title}\n\n{plain[:2500]}")
            if not is_rel:
                mark_seen(conn, url)
                continue

            date_m = re.search(r"(\d{4}-\d{2}-\d{2})", resp.text)
            date   = (
                datetime.strptime(date_m.group(1), "%Y-%m-%d") if date_m
                else datetime.utcnow()
            )
            paras   = extract_key_paragraphs(model, plain)
            summary = generate_summary(model, plain, "USTR")
            log.info(f"[ustr] Writing: {title}")
            add_release_entry(doc, date, summary, paras)
            doc.save(DOC_PATH)
            mark_seen(conn, url)

        except Exception as exc:
            log.error(f"[ustr] Error on {url}: {exc}")


# ── Source: Department of Defense (war.gov) ───────────────────────────────────

def scrape_wardept(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    list_url = "https://www.war.gov/News/Releases/"
    log.info(f"[wardept] Launching Playwright for {list_url}")

    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        log.error("[wardept] playwright not installed — run: playwright install chromium")
        return

    html_list = None
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            page    = browser.new_context(
                user_agent=BROWSER_UA,
                viewport={"width": 1280, "height": 800},
            ).new_page()
            try:
                page.goto(list_url, wait_until="networkidle", timeout=30_000)
                html_list = page.content()
            except Exception as exc:
                log.error(f"[wardept] Page load failed: {exc}")
            finally:
                browser.close()
    except Exception as exc:
        log.error(f"[wardept] Playwright error: {exc}")
        return

    if not html_list:
        return

    soup = BeautifulSoup(html_list, "html.parser")
    raw_links = []
    for a in soup.find_all("a", href=re.compile(r"/News/Releases/")):
        href = a.get("href", "")
        if href.rstrip("/") in ("/News/Releases", list_url.rstrip("/")):
            continue
        full  = urljoin("https://www.war.gov", href)
        title = a.get_text(strip=True)
        if title:
            raw_links.append((full, title))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:10]
    log.info(f"[wardept] {len(new_links)} new items")

    client = make_client()
    for url, title in new_links:
        try:
            resp = fetch(client, url)
            if not resp:
                continue

            plain  = BeautifulSoup(resp.text, "html.parser").get_text()
            is_rel, _ = classify_relevance(model, f"{title}\n\n{plain[:2500]}")
            if not is_rel:
                mark_seen(conn, url)
                continue

            date_m = re.search(
                r"(January|February|March|April|May|June|July|August|"
                r"September|October|November|December)\s+\d{1,2},\s+\d{4}",
                plain,
            )
            date  = (
                datetime.strptime(date_m.group(0), "%B %d, %Y") if date_m
                else datetime.utcnow()
            )
            paras   = extract_key_paragraphs(model, plain)
            summary = generate_summary(model, plain, "Department of Defense")
            log.info(f"[wardept] Writing: {title}")
            add_release_entry(doc, date, summary, paras)
            doc.save(DOC_PATH)
            mark_seen(conn, url)

        except Exception as exc:
            log.error(f"[wardept] Error on {url}: {exc}")


# ── Main ──────────────────────────────────────────────────────────────────────

SOURCES = {
    "fmprc_conf":    "MFA press conferences",
    "fmprc_remarks": "MFA spokesperson remarks",
    "mofcom":        "MOFCOM press conferences",
    "state":         "State Dept",
    "whitehouse":    "White House",
    "treasury":      "Treasury",
    "ustr":          "USTR",
    "wardept":       "Dept of Defense (war.gov)",
}


def main() -> None:
    parser = argparse.ArgumentParser(description="US-China tracker scraper")
    parser.add_argument(
        "--source",
        choices=list(SOURCES.keys()),
        help="Run a single source only (default: all)",
    )
    args = parser.parse_args()

    model  = init_llm()
    conn   = init_db()
    doc    = get_or_create_doc(DOC_PATH)
    client = make_client()

    run_all = args.source is None
    s = args.source

    def run(key: str, fn, *fn_args):
        if run_all or s == key:
            try:
                fn(*fn_args)
            except Exception as exc:
                log.error(f"[{key}] Unhandled error: {exc}")

    run("fmprc_conf",    scrape_fmprc,
        "https://www.fmprc.gov.cn/eng/xw/fyrbt/lxjzh/",
        "press conference", model, conn, doc, client)

    run("fmprc_remarks", scrape_fmprc,
        "https://www.fmprc.gov.cn/eng/xw/fyrbt/fyrbt/",
        "spokesperson remarks", model, conn, doc, client)

    run("mofcom",        scrape_mofcom,     model, conn, doc)
    run("state",         scrape_state,      model, conn, doc)
    run("whitehouse",    scrape_whitehouse, model, conn, doc)
    run("treasury",      scrape_treasury,   model, conn, doc)
    run("ustr",          scrape_ustr,       model, conn, doc)
    run("wardept",       scrape_wardept,    model, conn, doc)

    log.info("All done.")


if __name__ == "__main__":
    main()

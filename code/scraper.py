#!/usr/bin/env python3
from __future__ import annotations
"""
US-China Relations Tracker Scraper

Usage:
    python scraper.py                  # run all sources
    python scraper.py --source fmprc_conf
    python scraper.py --source fmprc_remarks
    python scraper.py --source mfa_leadership_speeches
    python scraper.py --source mfa_leadership_activity
    python scraper.py --source mofcom
    python scraper.py --source mofcom_daily
    python scraper.py --source scio
    python scraper.py --source mnd
    python scraper.py --source state
    python scraper.py --source whitehouse
    python scraper.py --source treasury
    python scraper.py --source ustr
    python scraper.py --source wardept

See SOURCES.md for the full source inventory (incl. what's intentionally
NOT covered — X/Twitter, TruthSocial, WeChat, YouTube — and why), and
NOTES.md for this project's running decision/assumption log.

Environment variables (put in .env or export directly):
    GEMINI_API_KEY   — from aistudio.google.com
    GROQ_API_KEY     — fallback, from console.groq.com
"""

import argparse
import io
import json
import logging
import os
import re
import sqlite3
import subprocess
import sys
import time
from collections import Counter
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from urllib.parse import urljoin

from dotenv import load_dotenv
load_dotenv()

import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from google import genai
from google.genai import types
from groq import Groq
from tqdm import tqdm
from tqdm.contrib.logging import logging_redirect_tqdm
from typing import Callable, Optional
from pydantic import BaseModel

# Full detail (every fetch, every LLM call, every "wrote N entries") always
# goes to a timestamped file under logs/ — but only WARNING/ERROR go to the
# terminal by default. Added 2026-09-02 alongside the progress bar: with
# the old single console handler, ~150 INFO lines per run scrolled the
# progress bar off-screen and buried the final "saved to" message,
# defeating the point of adding either for a non-technical user. `-v`/
# `--verbose` (wired up in main()) restores full INFO on-screen for
# debugging — this only changes what's PRINTED, not what's logged;
# logs/*.log always has the complete record either way.
os.makedirs("logs", exist_ok=True)
_log_file_path = os.path.join("logs", f"{datetime.now():%Y-%m-%d_%H%M%S}.log")
_file_handler = logging.FileHandler(_log_file_path)
_file_handler.setLevel(logging.INFO)
_console_handler = logging.StreamHandler()
_console_handler.setLevel(logging.WARNING)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-7s  %(message)s",
    datefmt="%H:%M:%S",
    handlers=[_file_handler, _console_handler],
)
log = logging.getLogger(__name__)


def _utcnow() -> datetime:
    """
    Naive UTC "now" — use this everywhere instead of the deprecated
    `datetime.utcnow()` (Python 3.12+ warns on every call, scheduled for
    removal). Found live, 2026-09-02: those warnings print straight to
    the terminal and land mid-redraw of the progress bar, visually
    splitting one continuously-updating bar into what looks like two.

    Deliberately still returns a NAIVE datetime (no tzinfo), not the
    "correct" fix of switching to timezone-aware datetimes everywhere —
    every date this pipeline stores, sorts, and compares (PENDING_ENTRIES'
    sort key, every process_*_item()'s date fallback, item_date()'s
    strptime results) is naive. Returning an aware datetime from just this
    one helper while everything else stays naive would raise `TypeError:
    can't compare offset-naive and offset-aware datetimes` the first time
    a fallback-dated item and a normally-dated item land in the same sort
    — so this gets the correct instant from the non-deprecated API, then
    immediately strips the tzinfo back off to match every existing
    call site exactly. Not a behavior change, just no longer calling a
    function that is going away.
    """
    return datetime.now(timezone.utc).replace(tzinfo=None)


# ── Constants ─────────────────────────────────────────────────────────────────

# Every file this pipeline generates or maintains (the dedup database, the
# output doc, cost/review logs, the X account cache) lives under output/,
# not the project root — keeps runtime state visually separate from
# source code, and the folder name says what it holds without needing to
# already know the project's conventions. Created here (not lazily at
# each write site) so it exists before ANY of the paths below are first
# used, regardless of which script imports this module (scraper.py,
# backtest.py, format_entry.py, seed_dedup_db.py all end up importing
# these constants). Resolved relative to the current working directory
# (NOT this file's own location) — same as every other relative path in
# this project (input/past_trackers/, etc.) — so run these scripts
# with the project root as your working directory, e.g.
# `python code/scraper.py` from the project root, not `cd code && python
# scraper.py`.
DATA_DIR = "output"
os.makedirs(DATA_DIR, exist_ok=True)

DB_PATH  = os.path.join(DATA_DIR, "tracker.db")
DOC_PATH = os.path.join(DATA_DIR, "tracker_output.docx")
FLAGGED_REVIEW_PATH = os.path.join(DATA_DIR, "flagged_for_review.md")

# 0.5 s between Gemini calls — was 16s (RPM=4, checked 2026-09-01), then
# 1s (after confirming a real tier upgrade to RPM=1000/TPM=1M/RPD=10,000,
# 2026-09-02). Cut further the same day: checked a real run's own log —
# most calls in a normal run are small (classify_relevance: ~200-250
# tokens), and at 1s they were STILL spaced ~1-2s apart, meaning the
# sleep itself, not real network/generation latency, was still the
# dominant cost for the common case. At RPM=1000 (16.7/sec) neither RPM
# nor a single-threaded script's realistic max request rate is a genuine
# constraint. TPM=1M/min (16,667 tok/sec) is the only real limit worth
# respecting — 0.5s allows ~20K tok/sec if EVERY call happened to be a
# maximal ~10K-token translate back-to-back, modestly over TPM in that
# specific worst case, but real runs mix in far more small calls than
# that, and _gemini_on_cooldown()'s 429-triggered backoff (below) is the
# actual safety net if that edge case is ever hit — a brief graceful
# fallback to Groq/OpenRouter, not a crash or overspend. Re-check the
# actual dashboard (aistudio.google.com) before changing this again —
# never infer a new value from behavior or a third-party aggregator
# figure alone; two different wrong assumptions already got corrected
# this way.
GEMINI_SLEEP  = 0.5
# 2 s between HTTP requests — polite crawl rate for every source.
REQUEST_SLEEP = 2

# How many new (not-yet-seen) items each source processes per run. Was a
# hardcoded 10 in 9 different places, sized for a DAILY cadence — the
# actual intended use is a WEEKLY tracker (run ~once a week, not once a
# day), and several sources (State, Treasury, USTR, MOFCOM) can easily
# publish more than 10 total items — not all China-relevant, but enough
# that a real China-relevant one could fall outside the most-recent-10
# window and get silently skipped. Raised to a week-appropriate number;
# each source's own relevance filtering (keyword pre-filters,
# classify_relevance) still does the real work of picking out what
# matters — this constant only bounds how far back into a list page we
# bother looking before giving up. See NOTES.md, 2026-09-01.
#
# Briefly raised to 150 (2026-09-02) after confirming USTR alone had 128
# genuinely unseen items, then reverted back to 30 the same day on user
# pushback — correctly: a backlog that size is a ONE-TIME catch-up
# problem (this project hadn't targeted USTR much during earlier
# development), not a recurring weekly need, since the user's actual
# cadence always targets "last week." A permanent global raise pays real
# ongoing cost (more page fetches, more time) as the default for every
# future run, for a problem that's actually temporary — the wrong lever.
# `--max-items` (see main()) is the right one: a deliberate, one-time
# override for an explicit catch-up run, not a change to what every
# normal week costs. This module-level value stays the cheap, correct
# default; main() may reassign it via `global` if that flag is passed.
MAX_NEW_ITEMS_PER_RUN = 30

BROWSER_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/125.0.0.0 Safari/537.36"
)

FMPRC_SPOKESPERSONS = {
    "Lin Jian", "Mao Ning", "Wang Wenbin", "Zhao Lijian",
    "Hua Chunying", "Liu Pengyu", "Guo Jiakun",
}

MND_SPOKESPERSONS = {
    "Wu Qian", "Zhang Xiaogang", "Jiang Bin", "Ren Guoqiang", "Tan Kefei",
}

# Chinese-transliteration -> correct English spelling for names that recur
# constantly in this content and that a weaker fallback model gets wrong.
# Added 2026-08-05 after the Groq fallback (see GROQ_MODEL) rendered
# 贝森特/格里尔 (Bessent/Greer) as "Bessenet"/"Gill" in a real translation —
# not a huge deal once, but these two names alone show up in a large share
# of MOFCOM/FMPRC entries about trade talks, so it's worth the one-line
# prompt hint. Passed into translate_to_english()'s prompt; update as the
# cast of recurring officials changes.
KNOWN_NAME_ROMANIZATIONS = {
    "贝森特": "Bessent", "格里尔": "Greer", "卢特尼克": "Lutnick",
    "鲁比奥": "Rubio", "万斯": "Vance", "特朗普": "Trump",
    "纳瓦罗": "Navarro", "耶伦": "Yellen", "王毅": "Wang Yi",
    "何立峰": "He Lifeng",
    # Added 2026-09-03 after a live run rendered 庞德伟 (US Ambassador to
    # China David Perdue's official Chinese name) as "David Pond" — a
    # naive phonetic guess at a name the model has no other way to know,
    # since 庞德伟 isn't a transliteration Perdue's actual English name
    # would predict. Same failure shape as Bessent/Greer above, just for a
    # name with even less phonetic overlap.
    "庞德伟": "Perdue",
    # Added 2026-09-03, same session — a broader pass over every US
    # official this tracker is likely to name, verified live against real
    # FMPRC/MFA/Xinhua/mainland-media usage (not guessed) — see
    # input/notes/US_OFFICIALS_CHINESE_NAMES.md for the full table,
    # sources, and how to extend this when the cast of officials changes.
    "赫格塞斯": "Hegseth", "加巴德": "Gabbard", "拉特克利夫": "Ratcliffe",
    "莱维特": "Leavitt", "米勒": "Miller",
    # Kept for the prior (Biden) administration too — older past-tracker
    # weeks used in backtest.py still reference these.
    "拜登": "Biden", "沙利文": "Sullivan", "布林肯": "Blinken",
}

# Keywords that mark a Q or A as US-China relevant.
# A Q+A pair is kept if ANY exchange in the pair matches.
# Bare alternatives only (no \b(?:...)\b wrapper) — kept separate from the
# compiled patterns below so BOTH RELEVANCE_KEYWORDS and
# US_SOURCE_RELEVANCE_KEYWORDS get a properly closed trailing \b, not just
# a leading one. The original single `\b(?:...)` (no closing \b) let a
# short alternative match as a mere PREFIX of a longer unrelated word —
# found live, 2026-09-01: bare "AI" (meant to catch "AI"/"artificial
# intelligence" export-control language) matched inside "UnmAnned AIrcraft
# Systems" — the leading boundary (space→"A") was satisfied, and with no
# trailing \b required, "AI" as the first two letters of "AIrcraft" was
# good enough. Same risk for other short alternatives (e.g. "chip" inside
# a longer word). This wasted an LLM call on the free pre-filter passing
# a document with nothing really relevant in it — not a correctness bug
# (classify_relevance's own LLM judgment still catches it), but a real
# cost one, and the same "match a boundary you didn't actually mean"
# category as several bugs already fixed today.
#
# `s?`/`(y|ies)` on the common-noun terms below: found via test_scraper.py
# — adding the trailing \b above (to fix the AI/aircraft bug) also broke
# matching ordinary PLURAL forms ("tariffs", "sanctions", "chips",
# "semiconductors", "export controls"), since "tariff\b" doesn't match
# inside "tariffs" (no boundary between "tariff" and its own "s"). These
# plurals are at least as common as the singular in real headlines, so
# this was a real regression traded for the earlier fix, caught only by
# the test suite, not by the live testing that found the original bug.
_RELEVANCE_ALTERNATIVES = (
    # US references — "America[n]" excludes "Latin/South/Central America"
    # (a region, not the US) via negative lookbehind, added 2026-09-02
    # after a real false positive: a Latin America item got matched and
    # counted as US-relevant purely because "America" is a substring.
    # Each lookbehind is fixed-width (Python's re requirement) — "Latin "/
    # "South " are 6 chars, "Central " is 8, each checked independently.
    #
    # "U.S"/"United States" also exclude a following "dollar(s)" via
    # negative lookahead, added 2026-09-02 after the same-shaped false
    # positive on the other side: a Chinese economic release quoting a
    # USD-equivalent figure (e.g. "8100亿美元") gets translated to English
    # as "810 billion U.S. dollars" — CHINESE_RELEVANCE_KEYWORDS below
    # deliberately never matches bare 美元 for exactly this reason, but
    # the translated English text was then still tripping THIS regex on
    # "U.S." alone, undoing that protection after translation. "U.S.
    # dollar(s)"/"United States dollar(s)" is just naming a currency unit,
    # not talking about the United States as a country/actor.
    #
    # Bare "US" (no periods), added 2026-09-03 — a REAL, high-impact gap
    # found live: an LLM translation of a real MOFCOM Q&A rendered the
    # abbreviation as "US government" (no periods) rather than "U.S.
    # government," and since this pattern's only US-abbreviation
    # alternative was the strictly-dotted "U\.S\b", the entire exchange
    # (a reporter's real question about a proposed US tariff hike) failed
    # the explicit-US-mention check and got silently dropped — not a
    # crash, not a log warning, just gone. Both spellings are extremely
    # common in real English text (translated or original) and must both
    # count. `(?-i:US)` scopes OFF this whole pattern's outer IGNORECASE
    # flag for just this one alternative — required so it matches
    # uppercase "US" only, NOT the lowercase pronoun "us" ("let us know",
    # "join us"), which the rest of this pattern's case-insensitivity
    # would otherwise happily match. Same dollar-amount exclusion as
    # "U.S" above, for the same reason ("US dollars" as a currency unit).
    r"U\.S(?!\.?\s*dollars?\b)\b|(?-i:US)(?!\s*dollars?\b)\b"
    r"|United States(?!\s+dollars?\b)"
    r"|(?<!Latin )(?<!South )(?<!Central )America[n]?|Washington"
    # US officials (update as administrations change)
    r"|Trump|Biden|Rubio|Bessent|Lutnick|Navarro|Sullivan|Blinken|Yellen"
    # Trade / sanctions / tech
    r"|tariffs?|trade wars?|sanctions?|export controls?|import dut(?:y|ies)|reciprocal"
    r"|semiconductors?|chips?|AI|artificial intelligence|technology transfers?"
    r"|Huawei|TikTok|CATL|BYD|COSCO|SMIC"
    # Territorial / political flashpoints
    r"|Taiwan|Hong Kong|Xinjiang|Tibet|South China Sea|East China Sea"
    # Finance / currency — NOT bare "yuan"/"RMB"/"currency": those are just
    # the unit a dollar figure gets quoted in and show up on nearly any
    # economic-statistics release regardless of subject (the same trap as
    # bare 美元/人民币 in Chinese — see CHINESE_RELEVANCE_KEYWORDS below).
    # "currency manipulation"/"exchange rate" as full phrases are the actual
    # policy-relevance signal.
    r"|currency manipulation|exchange rate polic(?:y|ies)|trade deficits?|trade surplus(?:es)?"
    # Other recurring topics
    r"|fentanyl|espionage|intellectual property|forced transfers?"
    r"|NATO|G7|G20|QUAD|AUKUS|decoupling|de-risk"
)
RELEVANCE_KEYWORDS = re.compile(rf"\b(?:{_RELEVANCE_ALTERNATIVES})\b", re.IGNORECASE)

# Narrower than RELEVANCE_KEYWORDS above — used specifically by
# filter_relevant_exchanges() to decide WHICH Q&A block(s) to keep from a
# multi-topic press conference, where the broader topic list (Taiwan/AI/
# semiconductor/etc., without requiring an actual US mention) is too
# loose: a single FMPRC daily conference covers many different bilateral
# stories, and a topic like "Taiwan" or "artificial intelligence" comes up
# just as often in a purely China-Japan or China-everyone context as a
# US-China one. Found live, 2026-09-01, via direct user inspection of a
# real generated tracker doc: a China-Japan exchange about "erroneous
# remarks...on Taiwan" (zero US mentions) and a humanoid-robots Q&A about
# "cooperation on artificial intelligence...with all other countries"
# (also zero US mentions) both got included solely because "Taiwan"/
# "artificial intelligence" are in the broad list. RELEVANCE_KEYWORDS
# itself stays as-is for other uses (built into US_SOURCE_RELEVANCE_
# KEYWORDS's free whole-document pre-filter), where looseness is fine
# because classify_relevance's own LLM judgment supplies the real
# precision downstream — filter_relevant_exchanges has no such downstream
# check, so it needs to be precise on its own. Matches the user's own
# stated rule for this exact judgment call: an explicit US/named-official
# mention, not a shared topic.
#
# "America[n]" excludes "Latin/South/Central America" via negative
# lookbehind, added 2026-09-02 after a real false positive: a Latin
# America item got counted as an explicit US mention purely because
# "America" is a substring — and THIS regex in particular has no LLM
# judgment downstream to catch it (filter_relevant_exchanges and PRC
# accounts' relevance check both decide directly off this match, no
# second look). Each lookbehind is fixed-width (Python's re requirement).
#
# "U.S"/"United States" also exclude a following "dollar(s)" via negative
# lookahead, added 2026-09-02 — same fix and same reasoning as
# _RELEVANCE_ALTERNATIVES above: a Chinese figure quoted in USD (美元)
# translates to English as "U.S. dollars," which is naming a currency
# unit, not an actual mention of the United States.
#
# Bare "US" (no periods), added 2026-09-03 — same real, high-impact bug
# and same fix as _RELEVANCE_ALTERNATIVES above: this regex's own
# docstring above says it decides Q&A block relevance with NO downstream
# LLM check, which made this gap especially damaging — a real MOFCOM
# exchange (a reporter's question about a proposed US tariff hike,
# translated as "US government" with no periods) got silently dropped
# with nothing to catch the miss. `(?-i:US)` scopes off the outer
# IGNORECASE flag for just this alternative so it matches uppercase "US"
# only, not the lowercase pronoun "us."
_EXPLICIT_US_MENTION_RE = re.compile(
    r"\b(?:U\.S(?!\.?\s*dollars?\b)\b|(?-i:US)(?!\s*dollars?\b)\b"
    r"|United States(?!\s+dollars?\b)"
    r"|(?<!Latin )(?<!South )(?<!Central )America[n]?|Washington"
    r"|Trump|Biden|Rubio|Bessent|Lutnick|Navarro|Sullivan|Blinken|Yellen)\b",
    re.IGNORECASE,
)

# Chinese-language counterpart, used as a free pre-filter on RAW Chinese
# text — BEFORE spending an LLM call on translation/classification — for
# sources that publish Chinese-only (fmprc's CJK branch, mfa_leadership,
# mnd's non-bilingual pages). Added 2026-08-04 after burning through both
# providers' free daily quota in one afternoon of testing; most items on
# these list pages are routine diplomatic/departmental readouts with no
# US-China angle at all, so skipping those for free instead of paying for
# a translation + classification call on every single one matters a lot.
#
# Deliberately does NOT include bare 美元 ("US dollar") or bare 美 alone:
# nearly every Chinese economic/statistics article includes a USD-equivalent
# parenthetical (e.g. "5.5万亿元（约8100亿美元）") regardless of subject, and 美
# alone is a common character in unrelated words (美丽="beautiful", 完美=
# "perfect", etc.) — either would make this pre-filter pass almost
# everything, defeating the point. 美国 ("[the] US[A]", two characters) is
# the safe token; standalone currency mentions don't carry it.
# classify_relevance() screens US-originated press releases (state/
# treasury/ustr/whitehouse/wardept/scio/mofcom-EN), so unlike
# RELEVANCE_KEYWORDS above (which checks for the US SIDE within
# inherently-China-focused FMPRC/MOFCOM/MND content) this needs an explicit
# "does this mention China at all" check too — RELEVANCE_KEYWORDS alone has
# no bare "China"/"Chinese"/"Beijing"/"Xi Jinping" token. Kept as a
# separate pattern rather than folding into RELEVANCE_KEYWORDS itself,
# because filter_relevant_exchanges() uses that one to find the US-relevant
# block WITHIN a Chinese press conference — every block there already
# "mentions China" trivially, so adding China-mention terms there would
# make it pass everything and defeat its purpose.
US_SOURCE_RELEVANCE_KEYWORDS = re.compile(
    rf"\b(?:{_RELEVANCE_ALTERNATIVES}|China\w*|Beijing|Xi Jinping)\b",
    re.IGNORECASE,
)

CHINESE_RELEVANCE_KEYWORDS = re.compile(
    "|".join([
        # US references (NOT bare 美元/美 — see note above)
        "美国", "中美", "美方", "华盛顿", "白宫",
        # US officials (update as administrations change)
        "特朗普", "拜登", "鲁比奥", "贝森特", "卢特尼克", "纳瓦罗",
        "布林肯", "耶伦", "万斯",
        # Trade / sanctions / tech
        "关税", "贸易战", "制裁", "出口管制", "芯片", "半导体",
        "华为", "TikTok", "抖音国际版", "比亚迪", "宁德时代", "中远海运", "中芯国际",
        # Territorial / political flashpoints
        "台湾", "台独", "香港", "新疆", "西藏", "南海", "东海",
        # Finance / currency — NOT bare 人民币 ("RMB/yuan"): it's just the
        # unit almost every Chinese economic-statistics release quotes an
        # amount in (e.g. "5.5万亿元人民币"), regardless of subject — the
        # exact same trap as bare 美元. "汇率操纵" etc. as full phrases are
        # the actual policy-relevance signal, not the currency name alone.
        "汇率操纵", "贸易逆差", "贸易顺差",
        # Other recurring topics
        "芬太尼", "间谍", "知识产权", "强制转让技术",
        "北约", "七国集团", "二十国集团", "脱钩", "去风险",
    ])
)

# Narrower than CHINESE_RELEVANCE_KEYWORDS above — this is ONLY "does this
# paragraph name the US," not "is this paragraph on some China-policy-
# adjacent topic." Added 2026-09-01 per user request to replace an LLM
# judgment call with the simple rule they already use by hand: a Chinese
# paragraph that names the US gets included, no LLM verdict needed. Same
# 美元-exclusion logic as CHINESE_RELEVANCE_KEYWORDS: every term here is a
# 2+ character compound ending in a character other than 元, so "美元" (US
# dollar) and "人民币" (RMB) never match — no separate exclusion regex
# needed, the term list itself just doesn't overlap with those spellings.
_CHINESE_US_MENTION_RE = re.compile(
    "|".join(["美国", "中美", "美方", "华盛顿", "白宫", "驻美", "访美", "对美", "赴美"])
)


def select_relevant_chinese_paragraphs(raw_zh_text: str) -> list[str]:
    """
    Keyword-only paragraph filter for Chinese-sourced release-type content —
    replaces extract_key_paragraphs()'s LLM judgment call for this case, per
    user request 2026-09-01. Operates on the ORIGINAL Chinese text, BEFORE
    translation, so only the paragraphs that actually matter get translated
    at all (cheaper — one translate_to_english call over just the matches,
    not the whole page) and nothing is paraphrased, dropped, or misjudged by
    an LLM call. Deliberately narrower than the general-purpose extraction
    prompt it replaces: "does this paragraph name the US" is a simpler,
    fully reproducible rule, at the cost of missing a paragraph that's
    relevant only by implication (no literal US mention). Used by every
    Chinese-source caller as of 2026-09-03, MFA leadership included — see
    finalize_release_item()'s `raw_zh_text` parameter.
    """
    paragraphs = [p.strip() for p in raw_zh_text.split("\n") if p.strip()]
    return [p for p in paragraphs if _CHINESE_US_MENTION_RE.search(p)]

# ── Database ──────────────────────────────────────────────────────────────────

def init_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        "CREATE TABLE IF NOT EXISTS seen_urls "
        "(url TEXT PRIMARY KEY, date_seen TEXT)"
    )
    # Durable, queryable record of every entry ever written — added
    # 2026-09-02 so a per-week document can be regenerated on demand from
    # exactly the entries that belong to it, whether they were queued in
    # THIS run or a previous one. Before this, the only place entry
    # CONTENT lived was the master tracker_output.docx itself (formatted
    # Word paragraphs, no queryable date/kind), so the "dated" per-week
    # file was actually just a renamed COPY of the entire cumulative
    # history — see render_doc_for_range()'s docstring for the full story.
    conn.execute(
        "CREATE TABLE IF NOT EXISTS entries ("
        "id INTEGER PRIMARY KEY AUTOINCREMENT, "
        "url TEXT, date TEXT NOT NULL, kind TEXT NOT NULL, "
        "summary TEXT, anchor TEXT, "
        "exchanges_json TEXT, paragraphs_json TEXT)"
    )
    # source_label added 2026-09-02, after this table already had real
    # rows on the user's machine — CREATE TABLE IF NOT EXISTS is a no-op
    # against an existing table even with a different schema, so a real
    # ALTER TABLE is needed to add the column to a database that already
    # exists. Wrapped in try/except since ALTER TABLE ADD COLUMN fails if
    # the column is already there — must stay safe on every init_db()
    # call, not just the first.
    try:
        conn.execute("ALTER TABLE entries ADD COLUMN source_label TEXT")
    except sqlite3.OperationalError:
        pass  # column already exists
    conn.execute("CREATE INDEX IF NOT EXISTS idx_entries_date ON entries(date)")
    conn.commit()
    return conn


def is_seen(conn: sqlite3.Connection, url: str) -> bool:
    return conn.execute(
        "SELECT 1 FROM seen_urls WHERE url = ?", (url,)
    ).fetchone() is not None


def mark_seen(conn: sqlite3.Connection, url: str) -> None:
    conn.execute(
        "INSERT OR IGNORE INTO seen_urls (url, date_seen) VALUES (?, ?)",
        (url, _utcnow().isoformat()),
    )
    conn.commit()


def flag_for_review(url: str, title: str, reason: str) -> None:
    """
    Append a borderline-rejection to FLAGGED_REVIEW_PATH instead of letting
    it disappear silently — added 2026-09-01 per user feedback: "always
    good to have more than less, leave a note for a human reviewer to
    judge." Only called at the LLM-judgment rejection points (an actual
    model call said "not relevant enough" or "no relevant paragraphs"),
    NOT the free keyword-prefilter skips (zero signal at all — those
    aren't real judgment calls, just nothing to review). The item is still
    excluded from tracker_output.docx and still marked seen — this is a
    side-channel audit trail for a human to periodically skim and manually
    add back anything that was cut too aggressively, not a second output
    path into the tracker itself.
    """
    is_new = not Path(FLAGGED_REVIEW_PATH).exists()
    with open(FLAGGED_REVIEW_PATH, "a", encoding="utf-8") as f:
        if is_new:
            f.write(
                "# Flagged for human review\n\n"
                "Items an LLM judgment call excluded from the tracker — not "
                "auto-included anywhere, just logged so a human can spot-check "
                "and manually add any of these back. Zero-signal rejections "
                "(no keyword match at all, never reached an LLM) are NOT "
                "logged here — only actual judgment calls.\n\n"
            )
        f.write(f"- [ ] **{_utcnow().date()}** — {title}\n")
        f.write(f"  {url}\n")
        f.write(f"  > {reason.strip()}\n\n")


# ── LLM ───────────────────────────────────────────────────────────────────────

GEMINI_MODEL = "gemini-3.6-flash"
# Migrated from gemini-2.5-flash 2026-09-04: Google retired it for this
# account ("This model models/gemini-2.5-flash is no longer available to
# new users"), a real 404 reported live by the user's worker. Confirmed
# via client.models.list() which model to move to (matches Google's own
# error message) and checked pricing at
# https://ai.google.dev/gemini-api/docs/pricing (checked 2026-09-04):
# gemini-3.6-flash is $0.75/M input, $3.75/M output through 2026-12-31,
# rising to $1.50/$7.50 on 2027-01-01 — noticeably pricier than
# gemini-2.5-flash's $0.30/$2.50 (2.5x input, 1.5x output). Re-check
# _USD_PER_MILLION below after 2027-01-01 for the price step-up.

# gemini-2.5-flash was a "thinking" model by default that could be fully
# disabled with thinking_budget=0 — see the fuller history in NOTES.md.
# gemini-3.6-flash does NOT support that: thinking_budget=0 is rejected
# outright with a 400 INVALID_ARGUMENT (confirmed live, 2026-09-04, same
# day as the model migration above). It behaves like gemini-2.5-pro
# always did — a minimum amount of thinking is mandatory. Tried budgets
# of 1/128/512/1024 live and all produced roughly the same ~80-100
# thinking tokens regardless, so there's no real lever left to pull here;
# thinking_budget=1 (the smallest legal value) is kept only to signal
# intent, not because it measurably reduces spend.
_GEMINI_THINKING_CONFIG = types.ThinkingConfig(thinking_budget=1)

# Groq retired llama-3.3-70b-versatile from this account's available models
# sometime between 2026-08-04 and 2026-08-05 — the fallback path started
# 404ing with "model does not exist" (a different failure mode than the
# 429s we were used to) instead of the usual rate-limit. Checked
# client.models.list() live and picked the largest general-purpose model
# still available. Re-check this if Groq calls start failing again with a
# 404 (vs. the expected 429) — that's the model-retired signature, not a
# quota problem.
GROQ_MODEL = "openai/gpt-oss-120b"


def init_llm() -> genai.Client:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY is not set")
    return genai.Client(api_key=api_key)


USAGE_LOG_PATH = os.path.join(DATA_DIR, "usage_log.jsonl")

# Real per-million-token USD pricing, paid tier — only Gemini has real
# money behind it right now (billing linked 2026-09-01). Groq's key on
# this account is the free tier (separate from Gemini's billing) and
# OpenRouter's OPENROUTER_MODEL is a ":free" slug — both genuinely $0
# regardless of volume, not "unknown," so they're priced at 0.0 rather
# than omitted. Update if either of those ever changes (a paid Groq plan,
# a non-free OpenRouter model). Cerebras omitted — never successfully
# billed a call (402 Payment Required all session).
# Source: https://ai.google.dev/gemini-api/docs/pricing (checked
# 2026-09-04, after the gemini-3.6-flash migration above) —
# gemini-3.6-flash: $0.75/M input, $3.75/M output through 2026-12-31,
# rising to $1.50/$7.50 on 2027-01-01. "Thinking" tokens are billed at
# the OUTPUT rate, not a separate premium. (Was $0.30/$2.50 under the
# retired gemini-2.5-flash, checked 2026-09-01 — kept here for context
# on why costs look different from earlier runs' [usage] logs.)
# XAI (grok-4.3) pricing from https://docs.x.ai/docs/models/grok-4.3
# (checked 2026-09-01): $1.25/M input, $2.50/M output for requests under
# 200K prompt tokens (our snippets never get remotely close) — real money,
# the user's repurposed-xAI-key credit.
_USD_PER_MILLION = {
    "Gemini":     {"input": 0.75, "output": 3.75},
    "Groq":       {"input": 0.0,  "output": 0.0},
    "OpenRouter": {"input": 0.0,  "output": 0.0},
    "XAI":        {"input": 1.25, "output": 2.50},
}


def _estimate_usd(provider: str, prompt_tokens: int, completion_tokens: int, reasoning_tokens: int) -> float | None:
    """USD estimate for one call, or None for a provider with no pricing
    entry (e.g. Cerebras, never successfully billed). `reasoning_tokens`
    is folded into the output side — Gemini bills thinking tokens at the
    output rate, not separately (see _USD_PER_MILLION's source note)."""
    rates = _USD_PER_MILLION.get(provider)
    if rates is None:
        return None
    output_tokens = (completion_tokens or 0) + (reasoning_tokens or 0)
    return (prompt_tokens or 0) / 1e6 * rates["input"] + output_tokens / 1e6 * rates["output"]


def _log_usage(
    provider: str,
    label: str,
    prompt_tokens: int | None,
    completion_tokens: int | None,
    total_tokens: int | None = None,
    reasoning_tokens: int | None = None,
) -> None:
    """
    One-line token-usage log for every LLM call, tagged by which caller
    made it (label) and which provider actually served it — added
    2026-09-01 per user request, to see where tokens actually go (e.g.
    "finding" calls like classify_relevance vs. "reproducing" calls like
    parse_qa_with_llm) instead of guessing. `reasoning_tokens` surfaces
    GROQ_MODEL's invisible chain-of-thought spend (see the reasoning-token-
    exhaustion bug, NOTES.md) when a provider reports it. Never raises —
    a usage-logging failure should never take down the actual LLM call.

    Also appends a structured record to USAGE_LOG_PATH (one JSON object
    per line) with a real USD estimate — added the same day billing was
    linked to Gemini, so cost can be tracked precisely across a whole
    session/day rather than re-parsing free-text log lines by hand each
    time. See summarize_usage_log() to aggregate it.
    """
    total = total_tokens if total_tokens is not None else (prompt_tokens or 0) + (completion_tokens or 0)
    usd = _estimate_usd(provider, prompt_tokens or 0, completion_tokens or 0, reasoning_tokens or 0)
    reasoning_part = f" reasoning={reasoning_tokens}" if reasoning_tokens else ""
    usd_part = f" usd=${usd:.6f}" if usd is not None else ""
    log.info(
        f"[usage] {label or '(unlabeled)'} via {provider}: "
        f"prompt={prompt_tokens} completion={completion_tokens} total={total}{reasoning_part}{usd_part}"
    )
    try:
        with open(USAGE_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps({
                "ts": _utcnow().isoformat(),
                "label": label or "(unlabeled)",
                "provider": provider,
                "prompt_tokens": prompt_tokens,
                "completion_tokens": completion_tokens,
                "reasoning_tokens": reasoning_tokens,
                "total_tokens": total,
                "usd": usd,
            }) + "\n")
    except Exception as exc:
        log.warning(f"Failed to append to {USAGE_LOG_PATH}: {exc}")


def summarize_usage_log(path: str = USAGE_LOG_PATH) -> dict:
    """
    Aggregate USAGE_LOG_PATH by label and by provider, with a real USD
    total. Returns the summary dict (also printed) so callers/tests can
    inspect it directly. Safe to call any time — reports {} if the log
    doesn't exist yet.
    """
    if not Path(path).exists():
        print(f"No usage log at {path} yet.")
        return {}

    by_label: dict[str, dict] = {}
    by_provider: dict[str, dict] = {}
    grand_total_usd = 0.0
    grand_total_tokens = 0

    with open(path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            rec = json.loads(line)
            for bucket, key in ((by_label, rec["label"]), (by_provider, rec["provider"])):
                d = bucket.setdefault(key, {"calls": 0, "prompt": 0, "completion": 0, "reasoning": 0, "total": 0, "usd": 0.0})
                d["calls"] += 1
                d["prompt"] += rec.get("prompt_tokens") or 0
                d["completion"] += rec.get("completion_tokens") or 0
                d["reasoning"] += rec.get("reasoning_tokens") or 0
                d["total"] += rec.get("total_tokens") or 0
                d["usd"] += rec.get("usd") or 0.0
            grand_total_usd += rec.get("usd") or 0.0
            grand_total_tokens += rec.get("total_tokens") or 0

    print("=== Usage by call type (label) ===")
    for label, d in sorted(by_label.items(), key=lambda kv: -kv[1]["total"]):
        print(f"{label:25s} calls={d['calls']:3d} prompt={d['prompt']:7d} completion={d['completion']:6d} "
              f"reasoning={d['reasoning']:6d} total={d['total']:8d} usd=${d['usd']:.6f}")
    print("\n=== Usage by provider ===")
    for provider, d in sorted(by_provider.items(), key=lambda kv: -kv[1]["total"]):
        print(f"{provider:12s} calls={d['calls']:3d} prompt={d['prompt']:7d} completion={d['completion']:6d} "
              f"reasoning={d['reasoning']:6d} total={d['total']:8d} usd=${d['usd']:.6f}")
    print(f"\nGRAND TOTAL: {grand_total_tokens} tokens, ${grand_total_usd:.6f}")

    return {"by_label": by_label, "by_provider": by_provider,
            "grand_total_tokens": grand_total_tokens, "grand_total_usd": grand_total_usd}


class _SourceErrorCapture(logging.Handler):
    """
    Temporarily attached to `log` for the duration of ONE source's
    run() call (see main()) to collect every ERROR-level message it logs
    — a real fetch failure, an HTTP error, an unhandled exception — into
    a plain list, with no changes needed to any individual scrape_*/
    process_*_item function. Added 2026-09-03 per user request: a run
    that silently logs "[scio] Failed to fetch ..." to a file nobody's
    watching (console output is capped at WARNING by default — see
    _console_handler) shouldn't just end with a cost/time summary as if
    everything went fine. Deliberately only ERROR+ (not INFO's routine
    "0 new items"/"No relevant exchanges" — those are healthy outcomes,
    not failures, and would drown out real signal if included).
    """
    def __init__(self):
        super().__init__(level=logging.ERROR)
        self.messages: list[str] = []

    def emit(self, record: logging.LogRecord) -> None:
        self.messages.append(record.getMessage())


def _summarize_run_usage(since_ts: str) -> tuple[int, int, float]:
    """
    Cost/token summary for JUST this run, not the whole project's history
    — USAGE_LOG_PATH is a running log across every run ever made, so this
    filters to records timestamped at or after `since_ts` (the calling
    run's own start time) rather than reusing summarize_usage_log()'s
    all-time aggregation. Added 2026-09-02 so a normal run shows its own
    cost automatically instead of requiring a separate manual command
    (see main()'s final summary print). Returns (llm_tokens, x_api_reads,
    total_usd) — X's pay-per-read charges are counted separately from LLM
    tokens since `total_tokens` on an X record is a read count, not an
    actual token count (see _log_x_cost). Never raises — a summary glitch
    at the very end of an otherwise-successful run shouldn't look like the
    run itself failed.
    """
    llm_tokens = 0
    x_reads = 0
    total_usd = 0.0
    try:
        if not Path(USAGE_LOG_PATH).exists():
            return 0, 0, 0.0
        with open(USAGE_LOG_PATH, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if rec.get("ts", "") < since_ts:
                    continue
                total_usd += rec.get("usd") or 0.0
                if rec.get("provider") == "X API":
                    x_reads += rec.get("total_tokens") or 0
                else:
                    llm_tokens += rec.get("total_tokens") or 0
    except Exception as exc:
        log.warning(f"Failed to summarize this run's usage: {exc}")
    return llm_tokens, x_reads, total_usd


def _format_duration(seconds: float) -> str:
    """'47s' / '16m 7s' / '1h 04m' — whichever unit fits, for the final
    run summary. Never needs to handle negative input (always measures a
    just-elapsed time.monotonic() delta)."""
    seconds = int(seconds)
    if seconds < 60:
        return f"{seconds}s"
    minutes, secs = divmod(seconds, 60)
    if minutes < 60:
        return f"{minutes}m {secs}s"
    hours, minutes = divmod(minutes, 60)
    return f"{hours}h {minutes:02d}m"


# GROQ_MODEL (openai/gpt-oss-120b)'s account tier caps at 8000 tokens/minute
# (TPM) — prompt + reserved max_tokens both count against it, checked
# BEFORE the call runs. A large translation chunk (this content runs
# multi-thousand-word position papers some days) plus a generous
# max_tokens can exceed that on its own, 413ing outright ("Request too
# large ... TPM: Limit 8000, Requested 8143") — found live via backtest.py,
# 2026-08-05, immediately after raising max_tokens to fix the reasoning-
# token-exhaustion bug (see _GROQ_MAX_TOKENS's other comment) — fixing one
# limit surfaced the other. _call_groq_with_retry below truncates harder
# and retries once on a 413 instead of just failing the item outright.
_GROQ_MAX_TOKENS = 2000        # headroom for reasoning + a real answer, without alone blowing the 8000 TPM cap
_GROQ_PROMPT_CHARS = 4000      # leaves room under 8000 TPM for _GROQ_MAX_TOKENS's reservation
_GROQ_PROMPT_CHARS_RETRY = 1500  # fallback truncation if even that 413s


def _call_groq_with_retry(client: Groq, prompt: str, suffix: str = ""):
    """
    Shared 413-aware call: truncate, call, and on a 413 truncate harder and
    retry once before giving up. `suffix` (e.g. _call_groq_json's "respond
    with this JSON schema" instructions) is appended AFTER truncating
    `prompt`, not before — truncating the whole already-combined string
    from the front risked cutting off exactly the instruction that tells
    the model to answer in JSON at all.
    """
    for max_chars in (_GROQ_PROMPT_CHARS, _GROQ_PROMPT_CHARS_RETRY):
        budget = max_chars - len(suffix)
        p = prompt if len(prompt) <= budget else prompt[:budget] + "\n\n[content truncated]"
        p += suffix
        try:
            return client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": p}],
                # GROQ_MODEL is a reasoning model that spends completion
                # tokens on an invisible chain-of-thought before the
                # visible answer — confirmed live: a one-sentence
                # translation burned 290 of a 600-token budget on
                # reasoning alone, hit the cap (finish_reason="length")
                # before emitting any answer, and came back as an EMPTY
                # string that then silently propagated (e.g. an empty
                # "translation" with no error) instead of raising — found
                # via backtest.py, 2026-08-05, on a real already-covered
                # entry (He Lifeng's video call with Secretary Bessent)
                # that a keyword/CJK check correctly flagged as worth
                # translating but then got nothing back. reasoning_effort=
                # "low" cuts that overhead for what are fundamentally
                # straightforward transform tasks (translate/extract/
                # summarize) that don't need heavy deliberation. Re-check
                # both if GROQ_MODEL changes to a non-reasoning model —
                # reasoning_effort will likely error on one.
                max_tokens=_GROQ_MAX_TOKENS,
                reasoning_effort="low",
            )
        except Exception as exc:
            if "413" in str(exc) and max_chars == _GROQ_PROMPT_CHARS:
                log.warning(f"Groq 413'd at {max_chars} chars — retrying at {_GROQ_PROMPT_CHARS_RETRY}")
                continue
            raise


def _log_groq_usage(label: str, completion) -> None:
    """Shared usage-log extraction for both _call_groq and _call_groq_json
    — completion.usage is an OpenAI-SDK-shaped object; reasoning_tokens
    (see the reasoning-token-exhaustion bug) lives one level down in
    completion_tokens_details."""
    usage = getattr(completion, "usage", None)
    if usage is None:
        return
    details = getattr(usage, "completion_tokens_details", None)
    reasoning = getattr(details, "reasoning_tokens", None) if details is not None else None
    _log_usage(
        "Groq", label,
        getattr(usage, "prompt_tokens", None),
        getattr(usage, "completion_tokens", None),
        getattr(usage, "total_tokens", None),
        reasoning,
    )


def _call_groq(prompt: str, label: str = "") -> str:
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("Gemini rate-limited and GROQ_API_KEY is not set")
    client = Groq(api_key=api_key)
    completion = _call_groq_with_retry(client, prompt)
    _log_groq_usage(label, completion)
    text = (completion.choices[0].message.content or "").strip()
    if not text:
        raise RuntimeError(
            f"Groq returned an empty response (finish_reason="
            f"{completion.choices[0].finish_reason!r}) — likely reasoning-"
            f"token exhaustion; do not silently treat this as a valid answer"
        )
    return text


# Gemini cooldown tracking — added 2026-09-01, a real speed fix with no
# accuracy tradeoff. Every call_llm/call_llm_json unconditionally slept
# GEMINI_SLEEP (16s) BEFORE even trying Gemini once — including calls made
# while Gemini's daily quota was already known-exhausted from an earlier
# 429 in the very same run, which then immediately fails over to the
# fallback chain anyway. Across a real ~150-call run that's 40+ minutes of
# pure sleep spent "confirming" something already learned minutes earlier.
# Once a 429 actually happens, skip the sleep AND the attempt entirely —
# straight to the fallback chain — for a cooldown window, instead of
# re-discovering the same exhaustion on every subsequent call. Rechecked
# periodically rather than disabled outright for the rest of the run, in
# case it was a transient RPM blip rather than the daily RPD cap (RPM
# clears in seconds; GEMINI_SLEEP already paces correctly for that case
# alone, so a 429 despite proper pacing is far more likely the daily cap,
# but this doesn't assume that — it just checks back every 2 minutes).
_GEMINI_COOLDOWN_SECONDS = 120
_gemini_retry_after: float = 0.0  # time.monotonic() timestamp; 0 = never rate-limited yet


def _gemini_on_cooldown() -> bool:
    return time.monotonic() < _gemini_retry_after


def _start_gemini_cooldown() -> None:
    global _gemini_retry_after
    _gemini_retry_after = time.monotonic() + _GEMINI_COOLDOWN_SECONDS


def call_llm(client: genai.Client, prompt: str, retries: int = 2, label: str = "") -> str:
    """Call Gemini with sleep; auto-fall back through _fallback_chain
    (Groq -> OpenRouter -> Cerebras, whichever have keys configured) on 429.
    `label` identifies the calling code path (e.g. "classify_relevance") for
    the [usage] token-count log — see _log_usage(). Skips straight to the
    fallback chain, no sleep, while Gemini is on cooldown from a recent
    429 — see _gemini_on_cooldown()'s docstring above."""
    if _gemini_on_cooldown():
        return _call_fallback_chain(prompt, label=label)
    time.sleep(GEMINI_SLEEP)
    for attempt in range(retries + 1):
        try:
            resp = client.models.generate_content(
                model=GEMINI_MODEL, contents=prompt,
                config=types.GenerateContentConfig(thinking_config=_GEMINI_THINKING_CONFIG),
            )
            usage = getattr(resp, "usage_metadata", None)
            if usage is not None:
                _log_usage(
                    "Gemini", label,
                    getattr(usage, "prompt_token_count", None),
                    getattr(usage, "candidates_token_count", None),
                    getattr(usage, "total_token_count", None),
                    getattr(usage, "thoughts_token_count", None),
                )
            return resp.text.strip()
        except Exception as exc:
            err = str(exc).lower()
            if "429" in err or "quota" in err or "resource_exhausted" in err:
                log.warning(f"Gemini rate-limited — falling back, cooling down {_GEMINI_COOLDOWN_SECONDS}s")
                _start_gemini_cooldown()
                return _call_fallback_chain(prompt, label=label)
            if attempt < retries:
                time.sleep(8 * (attempt + 1))
                continue
            raise


# Narrow, China-SPECIFIC pattern used only to pick where to center a
# relevance snippet — deliberately much narrower than
# US_SOURCE_RELEVANCE_KEYWORDS, which also contains generic terms
# (sanction, tariff, currency manipulation...) that show up early in
# plenty of documents with nothing to do with China (an Iran-sanctions
# release opens with "OFAC Sanctions..." at character 19, matching
# "sanction" long before the actual Chinese-company mention at character
# ~4300 — anchoring on that early, generic hit defeated the whole point of
# windowing). Found live via backtest.py, 2026-08-05.
_CHINA_MENTION_RE = re.compile(r"China\w*|Chinese|Beijing|Xi Jinping", re.IGNORECASE)


def _relevance_snippet(text: str, keyword_hit: re.Match, max_chars: int = 3000) -> str:
    """
    Build the snippet actually sent to the LLM relevance call, centered on
    wherever a real China-specific mention actually is — not just the
    first `max_chars` characters, and not wherever the (much broader,
    partly generic) US_SOURCE_RELEVANCE_KEYWORDS pattern happened to match
    first. Found live (backtest.py, 2026-08-05) that a real, already-
    covered Treasury sanctions release only names the actual Chinese/Hong-
    Kong-based shipping companies at character ~4300 (deep into a long
    press release that opens with generic Iran-sanctions boilerplate) — a
    naive prefix cutoff meant the LLM call never saw the part that
    actually made it relevant, even after the keyword pre-filter itself
    was fixed to scan the full text. Falls back to `keyword_hit` (e.g. a
    Taiwan/Hong Kong/named-official mention with no literal "China") only
    if no direct China mention exists at all.

    `max_chars` defaults to 3000, not something larger — found live that a
    5000-char snippet here, plus a couple hundred chars of instructions,
    could exceed `_call_groq`'s own 4000-char prompt cap (added for the
    TPM-limit fix) once it fell back to Groq, which then re-truncated the
    ALREADY-correctly-windowed snippet from the front and cut the
    China-relevant part right back out — two truncation layers stacking
    instead of cooperating. Keep this comfortably under
    `_GROQ_PROMPT_CHARS` including instruction overhead; do not raise it
    without checking that budget.
    """
    anchor = _CHINA_MENTION_RE.search(text) or keyword_hit
    if anchor.start() < max_chars:
        return text[:max_chars]
    lead  = text[:500]
    start = max(0, anchor.start() - 1000)
    window = text[start : start + max_chars - len(lead) - 10]
    return lead + "\n...\n" + window


def classify_relevance(
    model: genai.Client, text: str, chinese_origin: bool = False,
) -> tuple[bool, str]:
    """
    Strict gate: only pass items that explicitly and substantively involve
    China/PRC/Taiwan/Hong Kong or a named Chinese entity — not just adjacent
    foreign-policy topics (critical minerals, generic multilateral
    cooperation, an unrelated country's bilateral news, etc.) that happen to
    share a keyword with things China also does. Tightened 2026-08-04 after
    a live run let through items like a Cook Islands anniversary greeting
    and a US-Italy critical-minerals announcement that never mention China.

    Free keyword pre-filter first: most items on any given day's list page
    have nothing to do with China at all, and don't need an LLM call to
    figure that out. Only spend one if a keyword actually shows up — added
    2026-08-04 after burning through both providers' free daily quota in an
    afternoon of testing; skipping the obvious no's for free instead of
    paying for every single item matters a lot in practice.

    `chinese_origin=True` (added 2026-09-03, for SCIO): the "does this
    involve China" question above is meaningless for a CHINESE-government
    source — a State Council Information Office release trivially
    "explicitly and substantively involves China" almost by definition,
    regardless of subject. Live-caught the LLM being genuinely
    inconsistent about this: it correctly said NO to a purely domestic
    healthcare-access report and a housing-policy circular, but YES to an
    AI-medical-imaging ethics guideline and a bare manufacturing-PMI
    reading — all four are equally domestic-only, zero US mention
    anywhere in any of them. `chinese_origin=True` swaps in a prompt
    asking the actually-correct question for this direction: does this
    substantively involve THE US, US-China relations/trade tension, or
    (per a real, confirmed ground-truth SCIO entry with NO explicit US
    mention at all) an implicit US-driven geopolitical narrative like
    "China squeeze," "decoupling," or "overcapacity" — as opposed to
    ordinary domestic policy/regulation/statistics that only happens to
    be Chinese. Keeps the same keyword pre-filter (still includes bare
    "China"/"Beijing"/"Xi Jinping" — a Chinese source will always hit it,
    so this doesn't change the free/no-LLM-call fast path, only the
    judgment once an LLM call happens).
    """
    keyword_hit = US_SOURCE_RELEVANCE_KEYWORDS.search(text)
    if not keyword_hit:
        return False, "Keyword pre-filter: no US-China-relevant terms found — skipped LLM call."

    snippet = _relevance_snippet(text, keyword_hit)

    if chinese_origin:
        prompt = (
            "You are screening a CHINESE-government press release for a "
            "US-China relations tracker. Every release from this source is "
            "trivially 'about China' — that is NOT the question. Reply YES "
            "only if this release substantively involves the United States, "
            "US-China relations, trade tensions between the two, a named US "
            "person/entity, OR Taiwan/Hong Kong SPECIFICALLY IN CONNECTION "
            "WITH THE US (e.g. US arms sales to Taiwan, a US policy action "
            "toward Taiwan/HK, a US official's statement on either) — OR is "
            "a response to a US-driven geopolitical/economic narrative even "
            "without naming the US explicitly (e.g. rebutting 'China "
            "squeeze,' 'decoupling,' 'de-risking,' or 'overcapacity' "
            "framing, or addressing US tariff/export-control actions in "
            "context). Reply NO for: (1) ordinary domestic Chinese policy, "
            "regulation, or economic statistics with no such angle — a "
            "manufacturing PMI reading, a healthcare-access report, a "
            "housing-policy circular, an industry ethics guideline, etc. — "
            "even though it's genuinely about China's government doing "
            "something; (2) routine China-Taiwan cross-strait or "
            "territorial/sovereignty content with NO connection to the US "
            "at all — a coast guard patrol notice, a marine survey, a "
            "routine assertion of sovereignty over Taiwan/the South China "
            "Sea/disputed islands — these are genuine PRC-Taiwan or "
            "PRC-neighbor matters, not US-China relations, even though "
            "Taiwan is often discussed in a US-China context ELSEWHERE; "
            "bare mention of the word 'Taiwan' or 'Hong Kong' is not "
            "enough on its own. Ask yourself: would a human editor writing "
            "a US-CHINA relations tracker (not a China-news tracker or a "
            "cross-strait-relations tracker) include this, or is it just "
            "routine Chinese domestic or regional news that happens to "
            "come from a government source?\n\n"
            "Reply with exactly one word, YES or NO, on the first line, "
            "then one sentence of justification on the next line.\n\n" + snippet
        )
    else:
        prompt = (
            "You are screening press releases for a US-China relations tracker. "
            "Reply YES only if this text explicitly and SUBSTANTIVELY involves "
            "China, the People's Republic of China, Chinese government entities "
            "or officials, Chinese companies, Taiwan, or Hong Kong — or a direct "
            "US-China policy action (tariffs on Chinese goods, export controls, "
            "sanctions on Chinese entities, trade negotiations with China, etc.). "
            "A passing mention of 'global cooperation', general critical-minerals "
            "or supply-chain policy that does not name China or a Chinese entity, "
            "or an unrelated country's bilateral news, is NOT relevant — reply NO. "
            "Two more specific NO cases, both real false positives caught live: "
            "(1) China is named only ONCE, as a rhetorical comparison inside a "
            "story that is actually about a different country's bilateral "
            "relationship with the US (e.g., 'Country X, like China, chose "
            "retaliation over negotiation' inside a release that is otherwise "
            "entirely about Country X) — that single aside does not make the "
            "release substantively about China; reply NO. (2) China/Taiwan/Hong "
            "Kong appear only as ONE ROW in a statistical table or ranked list "
            "covering many countries (e.g., a country-by-country investment, "
            "trade, or holdings table), with no sentence of actual discussion "
            "about China specifically — reply NO. (3) Taiwan or Hong Kong is "
            "named only as one of several listed countries/participants/members "
            "in an unrelated multilateral program, coalition, or supply-chain "
            "initiative (e.g. a workforce-training program whose listed "
            "participants happen to include Taiwan alongside Japan, South "
            "Korea, India, etc.), with no discussion of Taiwan's political "
            "status, cross-strait tensions, or a China-specific angle — reply "
            "NO; a real, live false positive of exactly this shape got approved "
            "purely because 'Taiwan' appeared once in a list of program "
            "participants having nothing to do with China. In all three cases, "
            "would a human editor writing a US-China relations tracker actually "
            "consider this release to be ABOUT US-China relations, or does it "
            "just happen to contain the word 'China'/'Taiwan'/'Hong Kong' "
            "somewhere? If the latter, reply NO.\n\n"
            "Reply with exactly one word, YES or NO, on the first line, then one "
            "sentence of justification on the next line.\n\n" + snippet
        )

    result = call_llm(model, prompt, label="classify_relevance")
    first_line = result.strip().splitlines()[0].strip().upper() if result.strip() else ""
    return first_line.startswith("YES"), result


_TRANSLATION_PREAMBLE_RE = re.compile(
    r"^(here('?s| is)? (the |a )?translation|"
    r"sure[,!]?|certainly[,!]?|of course[,!]?|"
    r"translated (text|version)|below is)\b.*$",
    re.IGNORECASE,
)


def translate_to_english(model: genai.Client, text: str) -> str:
    """
    Translate Chinese text to English, preserving Q&A structure. Strips a
    leading meta-commentary line if the model adds one despite being told
    not to (e.g. "Here is the translation of the Chinese text to English:")
    — found 2026-08-04 when that exact line ended up written into the
    tracker doc as if it were a real CONT paragraph, because the regex
    Q&A-parser fallback has no way to tell it apart from real content.
    """
    # 10,000, not 5,500 — every real call site caps its input at 7000
    # chars before ever reaching this function (see the 4 process_*_item
    # call sites), so a 5500-char chunk size split almost every real
    # Chinese document translated by this pipeline into 2 chunks for no
    # reason: gemini-2.5-flash's real context window is enormous (~1M
    # tokens) relative to 7000 characters (~2000 tokens), so there was
    # never a technical need to chunk this small — found reviewing this
    # function, 2026-09-01, no NOTES.md entry ever documented a reason for
    # 5500 specifically. Splitting mid-document isn't just extra cost —
    # each chunk is translated independently and joined with "\n\n", so a
    # split landing right after an orphan "问："/"答：" label (with its
    # actual content starting on the far side of the boundary, in a
    # SEPARATE chunk) would silently orphan that label from its content
    # across the join, the same failure shape as bugs #13/#14, just at a
    # different layer. Raised comfortably above every current caller's max
    # input so none of them ever actually chunk in practice; kept as a
    # loop (not removed) so a future caller passing something longer still
    # degrades safely instead of hitting an unbounded single request.
    parts = []
    for i in range(0, len(text), 10000):
        chunk = text[i : i + 10000]
        glossary_hint = "".join(
            f"'{zh}' -> '{en}'; " for zh, en in KNOWN_NAME_ROMANIZATIONS.items() if zh in chunk
        )
        glossary_line = f"Use these exact spellings for these names: {glossary_hint}\n" if glossary_hint else ""
        translated = call_llm(
            model,
            "Translate the following Chinese text to English exactly. "
            "Preserve the Q&A structure. Do not summarize or paraphrase. "
            "Romanize speaker names (e.g. '林剑：' → 'Lin Jian:'). "
            + glossary_line +
            "Output ONLY the translation itself — no preamble like 'Here is "
            "the translation:', no commentary, nothing but the translated text.\n\n"
            + chunk,
            label="translate",
        )
        lines = translated.split("\n", 1)
        if lines and _TRANSLATION_PREAMBLE_RE.match(lines[0].strip()):
            translated = lines[1] if len(lines) > 1 else ""
        parts.append(translated)
    return "\n\n".join(parts)


class SummaryResponse(BaseModel):
    summary: str
    anchor: str


def generate_summary(
    model: genai.Client, text: str, source_name: str, extra_instruction: str = "",
) -> tuple[str, str]:
    """
    Write a tracker-style summary sentence matching the existing doc style, plus
    an "anchor" — the single verb/short-phrase within that sentence which should
    be hyperlinked to the source URL, matching how the past trackers embed links
    (e.g. "...Lin Jian **addressed** reporters' questions on..." with "addressed"
    hyperlinked to the FMPRC page).

    Returns (summary, anchor). `anchor` is guaranteed to be a verbatim substring
    of `summary` (falls back to the first word if the model's anchor doesn't
    match, so callers can always find something to hyperlink).

    `extra_instruction`: appended verbatim to the prompt — used by
    get_summary_and_anchor()'s retry path when a first attempt named an
    official the source text never actually mentions (see
    _hallucinated_officials()'s docstring for the real case that
    justified this).
    """
    # Same windowing as classify_relevance/extract_key_paragraphs — NOT a
    # plain prefix cutoff. Found live (backtest.py, 2026-08-05) that a real
    # Treasury sanctions release names the actual Chinese/Hong-Kong-based
    # companies at character ~4300; a flat text[:3500] here meant this call
    # never saw that part at all, so the generated summary was accurate but
    # framed entirely around the (dominant, but not tracker-relevant) Iran-
    # sanctions angle instead of the China angle that's the actual reason
    # the entry belongs in a US-China tracker. Falls back to a plain prefix
    # when there's no keyword hit to anchor on (e.g. already-filtered Q&A
    # work_text, which is short and pre-relevant either way).
    keyword_hit = US_SOURCE_RELEVANCE_KEYWORDS.search(text)
    snippet = _relevance_snippet(text, keyword_hit, max_chars=3500) if keyword_hit else text[:3500]
    prompt = (
        "Write ONE sentence for a US-China policy tracker, in the style already "
        "used in this tracker. State WHO (full name + title) did WHAT, and name "
        "the specific topic(s). Be concrete, use past tense, no fluff. Target "
        "roughly 15-25 words — real entries in this tracker run that short even "
        "when the topic is substantial; if a topic list is getting long, cut it "
        "to the 2-3 most important items rather than naming every single one. "
        "Do not start with 'This', 'The entry', or 'I'. If this document covers "
        "multiple topics, lead with whichever one is actually about China/"
        "Chinese entities/Taiwan/Hong Kong — that is the reason this belongs "
        "in a US-China tracker at all, even if it is not the document's main "
        "subject (e.g. a sanctions action framed around a third country that "
        "also names Chinese companies should be summarized around the "
        "Chinese companies, not the third country). Do NOT invent or guess a "
        "specific person's name or title if the text does not clearly name "
        "who wrote, signed, or announced it — attribute it to the "
        "institution instead (e.g. 'The White House released a National "
        "Security Strategy...', not a guessed official's name), even if a "
        "person with that kind of title would normally be associated with "
        "this type of document. Chinese officials often hold several "
        "simultaneous titles (Communist Party rank, a Central Commission "
        "office, a state ministry) — the text may lead with any of them. "
        "ALWAYS use the single most internationally-recognized functional "
        "title for these specific, frequently-appearing officials, "
        "regardless of which title the source text happens to lead with: "
        "Wang Yi is always 'Foreign Minister Wang Yi' (never 'Member of "
        "the Political Bureau...', never 'Director of the Office of the "
        "Central Commission for Foreign Affairs...'); Li Qiang is always "
        "'Premier Li Qiang'; Han Zheng is always 'Vice President Han "
        "Zheng'. For any OTHER Chinese official not on this list, drop "
        "Communist Party/Central Commission structure the same way and "
        "keep only their actual government/ministry role.\n\n"
        "Match this style:\n"
        "'Foreign Ministry Spokesperson Lin Jian addressed reporters' questions on "
        "U.S. tariffs on Chinese goods, Taiwan arms sales, and critical mineral "
        "export controls at the regular daily press conference.'\n"
        "'Treasury Secretary Scott Bessent testified before the House Financial "
        "Services Committee, discussing Chinese IPOs, rare earth leverage, and "
        "digital currency competition with China.'\n\n"
        "Also return 'anchor': the single action verb (or short verb phrase, 1-2 "
        "words, e.g. 'addressed', 'held', 'released', 'issued', 'testified', "
        "'answered') from your summary sentence that best represents the act of "
        "speaking/publishing — this word will be hyperlinked to the source. It "
        "MUST be copied verbatim from the 'summary' string.\n\n"
        + (f"{extra_instruction}\n\n" if extra_instruction else "")
        + f"Source: {source_name}\n\nContent:\n{snippet}"
    )
    try:
        result = call_llm_json(model, prompt, SummaryResponse, label="generate_summary")
        summary = (result.get("summary") or "").strip()
        anchor = (result.get("anchor") or "").strip()
        if summary and anchor and anchor in summary:
            return summary, anchor
        if summary:
            return summary, summary.split(" ", 1)[0]
    except Exception as exc:
        log.warning(f"Structured summary call failed, falling back to plain text: {exc}")

    # Fallback: plain-text summary, anchor = first verb-ish word after the subject.
    summary = call_llm(
        model,
        "Write ONE concise sentence (roughly 15-25 words) for a US-China policy "
        "tracker. State WHO (with their full name and title) said or did WHAT, "
        "and name the 2-3 most important topics covered — cut a long topic list "
        "down rather than naming everything. Be concrete. Use past tense. No "
        "fluff. Do not start with 'This', 'The entry', or 'I'. Chinese "
        "officials often hold several simultaneous titles — ALWAYS use the "
        "single most internationally-recognized one for these specific "
        "officials regardless of which title the source leads with: Wang "
        "Yi is always 'Foreign Minister Wang Yi', Li Qiang is always "
        "'Premier Li Qiang', Han Zheng is always 'Vice President Han "
        "Zheng'. For any other Chinese official, drop Communist Party/"
        "Central Commission structure the same way and keep only their "
        "actual government/ministry role.\n\n"
        f"Source: {source_name}\n\nContent:\n{snippet}",
        label="generate_summary",
    )
    return summary, summary.split(" ", 1)[0] if summary else ""


# Disabled 2026-09-01 per user request (less LLM dependency/cost/failure-
# surface while the rest of the pipeline was being tightened up); RE-
# ENABLED 2026-09-02 per user request, after directly checking the real
# past trackers first (not just re-reading this file's own docstrings) to
# confirm the existing prompt's style/anchor convention still matches:
# confirmed a single verb/short phrase hyperlinked within one concise
# sentence (13-19 words in real examples — "Lin Jian addressed reporters'
# questions on...", "Bessent testified before...") is exactly the real
# convention, and generate_summary()'s prompt (below) already produces
# that shape essentially verbatim. No prompt changes needed. See
# NOTES.md for the real cost/time estimate given before re-enabling
# (~$0.0006/entry, ~20-30s/entry from GEMINI_SLEEP pacing — now much less
# than that estimate implied, since GEMINI_SLEEP is down to 0.5s today).
ENABLE_LLM_SUMMARY = True

# Real officials this project has needed a "don't hallucinate a name the
# source doesn't actually mention" check for — checked as a plain
# surname substring against the ORIGINAL source text. Added 2026-09-03
# after a live run's generate_summary() invented "Treasury Secretary
# Janet Yellen" for a real G20 Finance Ministers statement that names NO
# individual official anywhere in its ~18,500 characters — Yellen hasn't
# been Treasury Secretary since January 2025 (Bessent is, and even HE
# isn't named in the source either). The prompt already has an explicit
# "do not invent a name" instruction (see generate_summary's docstring
# history) — this model still did it anyway, almost certainly pattern-
# matching "G20 Finance Ministers meeting" to whichever name is most
# strongly associated with that event type in its training data,
# overriding the instruction. A second, PROGRAMMATIC check catches what
# telling it nicely didn't.
_HALLUCINATION_CHECK_SURNAMES = [
    "Bessent", "Greer", "Lutnick", "Rubio", "Vance", "Navarro", "Yellen",
    "Perdue", "Hegseth", "Gabbard", "Ratcliffe", "Leavitt", "Miller",
    "Biden", "Sullivan", "Blinken",
]
# "Trump" deliberately excluded: he's named constantly as an adjective
# ("Trump administration," "Trump tariffs") in documents that never
# actually name HIM as the one who did the specific act being
# summarized — including him here would fire on nearly every US-origin
# document and defeat the whole point of this check.


def _hallucinated_officials(summary: str, source_text: str) -> list[str]:
    """
    Returns the (possibly empty) list of _HALLUCINATION_CHECK_SURNAMES
    entries that appear in `summary` but NOT anywhere in `source_text` —
    a plain substring check, not real NLP, so it can't distinguish "the
    text discusses Bessent's earlier remarks" from "Bessent personally
    announced this specific thing." That's fine: the one failure mode
    this exists to catch is a name in the summary that isn't in the
    source text AT ALL, which is unambiguous either way.

    Case-INSENSITIVE on purpose — found live, 2026-09-03: a real State
    Department interview transcript labels its speaker "SECRETARY
    RUBIO:" (all caps, a common transcript convention), which a
    case-sensitive `"Rubio" in source_text` check doesn't match at all —
    this falsely flagged a real, correctly-named summary as a
    hallucination and triggered a needless (and in that case actively
    worse — the retry ended up LESS specific, "The State Department
    addressed..." instead of naming Rubio) regeneration.
    """
    source_lower = source_text.lower()
    return [
        name for name in _HALLUCINATION_CHECK_SURNAMES
        if name in summary and name.lower() not in source_lower
    ]


def get_summary_and_anchor(
    model: genai.Client, text: str, source_name: str, url: str,
) -> tuple[str, str]:
    """
    Single choke point every call site uses instead of calling
    generate_summary() directly — see ENABLE_LLM_SUMMARY above. When
    disabled, the "summary" is just the bare URL: add_summary_para() then
    renders the whole line as the hyperlink (summary == anchor == url), so
    entries still get a date heading + link + body (paragraphs/exchanges)
    — just no prose summary sentence for now.

    Retries once, with an explicit correction, if the first attempt named
    an official _hallucinated_officials() can't find in the source text —
    see that function's docstring for the real case this exists for. If
    the retry STILL names a hallucinated official, keeps the retry's
    result (better odds than the original) and flags the URL for human
    review rather than silently shipping a specific, wrong factual claim.
    """
    if not ENABLE_LLM_SUMMARY:
        return url, url

    summary, anchor = generate_summary(model, text, source_name)
    bad_names = _hallucinated_officials(summary, text)
    if not bad_names:
        return summary, anchor

    log.warning(f"[get_summary_and_anchor] Named {bad_names} but source "
                f"text doesn't mention them — retrying: {url}")
    retry_summary, retry_anchor = generate_summary(
        model, text, source_name,
        extra_instruction=(
            f"IMPORTANT: your previous attempt incorrectly named "
            f"{', '.join(bad_names)} — the source text does NOT mention "
            f"{'this person' if len(bad_names) == 1 else 'these people'} "
            f"anywhere. Do not use that name. If the source text doesn't "
            f"clearly name who did this, attribute it to the institution "
            f"instead."
        ),
    )
    if not _hallucinated_officials(retry_summary, text):
        return retry_summary, retry_anchor

    flag_for_review(url, summary[:80],
                     f"generate_summary named {bad_names}, not found anywhere "
                     f"in the source text — possible hallucination, kept after retry")
    return retry_summary, retry_anchor


_REFUSAL_RE = re.compile(
    r"^(i am sorry|i'm sorry|there (is|are) no|this text does not|"
    r"does not (contain|mention|explicitly)|cannot (extract|find)|"
    r"no (relevant|china-related) (paragraphs|content|text)|"
    r"unable to (extract|find)|the (provided|given) text)",
    re.IGNORECASE,
)


def extract_key_paragraphs(
    model: genai.Client, text: str, n: int = 4,
) -> list[str]:
    """
    Extract the most important verbatim paragraphs from a release/statement.
    Returns a list of paragraph strings, or [] if none were found — callers
    MUST treat an empty list as "not actually relevant, skip this entry",
    not fall back to writing something else. (Previously a "no relevant
    paragraphs" refusal from the model — e.g. "I am sorry, but the provided
    text does not contain..." — was itself written into the tracker doc as
    if it were body content; see NOTES.md, 2026-08-04.)

    Used to take a `general=True` mode that dropped the "must be about
    China/US-China relations" requirement, for MFA leadership sources on
    the theory that their real editorial bar was broader than that.
    Retired 2026-09-03: the one past-tracker example that justified it
    turned out to be a human coding error in the original tracker, and
    with that gone, `general=True` was actively unsafe — a real live run
    used it to approve a Wang Yi/India-border readout with zero US
    mentions anywhere, purely because a free keyword pre-filter upstream
    matched "Tibet" as a shared topic. MFA leadership now uses
    select_relevant_chinese_paragraphs() via finalize_release_item()'s
    `raw_zh_text` parameter instead, same as every other Chinese source.
    See NOTES.md, 2026-09-03, for the full story.
    """
    # Same windowing fix as classify_relevance's _relevance_snippet(): a
    # naive prefix cutoff can miss the actual relevant paragraph in a long
    # document that opens with boilerplate.
    keyword_hit = US_SOURCE_RELEVANCE_KEYWORDS.search(text)
    # NOT max_chars=5000 — see _relevance_snippet's docstring on why that
    # collided with _call_groq's own prompt-size cap once Groq is the one
    # actually serving the call. Use the same 3000-char default as
    # classify_relevance.
    snippet = _relevance_snippet(text, keyword_hit) if keyword_hit else text[:3000]
    prompt = (
        f"Extract the {n} most important verbatim paragraphs about China, US-China "
        f"relations, trade, or Taiwan from this text. "
        f"Return each paragraph separated by the delimiter '|||'. "
        f"Do not summarize or paraphrase — use exact text. "
        f"If NO paragraph in the text actually discusses China, US-China relations, "
        f"trade, or Taiwan, return exactly: NONE\n\n{snippet}"
    )
    result = call_llm(model, prompt, label="extract_key_paragraphs")
    if re.match(r"^\s*NONE\b", result, re.IGNORECASE):
        # Model sometimes adds trailing commentary, e.g. "NONE (since only
        # three relevant paragraphs exist)" — a bare startswith is more
        # robust than requiring the whole reply to be exactly "NONE".
        return []
    paras = [p.strip() for p in result.split("|||") if p.strip()]
    paras = [
        p for p in paras
        if len(p) > 20 and not _REFUSAL_RE.match(p) and not re.match(r"^NONE\b", p, re.IGNORECASE)
    ]
    return paras[:n]


def extract_main_text(html: str) -> str:
    """
    Get the actual article body text, not the whole page. Plain
    `BeautifulSoup(html).get_text()` was pulling in nav menus and — on sites
    like ustr.gov, whose article pages open with a "recent headlines" widget
    — several KB of OTHER articles' titles before the real body even starts.
    Since classify_relevance/extract_key_paragraphs/generate_summary all
    truncate to the first 2500-5000 chars, that junk was pushing the actual
    China-relevant content out of the window entirely, causing false
    negatives (a USTR forced-labor Section 301 release was rejected as
    irrelevant because character 2500 was still inside the headline list).
    Found + fixed 2026-08-04 during live testing — see NOTES.md.
    """
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()

    content = (
        soup.select_one("article")
        or soup.find("div", class_=re.compile(r"content|article|entry|post|body|main", re.I))
        or soup.find("div", id=re.compile(r"content|article|entry|post|body|main", re.I))
        or soup.find("main")
    )
    if content and len(content.get_text(strip=True)) > 80:
        return content.get_text(separator="\n", strip=True)
    return soup.get_text(separator="\n", strip=True)


def extract_pdf_text(pdf_bytes: bytes, max_pages: int = 60) -> str:
    """
    Real text extraction for a PDF response, via pdfplumber. Added
    2026-09-01 after finding that a White House National Security Strategy
    PDF — a real past-tracker entry, substantively about China (21 "China"
    mentions across 33 pages) — was being run through extract_main_text()
    as if it were HTML: BeautifulSoup parsing raw PDF bytes (which start
    with binary stream data like `%PDF-1.6...FlateDecode...`) produces
    unreadable garbage, so classify_relevance/extract_key_paragraphs
    correctly said "no relevant content" about text that was never
    readable in the first place — a false negative that looked exactly
    like a real judgment call. `max_pages` bounds extraction time/size for
    an unusually long PDF; 60 pages is generous for anything this tracker
    is likely to encounter (a National Security Strategy runs ~30).
    """
    import pdfplumber
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        pages = pdf.pages[:max_pages]
        return "\n\n".join(p.extract_text() or "" for p in pages)


def extract_text_from_response(resp: "httpx.Response") -> str:
    """
    Content-type-aware dispatcher: routes a PDF response through
    extract_pdf_text() (real text extraction) and everything else through
    extract_main_text() (HTML body isolation) — use this instead of
    calling extract_main_text(resp.text) directly wherever a source might
    link to a PDF (press releases occasionally do: fact sheets, strategy
    documents, reports). Checked by Content-Type header first (most
    reliable), falling back to a `.pdf` URL suffix for a server that
    mislabels it.
    """
    content_type = resp.headers.get("content-type", "")
    if "application/pdf" in content_type or str(resp.url).lower().endswith(".pdf"):
        return extract_pdf_text(resp.content)
    return extract_main_text(resp.text)


# ── HTTP client ───────────────────────────────────────────────────────────────

def make_client(verify_ssl: bool = True) -> httpx.Client:
    return httpx.Client(
        headers={"User-Agent": BROWSER_UA, "Accept-Language": "en-US,en;q=0.9"},
        follow_redirects=True,
        timeout=30.0,
        verify=verify_ssl,
    )


# Set by fetch() on any definitive HTTP error status, None on success or a
# non-HTTP failure (timeout, connection error) — a lightweight side channel
# so a caller that cares WHICH status code a None return meant can check
# it right after calling fetch(), without changing fetch()'s simple
# Optional[Response] contract for its many other callers. Added 2026-09-02
# for war.gov specifically — see scrape_wardept()'s use of this.
_LAST_FETCH_STATUS: int | None = None


def fetch(client: httpx.Client, url: str, retries: int = 3) -> httpx.Response | None:
    """
    `retries` used to only apply to low-level network exceptions (timeouts,
    connection errors) — ANY HTTP error status (403/404/500/521/...) hit
    `return None` on the very first attempt, silently ignoring `retries`
    entirely for what's actually the most common failure shape. Found
    reviewing this function, 2026-09-01 — Treasury explicitly passes
    `retries=4` at its call site, which did nothing for an HTTP error
    response. Fixed to distinguish: a 5xx is frequently a transient
    server-side blip and worth retrying with backoff, same as a network-
    level failure; a 4xx (403 Forbidden, 404 Not Found, ...) is left as an
    immediate bail, since those are essentially always persistent (wrong
    URL, blocked, gone) and retrying just wastes a request for nothing.
    """
    global _LAST_FETCH_STATUS
    _LAST_FETCH_STATUS = None
    for attempt in range(retries):
        try:
            time.sleep(REQUEST_SLEEP)
            resp = client.get(url)
            if resp.status_code == 429:
                wait = int(resp.headers.get("Retry-After", 30))
                log.warning(f"429 on {url} — waiting {wait}s")
                time.sleep(wait)
                continue
            resp.raise_for_status()
            return resp
        except httpx.HTTPStatusError as exc:
            status = exc.response.status_code
            _LAST_FETCH_STATUS = status
            log.warning(f"HTTP {status}: {url}")
            if status >= 500 and attempt < retries - 1:
                time.sleep(6 * (attempt + 1))
                continue
            return None
        except Exception as exc:
            log.warning(f"Request failed ({attempt+1}/{retries}): {url} — {exc}")
            if attempt < retries - 1:
                time.sleep(6 * (attempt + 1))
    return None


# ── Document writer ───────────────────────────────────────────────────────────

def _set_doc_defaults(doc: Document) -> None:
    """
    Apply document-wide defaults to the Normal style:
      - Times New Roman 12 pt
      - 1.15 line spacing
      - 8 pt space after paragraph
    All paragraphs added afterwards inherit these automatically.
    """
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after  = Pt(8)


# Valid XML 1.0 text content: tab/newline/CR plus most of the Unicode
# range -- everything else (most commonly a stray C0 control character)
# makes lxml raise a hard ValueError ("All strings must be XML
# compatible...") and crash the WHOLE run, not just skip that one piece
# of text. Found live, 2026-09-02: a real Aug 18-24 MFA leadership run
# crashed mid-flush on exactly this, inside generate_summary()'s LLM
# output -- the first real-world exercise of that path since
# ENABLE_LLM_SUMMARY was turned back on. Scraped page text or any other
# text this pipeline handles could just as easily carry one.
_XML_INVALID_CHARS_RE = re.compile(
    "[^\t\n\r\u0020-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]"
)


def _xml_safe(text: str) -> str:
    """Strip characters illegal in XML text content -- see
    _XML_INVALID_CHARS_RE's comment. Applied at the two actual text-
    insertion points (_run(), add_hyperlink()) rather than at every
    individual call site, so nothing that flows into either one can ever
    crash the whole run again over one bad character in one entry."""
    return _XML_INVALID_CHARS_RE.sub("", text)


# Only tags with a recognized HTML name, not a blanket "<[^>]+>" -- a
# blanket strip would also eat a genuine "<5%" or similar comparison in
# real economic/policy text, mistaking it for a tag. Found live,
# 2026-09-02: a real White House/Ford entry came out as literal
# '**Ford Motor Company** <a href="...">**announced**</a> **it will
# reshore...' in the finished doc -- confirmed by re-fetching the actual
# source page that the RAW extracted text has NONE of this markup at
# all, so extract_key_paragraphs()'s LLM call added it despite being
# told to preserve exact text, not something carried through from the
# page. generate_summary() or any other free-text LLM call could just as
# easily do the same thing.
_HTML_TAG_RE = re.compile(
    r"</?(?:a|b|i|u|strong|em|span|p|br|div|ul|ol|li|h[1-6])\b[^>]*>", re.IGNORECASE
)
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MD_BOLD_UNDERSCORE_RE = re.compile(r"__(.+?)__")


def _strip_stray_markup(text: str) -> str:
    """Strip HTML tags and markdown bold markers that occasionally leak
    into LLM-generated text despite being told to produce/preserve plain
    text -- see _HTML_TAG_RE's comment. Applied at the same choke points
    as _xml_safe() (_run(), add_hyperlink()) so this is caught regardless
    of which LLM call produced it, not just the one case actually seen."""
    text = _HTML_TAG_RE.sub("", text)
    text = _MD_BOLD_RE.sub(r"\1", text)
    text = _MD_BOLD_UNDERSCORE_RE.sub(r"\1", text)
    return text


def _run(paragraph, text: str, bold: bool = False, italic: bool = False):
    """Add a run with explicit TNR 12 so theme fonts can't override."""
    r = paragraph.add_run(_xml_safe(_strip_stray_markup(text)))
    r.bold        = bold
    r.italic      = italic
    r.font.name   = "Times New Roman"
    r.font.size   = Pt(12)
    return r


def add_hyperlink(paragraph, text: str, url: str, bold: bool = False, italic: bool = False):
    """
    Insert a real external hyperlink run into `paragraph`, styled to match the
    past trackers: Times New Roman 12pt, Google-Docs-blue (#1155CC), single
    underline. python-docx has no built-in hyperlink API, so this builds the
    OOXML <w:hyperlink> element by hand (the standard workaround).
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # half-points -> 12pt
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = _xml_safe(_strip_stray_markup(text))
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_summary_para(doc: Document, summary: str, url: str | None = None, anchor: str | None = None):
    """
    Write the plain-text summary line, hyperlinking `anchor` (a verbatim
    substring of `summary`) to `url` — matching how the past trackers embed
    the source link on a single verb inside the summary sentence. Falls back
    to a fully plain-text paragraph if no url/anchor is available or anchor
    isn't found in the summary.
    """
    p = doc.add_paragraph()
    if url and anchor and anchor in summary:
        before, after = summary.split(anchor, 1)
        if before:
            _run(p, before)
        add_hyperlink(p, anchor, url)
        if after:
            _run(p, after)
    else:
        _run(p, summary)
    return p


def _indented_para(doc: Document) -> object:
    """
    Return a new paragraph with a full paragraph-level left indent of 0.5 inches.
    Every line — not just the first — is indented because we set left_indent,
    NOT first_line_indent.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent       = Inches(0.5)
    pf.first_line_indent = Pt(0)          # no extra first-line indent on top
    return p


def get_or_create_doc(path: str) -> Document:
    p = Path(path)
    if p.exists():
        doc = Document(str(p))
    else:
        doc = Document()
        for para in list(doc.paragraphs):
            if not para.text.strip():
                para._element.getparent().remove(para._element)
    _set_doc_defaults(doc)
    return doc


def _fmt_date(date: datetime) -> str:
    """Format as 'Monday, June 17, 2026' — cross-platform."""
    return date.strftime("%A, %B ") + str(date.day) + ", " + str(date.year)


def add_date_heading(doc: Document, date: datetime) -> None:
    """Bold normal paragraph for the date, matching Part 3 of the existing tracker."""
    p = doc.add_paragraph()
    _run(p, _fmt_date(date), bold=True)


def add_qa_entry_body(
    doc: Document,
    summary: str,
    exchanges: list[dict],
    url: str | None = None,
    anchor: str | None = None,
) -> None:
    """
    Write a Q&A entry's summary + exchanges — NOT the date heading (see
    add_qa_entry / queue_entry+flush_pending_entries for that) and NOT a
    blank separator paragraph. Confirmed against the past trackers' raw XML
    that they never insert a blank paragraph between entries — spacing
    comes entirely from each paragraph's own space_after (8pt, the Normal
    style default set in _set_doc_defaults) — so adding one here would be
    an extra gap the source docs don't have.

    exchanges: list of dicts with keys:
        type    — "Q" | "A" | "CONT"
        speaker — outlet/spokesperson name (None for CONT)
        text    — paragraph text

    Q paragraphs:   outlet name bold+italic, ": text" italic only — full paragraph indented
    A paragraphs:   "Speaker:" bold, " text" plain              — full paragraph indented
    CONT paragraphs: plain continuation                         — full paragraph indented
    All lines in every exchange paragraph are indented (left_indent, not first_line_indent).

    `url`/`anchor`: if given, `anchor` (a verbatim substring of `summary`) is
    hyperlinked to `url` in the summary line, matching the past trackers'
    inline-source-link convention.
    """
    add_summary_para(doc, summary, url, anchor)  # summary — plain, no indent, link on `anchor`

    for ex in exchanges:
        p = _indented_para(doc)

        if ex["type"] == "Q":
            _run(p, ex["speaker"], bold=True, italic=True)
            _run(p, ": " + ex["text"], italic=True)

        elif ex["type"] == "A":
            _run(p, ex["speaker"] + ":", bold=True)
            _run(p, " " + ex["text"])

        else:  # CONT — continuation paragraph, no speaker label
            _run(p, ex["text"])


def add_qa_entry(
    doc: Document,
    date: datetime,
    summary: str,
    exchanges: list[dict],
    url: str | None = None,
    anchor: str | None = None,
) -> None:
    """Convenience wrapper: date heading + add_qa_entry_body, for callers
    (format_entry.py, one-off manual entries) that add exactly one entry at
    a time and always want their own heading. The scraper's own pipeline
    uses queue_entry()/flush_pending_entries() instead so that multiple
    same-day entries share a single heading — see NOTES.md, 2026-08-04."""
    add_date_heading(doc, date)
    add_qa_entry_body(doc, summary, exchanges, url, anchor)


def add_release_entry_body(
    doc: Document,
    summary: str,
    body_paragraphs: list[str],
    url: str | None = None,
    anchor: str | None = None,
    source_label: str | None = None,
) -> None:
    """
    Write a press release / statement entry's summary + body — NOT the date
    heading or a blank separator; see add_qa_entry_body's docstring, same
    reasoning applies here.
    body_paragraphs: verbatim paragraphs indented as a block.
    If a paragraph starts with 'Speaker: text', the speaker name is bolded.
    Every line in each block paragraph is indented (paragraph-level, not first-line).

    `url`/`anchor`: see add_qa_entry_body — hyperlinks `anchor` within `summary`.

    `source_label`: who/what is actually speaking in the FIRST body
    paragraph, when that paragraph doesn't already carry its own natural
    "Name: text" shape (an X account's display name, or a plain
    institutional label like "Treasury Department") — bold-prefixed onto
    it, matching the real past trackers exactly (e.g. "State Council
    Press Release: A senior official..."). Added 2026-09-02 per user
    request. Only the first paragraph gets it — confirmed against a real
    multi-paragraph tracker entry that subsequent paragraphs of the same
    quoted release have NO repeated label. If the first paragraph already
    matches the speaker pattern below (e.g. a Q&A-shaped release
    fallback that extracted "Zhang Xiaogang: ..." verbatim), that
    existing label wins and `source_label` is not applied — don't stack
    two labels on one paragraph.
    """
    add_summary_para(doc, summary, url, anchor)

    # \w (not A-Za-z0-9) so accented names ("Arévalo") aren't silently
    # unmatched — found live, 2026-09-01, on a joint press availability
    # transcript with foreign leaders. {0,50} (not {1,50}) so a bare
    # single-letter label ("Q:"/"A:" — the standard shorthand a translated
    # 问：/答： comes through as) also matches. See NOTES.md. "()" added
    # 2026-09-03 for the same reason as _QA_RE's note — a real outlet name
    # can carry a parenthetical abbreviation ("International Market News
    # (U.S.)"), which \w alone doesn't cover.
    speaker_re = re.compile(r"^([A-Z][\w \-'\.\:()]{0,50}):\s+(.+)$", re.DOTALL)
    for i, para_text in enumerate(body_paragraphs):
        p = _indented_para(doc)
        m = speaker_re.match(para_text)
        if m:
            _run(p, m.group(1) + ":", bold=True)
            _run(p, " " + m.group(2))
        elif i == 0 and source_label:
            _run(p, source_label + ":", bold=True)
            _run(p, " " + para_text)
        else:
            _run(p, para_text)


def add_release_entry(
    doc: Document,
    date: datetime,
    summary: str,
    body_paragraphs: list[str],
    url: str | None = None,
    anchor: str | None = None,
) -> None:
    """Convenience wrapper: date heading + add_release_entry_body — see
    add_qa_entry's docstring, same reasoning."""
    add_date_heading(doc, date)
    add_release_entry_body(doc, summary, body_paragraphs, url, anchor)


# ── Pending-entry buffer ─────────────────────────────────────────────────────
#
# scrape_* functions no longer write straight to `doc`. They call
# queue_entry() instead, and main() calls flush_pending_entries() once all
# requested sources have run (or after each source, for a long multi-source
# run — see main()). This is what makes "one date heading, even when
# several entries from different sources land on the same day" possible:
# with immediate per-item writes there's no way to know, at write time,
# whether a later item will share today's date. Trade-off: entries are
# fully sorted+grouped WITHIN each flush, and a module-level "last date
# written" is carried across flushes so a later flush's first entry still
# correctly skips repeating today's heading — but two flushes' worth of
# entries are NOT re-interleaved with each other, so if source A (flushed
# first) has a July 30 item and source B (flushed after) has a July 28 item,
# the doc shows A's July 30 block before B's July 28 block, not fully
# chronological end-to-end. Full global chronological order would mean
# deferring every write to the very end of the whole run, which in turn
# means losing ALL of a run's progress if the process is killed partway
# through instead of just the current source's — not worth it given how
# often a long run already hits per-item errors (mostly LLM quota) that are
# already handled by skipping that one item, not the whole source.

PENDING_ENTRIES: list[dict] = []
_LAST_WRITTEN_DATE = None

# The current run's target date range — set by main() from week_start/
# week_end (the requested or default-computed week), None outside a
# main()-driven run (backtest.py, format_entry.py never set these, so
# queue_entry() below never filters for them — those tools legitimately
# want every entry regardless of date). A real, HARD filter, added
# 2026-09-02 per direct user request: previously the requested week was
# only ever a label, never a filter, so a run "targeting Aug 25-31" could
# still queue a genuinely-found Sept 1 item into that week's output —
# confirmed live (a real MFA leadership URL dated Sept 1 got queued
# during a run targeting Aug 25-31). An out-of-range item is dropped
# WITHOUT being marked seen, so a future run whose target actually covers
# its date will still find and correctly include it — not lost, just
# deferred to the right week.
_RUN_TARGET_START: date | None = None
_RUN_TARGET_END: date | None = None

# Guards against a real, live-observed source-side duplicate: state.gov's
# own WP-API listed TWO separate post IDs for the exact same headline
# ("Trump Administration Launches Foundry School...") two days apart,
# with URL slugs differing by exactly one missing hyphen
# ("workforce-behind" vs "workforcebehind") — apparently republished
# rather than edited in place on their end. Different URLs means our
# normal URL-based dedup (is_seen()) doesn't catch it; this is a second,
# narrower net that does. Reset once per main()-driven run (see
# _RUN_TARGET_START above) — a persistent, across-runs version would
# need to check the `entries` table instead, which no real case has
# justified building yet (this one surfaced within a single run).
_QUEUED_URL_SLUGS_THIS_RUN: set[str] = set()
_URL_SLUG_NON_ALNUM_RE = re.compile(r"[^a-z0-9]+")


def _url_dedup_slug(url: str) -> str:
    """Alphanumeric-only fingerprint of a URL's path — collapses
    "workforce-behind" and "workforcebehind" (or any other hyphen/
    underscore/casing variation) to the identical string, so a same-run
    republish under a near-identical slug can be recognized even though
    the exact URLs differ. Deliberately crude (not real similarity
    scoring) — genuinely different articles' slugs still differ enough in
    their actual words to never collide by accident."""
    return _URL_SLUG_NON_ALNUM_RE.sub("", url.lower())


def queue_entry(
    kind: str,               # "qa" | "release"
    date: datetime,
    summary: str,
    url: str | None,
    anchor: str | None,
    exchanges: list[dict] | None = None,
    paragraphs: list[str] | None = None,
    source_label: str | None = None,
) -> None:
    """
    `source_label` (release-kind entries only): who/what is actually
    speaking in `paragraphs` — an X account's display name ("Chinese
    Embassy", "Rapid Response 47"), or an institution ("State Council
    Information Office", "Treasury Department"). Added 2026-09-02 per
    user request, confirmed against the real past trackers first: every
    body paragraph there is bold-prefixed with who's speaking, whether
    that's a named official (Q&A, already handled elsewhere), an X
    account, or a plain institutional label for a release with no named
    speaker at all (e.g. "State Council Press Release: A senior
    official..."). See add_release_entry_body() for where this actually
    gets rendered.
    """
    entry_date = date.date() if hasattr(date, "date") else date
    if _RUN_TARGET_START is not None and not (_RUN_TARGET_START <= entry_date <= _RUN_TARGET_END):
        log.info(
            f"[queue_entry] {entry_date} is outside this run's target range "
            f"({_RUN_TARGET_START}..{_RUN_TARGET_END}) — leaving unseen for "
            f"a future run whose target actually covers it: {url}"
        )
        return
    if url:
        slug = _url_dedup_slug(url)
        if slug in _QUEUED_URL_SLUGS_THIS_RUN:
            log.warning(
                f"[queue_entry] Skipping likely source-side republish "
                f"duplicate (same slug, different URL, already queued "
                f"this run): {url}"
            )
            return
        _QUEUED_URL_SLUGS_THIS_RUN.add(slug)
    PENDING_ENTRIES.append({
        "kind": kind, "date": date, "summary": summary,
        "url": url, "anchor": anchor,
        "exchanges": exchanges, "paragraphs": paragraphs,
        "source_label": source_label,
    })


def flush_pending_entries(doc: Document, conn: sqlite3.Connection) -> int:
    """Sort pending entries by date, write them with one heading per date
    (collapsing repeats against the last date written by a PREVIOUS flush
    too), save, mark each entry's URL seen, and clear the buffer. Returns
    how many entries were written.

    `mark_seen()` happens HERE — after `doc.save()` succeeds — not at
    queue_entry() time, and this is deliberate: every process_*_item()
    function used to call mark_seen() immediately upon deciding an item
    was relevant, in the same breath as queue_entry() appending it to the
    in-memory PENDING_ENTRIES buffer. That buffer isn't written to disk
    until THIS function runs, once per source. A crash/kill between those
    two points — after mark_seen() durably committed to tracker.db, but
    before this source's flush — permanently marks the item seen (so it's
    never retried) while its actual content never made it into the doc:
    silent, permanent data loss, not just "deferred to next run" as the
    per-item exception handling elsewhere in this file correctly is. Found
    live, 2026-09-01 (a demo run killed mid-source lost that source's
    unflushed entries harmlessly only because it happened to be using a
    throwaway dedup database that session, not the real one). Fixed by
    moving mark_seen() to run only after the write it's meant to record
    has actually happened — if doc.save() raises, nothing below it runs,
    so nothing gets marked seen for a flush that didn't happen either.
    """
    global _LAST_WRITTEN_DATE
    count = 0

    if PENDING_ENTRIES:
        ordered = sorted(PENDING_ENTRIES, key=lambda e: e["date"])
        for e in ordered:
            d = e["date"].date() if hasattr(e["date"], "date") else e["date"]
            if d != _LAST_WRITTEN_DATE:
                add_date_heading(doc, e["date"])
                _LAST_WRITTEN_DATE = d
            if e["kind"] == "qa":
                add_qa_entry_body(doc, e["summary"], e["exchanges"], e["url"], e["anchor"])
            else:
                add_release_entry_body(doc, e["summary"], e["paragraphs"], e["url"], e["anchor"],
                                        source_label=e.get("source_label"))

        doc.save(DOC_PATH)
        for e in ordered:
            if e["url"]:
                mark_seen(conn, e["url"])
            # Durable, queryable copy of this entry's content — same
            # after-doc.save() timing/safety as mark_seen() above, and
            # for the same reason: nothing below should run for a flush
            # that didn't actually happen. This is what render_doc_for_
            # range() reads to rebuild a specific week's document later,
            # independent of the ever-growing master doc.
            conn.execute(
                "INSERT INTO entries (url, date, kind, summary, anchor, exchanges_json, paragraphs_json, source_label) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    e["url"],
                    (e["date"].date() if hasattr(e["date"], "date") else e["date"]).isoformat(),
                    e["kind"], e["summary"], e["anchor"],
                    json.dumps(e["exchanges"]) if e["exchanges"] else None,
                    json.dumps(e["paragraphs"]) if e["paragraphs"] else None,
                    e.get("source_label"),
                ),
            )
        conn.commit()
        count = len(PENDING_ENTRIES)
        PENDING_ENTRIES.clear()

    # Persist any staged X since_id advances too (see scrape_x()'s
    # docstring) — safe regardless of whether there were doc entries this
    # flush: either there were none (nothing at risk either way), or
    # doc.save() above already succeeded. Done unconditionally, not inside
    # the `if PENDING_ENTRIES` block above, because a run where every
    # fetched tweet was correctly judged irrelevant still deserves its
    # since_id advanced — otherwise those same tweets get needlessly
    # re-fetched (re-billed) forever. Nested under "_search_groups" since
    # the 2026-09-02 search-based redesign — see that section's own
    # comment — tracks since_id per search GROUP, not per account.
    if _PENDING_X_SINCE_IDS:
        x_state = _load_x_state()
        x_state.setdefault("_search_groups", {})
        for group_key, new_since_id in _PENDING_X_SINCE_IDS.items():
            x_state["_search_groups"].setdefault(group_key, {})["since_id"] = new_since_id
        _save_x_state(x_state)
        _PENDING_X_SINCE_IDS.clear()

    return count


def render_doc_for_range(conn: sqlite3.Connection, start: date, end: date) -> Document:
    """
    Builds a FRESH document containing ONLY entries whose date falls in
    [start, end], reading from the durable `entries` table (populated by
    flush_pending_entries() above) — NOT a copy of the master
    tracker_output.docx, which accumulates every entry from every run
    ever made.

    Added 2026-09-02 to fix a real design gap found live: the dated
    per-week output file was previously just `shutil.copyfile(DOC_PATH,
    ...)` — a full copy of the ENTIRE cumulative history, under a
    filename that implied it held just one week. Combined with
    queue_entry()'s new hard date-range filter (see its own comment),
    this is what actually makes "the document for week X" mean the
    entries from week X: whether an entry was queued in THIS run or
    written weeks ago, it shows up here if its date is in range — so
    re-running an already-fully-covered week is fast (nothing new to
    queue) AND still produces a complete, correct document for it.
    """
    rows = conn.execute(
        "SELECT date, kind, summary, anchor, exchanges_json, paragraphs_json, url, source_label "
        "FROM entries WHERE date >= ? AND date <= ? ORDER BY date, id",
        (start.isoformat(), end.isoformat()),
    ).fetchall()

    doc = Document()
    for para in list(doc.paragraphs):
        if not para.text.strip():
            para._element.getparent().remove(para._element)
    _set_doc_defaults(doc)

    # Grouped by date first, then walked day-by-day across the WHOLE
    # requested range (not just the dates with rows) — added 2026-09-02
    # per user request, confirmed against the real past trackers first:
    # a day with nothing found still gets its own date heading with no
    # body under it (e.g. "Sunday, August 2, 2026" followed immediately
    # by the next heading), not silently skipped. A quiet day is still
    # real information — "nothing happened" is different from "we didn't
    # check."
    by_date: dict[date, list] = {}
    for date_str, kind, summary, anchor, exchanges_json, paragraphs_json, url, source_label in rows:
        by_date.setdefault(date.fromisoformat(date_str), []).append(
            (kind, summary, anchor, exchanges_json, paragraphs_json, url, source_label)
        )

    d = start
    while d <= end:
        add_date_heading(doc, datetime(d.year, d.month, d.day))
        for kind, summary, anchor, exchanges_json, paragraphs_json, url, source_label in by_date.get(d, []):
            if kind == "qa":
                exchanges = json.loads(exchanges_json) if exchanges_json else []
                add_qa_entry_body(doc, summary, exchanges, url, anchor)
            else:
                paragraphs = json.loads(paragraphs_json) if paragraphs_json else []
                add_release_entry_body(doc, summary, paragraphs, url, anchor, source_label=source_label)
        d += timedelta(days=1)
    return doc


# ── Q&A parser ────────────────────────────────────────────────────────────────

# \w, not A-Za-z0-9 — see the note on the other speaker-label regexes:
# accented names (e.g. "PRESIDENT ARÉVALO:", from a real joint press
# availability with a Latin American head of state) silently failed to
# match with a Latin-only character class, which meant this content was
# never even recognized as having 2+ labeled paragraphs at all — not a
# wrong classification, no classification. Found live, 2026-09-01.
# {0,40}, not {1,40}: a bare single-letter label ("Q:"/"A:") has ZERO
# characters between the letter and the colon — this is the STANDARD
# shorthand a translated 问：/答： comes through as (MOFCOM's Chinese Q&A
# pages), so requiring at least 1 extra character silently missed exactly
# the most common case. Also found live, 2026-09-01, same session.
#
# {0,60}, not {0,40} — raised 2026-09-03 after a real live translation
# rendered a reporter's label as "International Market News Agency
# (IMNA) Reporter" (47 characters), which the old {0,40} cap couldn't
# match at all. This silently reclassified that reporter's real question
# (about a proposed US tariff hike) as a CONT paragraph of the PRECEDING
# answer instead of a new "Q" — which then shifted filter_relevant_
# exchanges' Q→A block boundaries enough that the whole real,
# US-mentioning exchange got folded into an earlier, unrelated block and
# silently dropped. A real MOFCOM entry, invisible end to end (not a
# crash, not an error — just missing). Outlet names genuinely vary this
# much in the wild; {0,60} is a deliberate generous bump, not a
# one-off patch for this exact outlet.
#
# Added "()" to the character class too, same fix/same live case — the
# actual label text was "International Market News (U.S.) Reporter:",
# and \w does NOT include parentheses, so the {0,40}->{0,60} bump ALONE
# still didn't match; the parenthetical abbreviation was the real blocker,
# the length was a red herring found first. A label with a parenthetical
# like "(IMNA)"/"(U.S.)"/"(Reuters)" is common enough in real outlet
# names that this needs to be a permanent part of the character class,
# not a one-off allowance.
_QA_RE = re.compile(r"^([A-Z][\w \-'\.()]{0,60}):\s+(.+)$", re.DOTALL)


# Generic "someone is asking a question" labels/patterns — reporters, wire
# services, outlets, interviewers — as opposed to a named official being
# quoted. Ported from the same distinction googledoc_autoformat_extension/
# Code.gs's KNOWN_Q_/KNOWN_A_ tables make, trimmed to what's needed for detection
# (not full speaker-title lookup). "moderator" added 2026-09-01: a joint
# press availability with a foreign head of state (found live, State
# Dept/Rubio transcripts) uses a moderator to introduce speakers and field
# questions — functionally an asker role, just not a journalist.
#
# The second alternation's role words (reporter/journalist/correspondent/
# anchor/interviewer/host) were only in the FIRST (exact-whole-label)
# alternation, so a compound label like "CGTN Reporter" or "Kyodo News
# Correspondent" — outlet name + role word together, a very ordinary
# transcript-label shape — matched neither: not the exact-match branch
# (it's not literally just "reporter"), and not the outlet-suffix branch
# either unless the specific outlet name happened to also be a keyword in
# that list. Found live, 2026-09-01, verifying the distinct-non-asker-name
# heuristic below: "CGTN Reporter"/"CNBC Reporter" (neither "CGTN" nor
# "CNBC" is an outlet-suffix keyword) were falling through as if they were
# ADDITIONAL answerers, inflating the non-asker count on a real MOFCOM
# transcript and wrongly flipping it back to "release."
_ASKER_LABEL_RE = re.compile(
    r"^(reporter|the press|press|question|q|journalist|host|interviewer|"
    r"member of the press|correspondent|anchor|moderator)$"
    r"|(times|daily|news|agency|post|journal|tv|herald|tribune|gazette|"
    r"wire|press|network|broadcast|reporter|journalist|correspondent|"
    r"interviewer|anchor)\b",
    re.IGNORECASE,
)
_KNOWN_OUTLETS_RE = re.compile(
    r"^(afp|ap|reuters|bloomberg|cnn|cnbc|nbc|abc|cbs|bbc|cctv|nhk|"
    r"ria novosti|fox( news| business)?|wsj|ft|nyt)$",
    re.IGNORECASE,
)
_LABEL_RE = re.compile(r"^([A-Z][\w \-'\.()]{0,60}):\s+")  # see _QA_RE's note


def content_type_from_paragraphs(paragraphs: list[str]) -> str:
    """
    'qa' if this looks like a genuine back-and-forth press conference
    (a reporter/outlet asking, an official answering), 'release' if it's a
    statement/document — even one that quotes several different named
    officials, which superficially also produces multiple "Label: text"
    paragraphs but with no one actually asking a question.

    Matters because content type is a property of what a given page
    actually IS, not of which source published it: FMPRC/MOFCOM/MND publish
    both real press-conference transcripts (reporter outlet names asking,
    a spokesperson answering — 'qa') AND plain document releases (multiple
    officials each quoted once, in their own paragraph, no reporter turn at
    all — 'release'); conversely SCIO/State/etc. occasionally publish
    press-conference readouts too. Confirmed against the past trackers:
    e.g. a MOFCOM "position paper" release has ~10 "Section N.M (topic):
    text" paragraphs (not even speaker-shaped) that must render as
    add_release_entry, while an SCIO press conference with 4 different
    officials (Yan Dong, Lin Weilong, Han Yong, He Shaojun) each quoted
    once — no reporter label anywhere — is ALSO 'release' (that's exactly
    how the past tracker formats it: bold name + trimmed quote, no "Q:"),
    while an FMPRC transcript alternating "Reuters: ... / Lin Jian: ..."
    is 'qa'. The discriminator isn't "are there 2+ labeled paragraphs" (that
    alone is format_entry.py's simpler heuristic, which both match) — it's
    "does at least one label look like it's ASKING rather than being
    quoted."
    """
    labels = [m.group(1).strip() for p in paragraphs if (m := _LABEL_RE.match(p))]
    return _classify_by_labels(labels)


def content_type_from_exchanges(exchanges: list[dict]) -> str:
    """
    Same classification as content_type_from_paragraphs(), but from an
    already-built exchanges list (see _build_exchanges) — lets callers that
    already parsed Q&A-style (every FMPRC/MOFCOM/MND page, by default)
    re-check whether the result was actually a genuine Q&A without
    re-deriving the raw paragraph list from scratch. Uses each exchange's
    raw `speaker` field directly, which is the verbatim label regardless of
    whether _build_exchanges (mis-)classified it as Q or A.
    """
    labels = [ex["speaker"] for ex in exchanges if ex.get("speaker")]
    return _classify_by_labels(labels)


def _classify_by_labels(labels: list[str]) -> str:
    """
    'qa' if the labels look like a real back-and-forth, 'release' if not
    — see content_type_from_paragraphs()'s docstring for the full
    discriminator this implements.

    An asker-shaped label (outlet name, "reporter", "Q", ...) is
    NECESSARY but not SUFFICIENT: an ordinary press release can cite media
    coverage ("Fox News: <headline>", "Bloomberg: <headline>") or quote
    several outlets' names without anyone actually being interviewed —
    found live, 2026-09-01, on a real USTR release compiling news-coverage
    headlines plus reaction quotes from EIGHT different named officials/
    orgs (a Senator, several Representatives, trade-association
    presidents...), where "Bloomberg"/"CNBC" (both literal entries in
    _KNOWN_OUTLETS_RE) were each enough to flip the whole release to "qa"
    under the old any-hit-at-all rule. The real discriminator: a genuine
    Q&A press conference has ONE (occasionally two, for a joint briefing)
    person actually answering, no matter how many different reporters ask
    — a "roundup of individual reactions" cites MANY different people
    symmetrically, none of them a repeat answerer.

    So: "release" only if BOTH (a) there are 3+ distinct non-asker-shaped
    names AND (b) none of them repeats. Either condition alone being false
    is enough for "qa": a single answerer (1-2 distinct non-asker names,
    regardless of repeats — covers a single-exchange transcript where the
    answerer's whole turn is one long block and never repeats as a
    separate label) is a real conversation; so is a REPEATING non-asker
    name even among 3+ total, since that repeat IS the answerer regardless
    of how many one-off outlet names sit alongside them.

    This two-part shape (not just "count <= 2") was needed after
    test_scraper.py caught a real regression the count-only version
    introduced: a genuine historical FMPRC transcript (A Tarde / The New
    York Times / Antara asking, Lin Jian answering — the exact transcript
    that motivated the .search()-vs-.match() outlet fix on 2026-08-04) has
    3 distinct non-asker names because "A Tarde" and "Antara" (both real
    foreign wire/outlet names) aren't in _KNOWN_OUTLETS_RE's necessarily-
    incomplete curated list — the count-only rule wrongly called this
    "release" despite Lin Jian clearly repeating 3 times. Verified against
    every real case from this session: MOFCOM's He Yadong (1 distinct
    answerer, repeats — qa), Rubio/Arévalo (1 distinct answerer, doesn't
    repeat — qa via the count alone), this FMPRC transcript (3 distinct
    non-asker names, but Lin Jian repeats — qa via the repeat), and the
    USTR release (8 distinct non-asker names, none repeats — release,
    correctly the only case where neither condition saves it).
    """
    if len(labels) < 2:
        return "release"

    has_asker = any(_ASKER_LABEL_RE.search(l) or _KNOWN_OUTLETS_RE.match(l) for l in labels)
    if not has_asker:
        return "release"

    non_asker_labels = [l for l in labels if not (_ASKER_LABEL_RE.search(l) or _KNOWN_OUTLETS_RE.match(l))]
    too_many_distinct = len(set(non_asker_labels)) >= 3
    someone_repeats = any(non_asker_labels.count(l) >= 2 for l in set(non_asker_labels))
    return "release" if (too_many_distinct and not someone_repeats) else "qa"


def exchanges_to_paragraphs(exchanges: list[dict]) -> list[str]:
    """Flatten a Q&A-shaped exchanges list back into plain paragraphs
    ('Speaker: text', or bare text for CONT) — used when
    content_type_from_exchanges() determines a page parsed as Q&A-shaped
    was actually a release (a statement quoting several officials, no one
    asking a question), so it can be routed through add_release_entry_body
    instead."""
    paras = []
    for ex in exchanges:
        if ex.get("speaker"):
            paras.append(f"{ex['speaker']}: {ex['text']}")
        else:
            paras.append(ex["text"])
    return paras


# {0,40}, not {1,40} — a bare single-letter label like "Q:"/"A:" (the
# standard journalism-transcript shorthand; a MOFCOM Q&A translated from
# 问：/答： comes through exactly this way) has ZERO characters between the
# letter and the colon, so a {1,40} minimum silently failed to match it.
# Found live, 2026-09-01, on a real MOFCOM Q&A about blocking US sanctions
# on 5 Chinese companies — "Q:" alone on its own line was invisible to
# this regex, so the orphan-label merge never ran and the exchange fell
# through to a (correct-outcome-but-avoidable) LLM parse instead of the
# free regex path.
#
# {0,60}, not {0,40} — raised 2026-09-03, same fix and same reasoning as
# _QA_RE above: a real orphan label ("International Market News Agency
# (IMNA) Reporter:", 47 chars, alone on its own line) was too long for
# the old {0,40} cap, so this merge never ran for it either.
_ORPHAN_LABEL_RE = re.compile(r"^[A-Z][\w \-'\.()]{0,60}:$")

# MOFCOM's *regular* press-conference transcripts (as opposed to its
# spokesperson-remarks pages) wrap each speaker name in Chinese brackets —
# "【何亚东】：" — which translate_to_english carries straight through as
# English square brackets ("[He Yadong]:") rather than stripping them.
# Every other label regex in this file requires the label to START with a
# plain Latin/accented letter (`^[A-Z]...`), so a leading "[" made the
# whole label invisible — not misclassified, unrecognized, same failure
# shape as the orphan-label and {0,40} bugs above. Found live, 2026-09-01,
# on a real He Yadong rare-earth-export Q&A. Half-width and full-width
# brackets both handled since a translation occasionally leaves the
# original CJK bracket characters in place instead of converting them —
# and for the same reason, both the ASCII colon (:) and the fullwidth
# Chinese colon (：, a DIFFERENT character, U+FF1A not U+003A) are
# accepted after the closing bracket: found via test_scraper.py that if a
# translation leaves 【】 untranslated, it may just as plausibly leave the
# adjacent ： untranslated too, and the original version of this pattern
# only recognized the ASCII colon, silently failing on genuinely
# untranslated Chinese punctuation despite explicitly trying to handle it.
_BRACKETED_LABEL_RE = re.compile(r"^[\[【]([^\]】]{1,50})[\]】][:：]\s*")


def _unbracket_label(para: str) -> str:
    """Rewrite a leading "[Name]:"/"【Name】:" into plain "Name:" so the
    existing "Label: text" regexes can see it — see _BRACKETED_LABEL_RE.
    Preserves a clean "Name:" (no trailing space) when nothing follows on
    the same paragraph — an orphan label ("[He Yadong]:" alone, content on
    the next line) must still end exactly at the colon for
    _ORPHAN_LABEL_RE's `:$` anchor to match; appending ": " unconditionally
    would silently break that."""
    m = _BRACKETED_LABEL_RE.match(para)
    if not m:
        return para
    name, rest = m.group(1).strip(), para[m.end():]
    return f"{name}: {rest}" if rest else f"{name}:"


def _merge_orphan_speaker_labels(paragraphs: list[str]) -> list[str]:
    """
    Merge a standalone "SPEAKER:" paragraph (nothing after the colon — the
    speaker's first line of remarks lands in the NEXT paragraph instead)
    with the paragraph that follows it, so downstream label detection
    (_LABEL_RE/_QA_RE, both of which require "Label: text" on one line)
    actually sees it. Without this, a transcript that structurally
    alternates speakers is invisible to content_type_from_paragraphs/
    _build_exchanges — not misclassified as "release", just unrecognized
    as having any labels at all. Found live, 2026-09-01, on a Rubio/
    Arévalo joint press availability transcript from state.gov, where
    "MODERATOR:" and "PRESIDENT ARÉVALO:" each sit alone in their own
    <p> tag, one per speaker turn.
    """
    orphan_free: list[str] = []
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if _ORPHAN_LABEL_RE.match(para) and i + 1 < len(paragraphs):
            orphan_free.append(f"{para} {paragraphs[i + 1]}")
            i += 2
        else:
            orphan_free.append(para)
            i += 1
    return orphan_free


_METADATA_LABEL_RE = re.compile(
    r"^(category|source|type|author|date|tags?|editor|reprint|"
    r"来源|类型|分类|作者|日期|标签|编辑|转载)$",
    re.IGNORECASE,
)

# A visual divider between unrelated topics within the same press
# conference transcript — FMPRC's own pages use a bare line of repeated
# asterisks for this (confirmed live, 2026-09-03: a real page had
# "**************************************************" between Guo
# Jiakun's trade-surplus answer and the next reporter's unrelated
# question about Nepal). Not a "Label: text" match, so it fell through to
# a plain CONT paragraph and got written into the tracker as if it were
# more of the preceding answer. Same shape as other junk-in-body-text
# bugs this project has already fixed (raw markup, control characters) —
# strip it before it can become a paragraph at all. `\1{3,}` requires the
# SAME character repeated 4+ times, so it can't accidentally eat a real
# short dash-joined phrase.
_SEPARATOR_LINE_RE = re.compile(r"^([*\-_=~])\1{3,}$")


def _build_exchanges(paragraphs: list[str], spokespersons: set[str]) -> list[dict]:
    """
    Convert a list of text paragraphs into exchange dicts.
    Continuation paragraphs (no 'Name: text' pattern) become CONT entries
    so they render as separate indented paragraphs — matching the existing docs.

    Skips a "Label: text" match whose label is a page-metadata field
    (Category/Source/Type/...), not a real speaker — found live, 2026-09-01,
    on a MOFCOM position-paper page whose translated boilerplate ("Category:
    News", "Source: Xinhua News Agency", "Type: Reprint") matched _QA_RE just
    as well as a real speaker label, tricking _build_exchanges into treating
    a plain document as if it had Q&A structure and dumping the entire
    translated page as undifferentiated exchanges instead of falling through
    to the (correct, keyword-filtered) release path.
    """
    exchanges: list[dict] = []
    in_qa = False

    for para in paragraphs:
        para = para.strip()
        if not para or len(para) < 8:
            continue
        if _SEPARATOR_LINE_RE.match(para):
            continue
        m = _QA_RE.match(para)
        if m and _METADATA_LABEL_RE.match(m.group(1).strip()):
            continue
        if m:
            speaker = m.group(1).strip()
            text    = m.group(2).strip()
            # `sp.lower() in speaker.lower()` only catches a spokespersons
            # entry (e.g. "Answer") that's a SUBSTRING of the label — which
            # fails for the bare single-letter "A:" shorthand a translated
            # 答： often comes through as (there's no substring relation
            # between "answer" and "a"). Recognize a bare "A" as an answer
            # role directly, matching the equally-standard bare "Q" already
            # covered on the asker side by _ASKER_LABEL_RE. Found live,
            # 2026-09-01, on the same real MOFCOM Q&A as the {0,40} fix
            # above — without this, a correctly-detected Q&A exchange still
            # mislabeled its answer half as another question.
            is_sp = speaker.lower() == "a" or any(sp.lower() in speaker.lower() for sp in spokespersons)
            exchanges.append({
                "type":    "A" if is_sp else "Q",
                "speaker": speaker,
                "text":    text,
            })
            in_qa = True
        elif in_qa:
            exchanges.append({"type": "CONT", "speaker": None, "text": para})

    # Second pass: catch a NAMED spokesperson whose personal name isn't in
    # the passed-in `spokespersons` set (a role-word list like
    # {"Spokesperson", "Minister", "Deputy"}). MOFCOM's *regular* press
    # conferences label each turn with the day's actual named official
    # (何亚东/He Yadong, 束珏婷/Shu Jueting, ...) instead of a generic
    # "Spokesperson:" prefix, and the roster rotates too often to maintain
    # a name list for the way FMPRC_SPOKESPERSONS does for FMPRC. Without
    # this, every one of that person's turns above got typed "Q" (the
    # substring check found no match), including the actual answers.
    # Heuristic: a real press conference has exactly one answerer and many
    # different askers, so whichever non-asker-shaped label repeats 2+
    # times is almost certainly the answerer — retype all their turns "A".
    # Found live, 2026-09-01, on a real He Yadong rare-earth-export Q&A.
    label_counts = Counter(ex["speaker"] for ex in exchanges if ex.get("speaker"))
    for ex in exchanges:
        speaker = ex.get("speaker")
        if (
            speaker and ex["type"] == "Q" and label_counts[speaker] >= 2
            and not _ASKER_LABEL_RE.search(speaker)
            and not _KNOWN_OUTLETS_RE.match(speaker)
        ):
            ex["type"] = "A"

    return exchanges


def filter_relevant_exchanges(exchanges: list[dict], require_china_mention: bool = False) -> list[dict]:
    """
    Group exchanges into Q→A blocks, then keep only blocks where any
    exchange text has an explicit US mention (_EXPLICIT_US_MENTION_RE) —
    NOT the broader RELEVANCE_KEYWORDS topic list (Taiwan/AI/semiconductor/
    etc.), which is too loose here: those topics come up in plenty of
    purely-China-and-a-third-country exchanges within the same press
    conference. See _EXPLICIT_US_MENTION_RE's docstring for the real case
    that surfaced this, found live 2026-09-01.

    This ensures that if a journalist asks about Trump/tariffs but the answer
    only mentions China's position (or vice versa), both Q and A are included.
    A block is: one Q + all following A/CONT until the next Q.
    A-only blocks (no preceding Q) are treated as their own block.

    `require_china_mention`: check for an explicit CHINA mention
    (_CHINA_MENTION_RE) instead of a US mention — for Q&A content
    rerouted from a US-origin source (finalize_release_item()'s reroute
    path, e.g. a State Department interview transcript). "Does this
    mention the US" is meaningless there: a US official being
    interviewed mentions the US constantly regardless of topic, so EVERY
    block passed that check. Found live, 2026-09-03: a real Rubio/Brian
    Kilmeade interview about Iran/Cuba/Ukraine/Venezuela got included
    almost entirely (33 of 34 exchanges) even though only ONE exchange
    actually mentioned China — the default US-mention check is exactly
    right for a Chinese-origin multi-topic press conference (FMPRC/
    MOFCOM/MND, where the whole thing is trivially "about China" and the
    real question is whether a given block ALSO touches the US), but is
    the wrong direction entirely once this same function got reused for
    English-origin Q&A rerouted here — same class of directional mismatch
    already fixed once for classify_relevance()'s chinese_origin
    parameter, just in this different function.
    """
    if not exchanges:
        return []

    blocks: list[list[dict]] = []
    current: list[dict] = []

    for ex in exchanges:
        if ex["type"] == "Q":
            if current:
                blocks.append(current)
            current = [ex]
        else:                          # A or CONT
            if current:
                current.append(ex)
            else:
                current = [ex]        # A with no preceding Q

    if current:
        blocks.append(current)

    check_re = _CHINA_MENTION_RE if require_china_mention else _EXPLICIT_US_MENTION_RE
    result = []
    for block in blocks:
        combined = " ".join(ex.get("text", "") for ex in block)
        if check_re.search(combined):
            result.extend(block)

    return result


def parse_qa(html: str, spokespersons: set[str]) -> list[dict]:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()

    content = (
        soup.find("div", class_=re.compile(r"content|article|detail|main", re.I))
        or soup.find("div", id=re.compile(r"content|article|detail|main", re.I))
    )
    if not content:
        divs = soup.find_all("div")
        content = max(divs, key=lambda d: len(d.get_text()), default=soup)

    paragraphs = [p.get_text(separator=" ").strip() for p in content.find_all("p")]
    if not paragraphs:
        paragraphs = [p.get_text(separator=" ").strip() for p in soup.find_all("p")]

    return _build_exchanges(paragraphs, spokespersons)


# ── Q&A parsing schemas and helpers ──────────────────────────────────────────

class QAExchange(BaseModel):
    type: str  # "Q" | "A" | "CONT"
    speaker: Optional[str] = None
    text: str


class QAResponse(BaseModel):
    exchanges: list[QAExchange]


def _call_groq_json(prompt: str, schema: type[BaseModel], label: str = "") -> dict:
    """
    JSON-mode fallback for when Gemini's structured-output call is
    rate-limited/quota-exhausted. Groq has no first-class response_schema
    API here, so we ask for JSON matching the Pydantic schema in the prompt
    and parse the reply, stripping markdown fences if present.
    """
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise RuntimeError("Gemini rate-limited and GROQ_API_KEY is not set")

    json_suffix = (
        "\n\nRespond with ONLY a single valid JSON object matching this "
        "JSON schema — no markdown fences, no commentary, no extra text:\n"
        + json.dumps(schema.model_json_schema())
    )

    client = Groq(api_key=api_key)
    # See _call_groq_with_retry — same reasoning-token exhaustion risk and
    # same 8000-TPM-cap risk apply here too (arguably worse: a JSON object
    # needs more completion tokens than plain text for the same content).
    # `json_suffix` is passed separately so truncation (if `prompt` is
    # long) never eats the part that tells the model to answer in JSON.
    completion = _call_groq_with_retry(client, prompt, suffix=json_suffix)
    _log_groq_usage(label, completion)
    text = (completion.choices[0].message.content or "").strip()
    if not text:
        raise RuntimeError(
            f"Groq returned an empty response (finish_reason="
            f"{completion.choices[0].finish_reason!r}) — likely reasoning-"
            f"token exhaustion"
        )
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.MULTILINE).strip()
    m = re.search(r"\{.*\}", cleaned, re.DOTALL)
    return json.loads(m.group(0) if m else cleaned)


# ── Additional fallback providers (OpenRouter, Cerebras) ────────────────────
#
# Added 2026-08-05 after burning through both Gemini's (20 req/day) and
# Groq's (200k tokens/day) free daily quotas in one day of heavy backtest
# testing — a third and fourth tier so a quota wall on any one provider
# doesn't stop the whole pipeline. Both are plain OpenAI-compatible chat-
# completions endpoints, so one shared helper covers both.
#
# OPENROUTER_MODEL: tried `google/gemma-4-31b-it:free` first, deliberately
# avoiding a reasoning model (see the _call_groq comments on the
# reasoning-token-exhaustion bug) — but live-tested 429 "temporarily
# rate-limited upstream", because OpenRouter routes that `:free` slug
# through Google AI Studio's shared free pool, the SAME contended resource
# behind our own GEMINI_API_KEY's exhaustion — so it was never going to
# help today regardless of which key hits it. Switched to
# `minimax/minimax-m3:free`, routed through a different upstream (GMICloud)
# — live-verified: correct translation, `reasoning_tokens: 0` in the
# response's usage breakdown (it exposes a `reasoning` field but didn't
# spend anything on it for a plain translation prompt) — good sign, but
# `_openai_compatible_chat`'s empty-response guard stays in place as a
# safety net in case a longer/harder prompt ever does trigger reasoning.
#
# CEREBRAS_MODEL: `gemma-4-31b` is listed as available, but live-tested
# 402 "Payment required to access this resource" — an ACCOUNT-level gate
# (the key works, per the earlier 401-with-a-fake-key test; a real key
# still gets 402), not a code issue. Cerebras's "free tier" appears to
# require billing verification on cloud.cerebras.ai before any inference
# call succeeds — left as configured for whenever that's resolved; every
# call through the fallback chain will simply fail over past it in the
# meantime (see _fallback_chain), so it costs nothing to leave wired up.
OPENROUTER_MODEL = "minimax/minimax-m3:free"
CEREBRAS_MODEL   = "gemma-4-31b"


def _openai_compatible_chat(
    base_url: str, api_key: str, model: str, prompt: str, max_tokens: int = 2000,
    label: str = "", provider_name: str = "", extra_body: dict | None = None,
) -> str:
    """Shared plain-text call for any OpenAI-compatible /chat/completions
    endpoint. Raises on an empty response rather than returning it — same
    reasoning as _call_groq's empty-response guard, since we have no way to
    know in advance whether a given free model on a given provider has the
    same invisible-reasoning-token failure mode until we've seen it not
    happen. `extra_body` merges in provider-specific params (e.g. xAI's
    `reasoning_effort: "none"`, same reasoning-token-disable idea as
    Groq's `reasoning_effort="low"`, added 2026-09-01)."""
    resp = httpx.post(
        f"{base_url}/chat/completions",
        headers={"Authorization": f"Bearer {api_key}"},
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            **(extra_body or {}),
        },
        timeout=60.0,
    )
    resp.raise_for_status()
    # NOT resp.json() — OpenRouter prefixes the body with blank/keep-alive
    # lines (an SSE-style anti-timeout comment pattern) even on a
    # non-streaming request, which breaks a plain json.loads(resp.text).
    # Found live, 2026-08-05. Locating the first "{" is robust to that
    # without assuming which providers do it and which don't.
    body = resp.text
    data = json.loads(body[body.index("{"):])
    usage = data.get("usage")
    if usage:
        details = usage.get("completion_tokens_details") or {}
        _log_usage(
            provider_name or base_url, label,
            usage.get("prompt_tokens"), usage.get("completion_tokens"),
            usage.get("total_tokens"), details.get("reasoning_tokens"),
        )
    text = (data["choices"][0]["message"].get("content") or "").strip()
    if not text:
        finish_reason = data["choices"][0].get("finish_reason")
        raise RuntimeError(
            f"{base_url} returned an empty response (finish_reason={finish_reason!r})"
        )
    return text


def _openai_compatible_chat_json(
    base_url: str, api_key: str, model: str, prompt: str, schema: type[BaseModel],
    max_tokens: int = 2000, label: str = "", provider_name: str = "",
    extra_body: dict | None = None,
) -> dict:
    """Shared JSON-via-prompt call — same technique as _call_groq_json,
    since free-tier OpenAI-compatible endpoints don't reliably support a
    first-class response_schema/json_schema mode across providers."""
    full_prompt = (
        prompt
        + "\n\nRespond with ONLY a single valid JSON object matching this "
          "JSON schema — no markdown fences, no commentary, no extra text:\n"
        + json.dumps(schema.model_json_schema())
    )
    text = _openai_compatible_chat(
        base_url, api_key, model, full_prompt, max_tokens=max_tokens,
        label=label, provider_name=provider_name, extra_body=extra_body,
    )
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.MULTILINE).strip()
    m = re.search(r"\{.*\}", cleaned, re.DOTALL)
    return json.loads(m.group(0) if m else cleaned)


def _call_openrouter(prompt: str, label: str = "") -> str:
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        raise RuntimeError("OPENROUTER_API_KEY is not set")
    return _openai_compatible_chat(
        "https://openrouter.ai/api/v1", api_key, OPENROUTER_MODEL, prompt,
        label=label, provider_name="OpenRouter",
    )


def _call_openrouter_json(prompt: str, schema: type[BaseModel], label: str = "") -> dict:
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        raise RuntimeError("OPENROUTER_API_KEY is not set")
    return _openai_compatible_chat_json(
        "https://openrouter.ai/api/v1", api_key, OPENROUTER_MODEL, prompt, schema,
        label=label, provider_name="OpenRouter",
    )


def _cerebras_api_key() -> str | None:
    # User-set name was CEREBAS_API_KEY (no R) — accept both spellings
    # rather than silently doing nothing over a typo.
    return os.environ.get("CEREBRAS_API_KEY") or os.environ.get("CEREBAS_API_KEY")


def _call_cerebras(prompt: str, label: str = "") -> str:
    api_key = _cerebras_api_key()
    if not api_key:
        raise RuntimeError("CEREBRAS_API_KEY is not set")
    return _openai_compatible_chat(
        "https://api.cerebras.ai/v1", api_key, CEREBRAS_MODEL, prompt,
        label=label, provider_name="Cerebras",
    )


def _call_cerebras_json(prompt: str, schema: type[BaseModel], label: str = "") -> dict:
    api_key = _cerebras_api_key()
    if not api_key:
        raise RuntimeError("CEREBRAS_API_KEY is not set")
    return _openai_compatible_chat_json(
        "https://api.cerebras.ai/v1", api_key, CEREBRAS_MODEL, prompt, schema,
        label=label, provider_name="Cerebras",
    )


# xAI (Grok) — added 2026-09-01 as a 5th fallback tier, after the user's
# X_API_KEY turned out to be an xAI credential (key prefix "xai-...") set
# up by mistake for the actual X/Twitter API — X and xAI are separate
# products under different auth systems, confirmed live (the key 401'd
# against api.x.com no matter what). Repurposed rather than wasted: it's a
# real, working LLM credential, so it slots into the same OpenAI-
# compatible-endpoint pattern as OpenRouter/Cerebras. Endpoint/model/
# pricing confirmed live via docs.x.ai, 2026-09-01. User has since renamed
# this credential to GROK_API_KEY and set up a separate, real X_API_KEY
# (an actual X/Twitter Bearer token) for the tweet-reading feature — the
# two are no longer interchangeable, so this reads GROK_API_KEY only.
XAI_MODEL = "grok-4.3"
# grok-4.3 is a reasoning model by default — same invisible-token-spend
# shape as GROQ_MODEL/gemini-2.5-flash (see both of those fixes above).
# reasoning_effort="none" disables it outright, per xAI's own docs.
_XAI_EXTRA_BODY = {"reasoning_effort": "none"}


def _xai_api_key() -> str | None:
    return os.environ.get("GROK_API_KEY") or os.environ.get("XAI_API_KEY")


def _call_xai(prompt: str, label: str = "") -> str:
    api_key = _xai_api_key()
    if not api_key:
        raise RuntimeError("XAI_API_KEY/X_API_KEY is not set")
    return _openai_compatible_chat(
        "https://api.x.ai/v1", api_key, XAI_MODEL, prompt,
        label=label, provider_name="XAI", extra_body=_XAI_EXTRA_BODY,
    )


def _call_xai_json(prompt: str, schema: type[BaseModel], label: str = "") -> dict:
    api_key = _xai_api_key()
    if not api_key:
        raise RuntimeError("XAI_API_KEY/X_API_KEY is not set")
    return _openai_compatible_chat_json(
        "https://api.x.ai/v1", api_key, XAI_MODEL, prompt, schema,
        label=label, provider_name="XAI", extra_body=_XAI_EXTRA_BODY,
    )


def _fallback_chain(kind: str) -> list[tuple[str, Callable]]:
    """
    Ordered list of (name, fn) fallback providers to try after Gemini,
    built from whichever API keys are actually present in the environment
    — so adding/removing a key in .env changes the chain with no code
    edit. `kind` is "text" or "json".
    """
    text_fns = {"Groq": _call_groq, "OpenRouter": _call_openrouter, "XAI": _call_xai, "Cerebras": _call_cerebras}
    json_fns = {"Groq": _call_groq_json, "OpenRouter": _call_openrouter_json, "XAI": _call_xai_json, "Cerebras": _call_cerebras_json}
    fns = text_fns if kind == "text" else json_fns
    key_present = {
        "Groq": bool(os.environ.get("GROQ_API_KEY")),
        "OpenRouter": bool(os.environ.get("OPENROUTER_API_KEY")),
        "XAI": bool(_xai_api_key()),
        "Cerebras": bool(_cerebras_api_key()),
    }
    # XAI before Cerebras: Cerebras has never successfully billed a call
    # all session (402 Payment Required throughout) — put the known-broken
    # tier last rather than ahead of a real, working credential.
    return [(name, fns[name]) for name in ("Groq", "OpenRouter", "XAI", "Cerebras") if key_present[name]]


def _call_fallback_chain(prompt: str, label: str = "") -> str:
    """Try each configured fallback provider in turn (see _fallback_chain);
    raise the last error if every one fails, or a clear error if none are
    configured at all."""
    providers = _fallback_chain("text")
    if not providers:
        raise RuntimeError("Gemini rate-limited and no fallback API keys are configured")
    last_exc: Exception | None = None
    for name, fn in providers:
        try:
            return fn(prompt, label=label)
        except Exception as exc:
            log.warning(f"{name} fallback failed, trying next: {exc}")
            last_exc = exc
    raise last_exc


def _call_fallback_chain_json(prompt: str, schema: type[BaseModel], label: str = "") -> dict:
    """JSON counterpart to _call_fallback_chain — see its docstring."""
    providers = _fallback_chain("json")
    if not providers:
        raise RuntimeError("Gemini rate-limited and no fallback API keys are configured")
    last_exc: Exception | None = None
    for name, fn in providers:
        try:
            return fn(prompt, schema, label=label)
        except Exception as exc:
            log.warning(f"{name} JSON fallback failed, trying next: {exc}")
            last_exc = exc
    raise last_exc


def call_llm_json(
    client: genai.Client,
    prompt: str,
    schema: type[BaseModel],
    retries: int = 2,
    label: str = "",
) -> dict:
    """
    Call Gemini for a structured JSON response using a Pydantic schema.
    Falls back to Groq (JSON-via-prompt, see _call_groq_json) on 429/quota
    errors — mirrors call_llm's fallback. Added 2026-08-04: Gemini's
    free-tier quota for gemini-2.5-flash turned out to be a hard 20
    requests/DAY cap (not just a 15 RPM rate limit as the original comments
    assumed), so on any real daily run this structured path was hitting
    quota almost immediately and either eating 3 slow retries before
    degrading (generate_summary) or failing the whole entry outright
    (parse_qa_with_llm, which has no other fallback). See NOTES.md.

    `label` identifies the calling code path for the [usage] token-count
    log — see _log_usage(). Skips straight to the fallback chain, no
    sleep, while Gemini is on cooldown from a recent 429 — see
    _gemini_on_cooldown()'s docstring (defined above call_llm).
    """
    if _gemini_on_cooldown():
        return _call_fallback_chain_json(prompt, schema, label=label)
    time.sleep(GEMINI_SLEEP)
    for attempt in range(retries + 1):
        try:
            resp = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=schema,
                    thinking_config=_GEMINI_THINKING_CONFIG,
                ),
            )
            usage = getattr(resp, "usage_metadata", None)
            if usage is not None:
                _log_usage(
                    "Gemini", label,
                    getattr(usage, "prompt_token_count", None),
                    getattr(usage, "candidates_token_count", None),
                    getattr(usage, "total_token_count", None),
                    getattr(usage, "thoughts_token_count", None),
                )
            return json.loads(resp.text)
        except Exception as exc:
            err = str(exc).lower()
            if "429" in err or "quota" in err or "resource_exhausted" in err:
                log.warning(f"Gemini rate-limited/quota exceeded — falling back for JSON, cooling down {_GEMINI_COOLDOWN_SECONDS}s")
                _start_gemini_cooldown()
                return _call_fallback_chain_json(prompt, schema, label=label)
            log.warning(f"Gemini JSON API call failed (attempt {attempt+1}/{retries+1}): {exc}")
            if attempt < retries:
                time.sleep(8 * (attempt + 1))
                continue
            raise


def split_single_paragraph(text: str) -> str:
    """
    If the text has 1 or 2 lines but contains multiple speaker colon patterns,
    insert newlines before the speaker labels so that each turn is on its own line.
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if len(lines) <= 2:
        # Match word boundaries for uppercase speaker names (like 'Lin Jian:', 'Q:', 'REPORTER:')
        # which are not preceded by 'http' or 'https'.
        # \w, not A-Za-z0-9 — same accented-name fix as the other speaker-
        # label regexes (see NOTES.md, bug #13): a Latin-only class silently
        # fails to match names like "ARÉVALO". {0,60}, not {1,40} — a bare
        # "Q:"/"A:" label needs 0 characters between the letter and colon,
        # and a long outlet name needs more than 40 (see _QA_RE's note,
        # 2026-09-03). "()" in the class too — a parenthetical abbreviation
        # in a real outlet name ("International Market News (U.S.)") isn't
        # a \w character either; see _QA_RE's note for the live case.
        pattern = re.compile(
            r"(?<!https)(?<!http)\b([A-Z][\w \-'\.()]{0,60}):\s+",
            re.MULTILINE
        )
        matches = list(pattern.finditer(text))
        if len(matches) >= 2:
            parts = []
            last_idx = 0
            for m in matches:
                start = m.start()
                if start > last_idx:
                    segment = text[last_idx:start].strip()
                    if segment:
                        parts.append(segment)
                last_idx = start
            parts.append(text[last_idx:].strip())
            return "\n\n".join(parts)
    return text


def parse_qa_with_llm(model: genai.Client, text: str, spokespersons: set[str]) -> list[dict]:
    """
    Use Gemini to parse and classify raw text into Q&A exchanges.
    Returns a list of dicts with keys: 'type', 'speaker', 'text'.
    """
    sp_list = ", ".join(sorted(spokespersons)) if spokespersons else "None specified"
    
    prompt = f"""You are an expert transcript parser for a foreign policy tracker.
Your task is to parse and segment the following transcript text into a structured JSON list of dialogue segments.

The known official spokespersons / government officials for this source are: [{sp_list}].

Each item in the output JSON list MUST have exactly these three keys:
- "type": "Q" or "A" or "CONT"
- "speaker": The name of the speaker (a string, or null if type is "CONT")
- "text": The exact verbatim text spoken (string)

Rules for "type":
- "Q" (Question): Spoken by a journalist, reporter, media outlet, or other questioning party. Usually contains a question, and Q always precedes A.
- "A" (Answer): Spoken by a government official, spokesperson, or official speaker. If there are no questions in the text at all (e.g., a white paper, unilateral statement, or reading of a document), everything should be classified as "A".
- "CONT" (Continuation): Use this for any subsequent/continuation paragraphs spoken by the SAME speaker as the previous paragraph, when there is a paragraph break but no new speaker label in the raw text. For CONT, "speaker" must be null.

Important structural guidelines:
- "A" always follows "Q" (or "CONT" of "Q"). "A" can also follow "A" (or "CONT" of "A") if another official speaks.
- Do NOT translate, summarize, or paraphrase. The "text" field MUST contain the exact verbatim English text from the transcript.
- If the speaker name is provided, extract it exactly (e.g. "Lin Jian", "Reporter", "Journalist", "Wall Street Journal").

Transcript text to parse:
\"\"\"
{text}
\"\"\"
"""
    result = call_llm_json(model, prompt, QAResponse, label="parse_qa_with_llm")
    
    exchanges = []
    for item in result.get("exchanges", []):
        ex_type = item.get("type", "CONT")
        if ex_type not in ("Q", "A", "CONT"):
            ex_type = "CONT"
        
        exchanges.append({
            "type": ex_type,
            "speaker": item.get("speaker"),
            "text": item.get("text", "").strip(),
        })
    return exchanges


def parse_qa_from_plaintext(
    text: str,
    spokespersons: set[str],
    model: genai.Client | None = None,
) -> list[dict]:
    """
    Regex-first, LLM-as-last-resort — changed 2026-09-01 per user request.
    _build_exchanges() slices the ORIGINAL text verbatim at each detected
    "Label: text" boundary (it's a formatter, not a rewriter): no
    reproduction, so no risk of an LLM silently paraphrasing, dropping, or
    garbling a line — which is exactly the failure mode the user reported
    seeing downstream (occasional errors piping LLM-reconstructed text
    through the Google Docs formatting step). After bug #13's fixes
    (Unicode-safe speaker-label regex, orphan-label-paragraph merging,
    "moderator" recognized as an asker role), the regex parser correctly
    handles the large majority of real transcripts on its own. The LLM
    path (parse_qa_with_llm, which DOES reproduce the text — see its own
    docstring) is now only a fallback for genuinely unlabeled prose with
    no "Speaker: text" structure at all for the regex to anchor on.
    """
    text = split_single_paragraph(text)

    paragraphs = [p for p in text.split("\n") if p.strip()]
    paragraphs = [_unbracket_label(p) for p in paragraphs]
    # Same fix as finalize_release_item's — MOFCOM's Chinese Q&A pages use
    # a standalone "问："/"答：" ("Question:"/"Answer:") label on its own
    # line, with the actual content starting on the NEXT line; translation
    # (told to preserve structure) carries that shape straight through as
    # "Question:" / "Answer:" alone on a line. Without merging, _QA_RE
    # never matches (no text after the colon on that line) and the whole
    # exchange is invisible to the regex parser — found live, 2026-09-01,
    # on a real MOFCOM Q&A about blocking US sanctions on 5 Chinese
    # companies, misclassified as "release" for exactly this reason.
    paragraphs = _merge_orphan_speaker_labels(paragraphs)
    exchanges = _build_exchanges(paragraphs, spokespersons)
    if exchanges:
        return exchanges

    if model:
        try:
            return parse_qa_with_llm(model, text, spokespersons)
        except Exception as exc:
            log.warning(f"LLM-driven Q&A classification failed: {exc}")

    return []


# ── Shared per-item finalizers ───────────────────────────────────────────────
#
# Content type (Q&A vs. release) is a property of what a given page IS, not
# of which source published it — see content_type_from_paragraphs()'s
# docstring. These two functions are the shared tail every scrape_* function
# calls once it has either an `exchanges` list (sources that default to
# parsing as Q&A: FMPRC/MOFCOM/MND) or a plain-text blob (sources that
# default to release: State/Treasury/USTR/WhiteHouse/DoW/SCIO/MFA
# leadership) — each re-checks that assumption against the actual content
# and reroutes if it guessed wrong, instead of forcing everything through
# one shape. Added 2026-08-04 after finding real past-tracker entries on
# both sides of this (an MOFCOM document release with no Q&A at all, and an
# SCIO Q&A-style press conference) — see NOTES.md.

def finalize_qa_item(
    model: genai.Client,
    tag: str,
    url: str,
    date: datetime,
    exchanges: list[dict],
    work_text: str,
    source_name: str,
    conn: sqlite3.Connection,
    raw_zh_text: str | None = None,
) -> bool:
    """Filter to the relevant block(s), re-check whether this is actually a
    genuine back-and-forth or a statement quoting several officials with no
    one asking a question, generate the summary+anchor, and queue the
    entry. Returns True if something was queued.

    `raw_zh_text`: the ORIGINAL, untranslated Chinese text, passed by
    Chinese-source callers (fmprc/mofcom/mnd) so the no-exchanges fallback
    below can use select_relevant_chinese_paragraphs() (keyword-only, no
    LLM judgment call) instead of extract_key_paragraphs() on the already-
    translated `work_text`. Leave unset for English-original sources."""
    if not exchanges:
        # Zero labeled paragraphs found at all — not necessarily "nothing
        # here": a plain position-paper/document release (e.g. a MOFCOM
        # "Section I.1 (topic): text" breakdown) has NO "Speaker: text"
        # structure anywhere, so parse_qa()/_build_exchanges() correctly
        # comes back empty, but the page can still be entirely relevant —
        # just not Q&A-shaped. Try the release path on the raw text before
        # giving up. Found via backtest.py against a real MOFCOM position
        # paper that was being dropped outright; see NOTES.md, 2026-08-04.
        log.info(f"[{tag}] No Q&A structure found — trying release-style extraction: {url}")
        if raw_zh_text is not None:
            zh_paras = select_relevant_chinese_paragraphs(raw_zh_text)
            if not zh_paras:
                log.info(f"[{tag}] No US-mentioning paragraphs found — skipping: {url}")
                mark_seen(conn, url)
                return False
            # One translation call over just the matched paragraphs, not
            # the whole page — cheaper than translating everything, and
            # skips extract_key_paragraphs()'s LLM judgment call entirely.
            translated_block = translate_to_english(model, "\n\n".join(zh_paras))
            paras = [p.strip() for p in translated_block.split("\n\n") if p.strip()]
        else:
            paras = extract_key_paragraphs(model, work_text)
        if not paras:
            log.info(f"[{tag}] No China-relevant paragraphs found either — skipping: {url}")
            # extract_key_paragraphs is an LLM judgment call ("is anything
            # here relevant enough"), unlike select_relevant_chinese_
            # paragraphs's plain keyword check just above — worth a human
            # second look. See flag_for_review()'s docstring.
            flag_for_review(url, work_text[:80], "extract_key_paragraphs found no relevant paragraphs")
            mark_seen(conn, url)
            return False
        summary, anchor = get_summary_and_anchor(model, work_text, source_name, url)
        queue_entry("release", date, summary, url, anchor, paragraphs=paras, source_label=source_name)
        # mark_seen() happens in flush_pending_entries() — see its docstring.
        return True

    exchanges = filter_relevant_exchanges(exchanges)
    if not exchanges:
        log.info(f"[{tag}] No relevant exchanges — skipping: {url}")
        mark_seen(conn, url)
        return False

    summary, anchor = get_summary_and_anchor(model, work_text, source_name, url)

    if content_type_from_exchanges(exchanges) == "release":
        log.info(f"[{tag}] Content is a statement/release quoting officials, not a Q&A — rerouting: {url}")
        queue_entry("release", date, summary, url, anchor,
                    paragraphs=exchanges_to_paragraphs(exchanges))
    else:
        queue_entry("qa", date, summary, url, anchor, exchanges=exchanges)

    # mark_seen() happens in flush_pending_entries(), after the entry is
    # actually written to disk — not here. See flush_pending_entries()'s
    # docstring for why (bug found live, 2026-09-01).
    return True


def process_release_common(
    tag: str,
    url: str,
    title: str,
    date: datetime,
    plain: str,
    source_name: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    chinese_origin: bool = False,
) -> bool:
    """
    Shared tail for every "defaults to release" source once it has the raw
    page text in hand: the classify_relevance() gate (itself keyword-
    pre-filtered, see US_SOURCE_RELEVANCE_KEYWORDS) followed by
    finalize_release_item(). One place for backtest.py to call regardless
    of how a given source originally obtained `plain` (RSS-embedded content
    for state/whitehouse; a direct page fetch + extract_main_text()
    everywhere else).

    `chinese_origin`: pass through to classify_relevance() AND
    finalize_release_item() — see their own docstrings. True for SCIO
    (a Chinese-government source, where "does this involve China" is a
    meaningless question for classify_relevance, and "does this Q&A
    block mention the US" is the correct direction if it ever turns out
    to be Q&A-shaped); leave False for every US-origin caller
    (state/treasury/ustr/whitehouse), where both are exactly the right
    direction already.
    """
    # Pass the FULL text, not a pre-truncated slice — found live (backtest.py,
    # 2026-08-05) that a real, already-covered Treasury sanctions release only
    # names the actual Chinese/Hong-Kong-based shipping companies at
    # character ~4300; a 2500-char cutoff here meant classify_relevance's
    # free keyword pre-filter never even SAW the part that made it relevant.
    # classify_relevance() does its own (larger) truncation internally for
    # the LLM call itself, to bound cost — this just stops truncating BEFORE
    # the free part gets a look.
    is_rel, reason = classify_relevance(model, f"{title}\n\n{plain}", chinese_origin=chinese_origin)
    if not is_rel:
        if not reason.startswith("Keyword pre-filter:"):
            flag_for_review(url, title, reason)
        mark_seen(conn, url)
        return False

    queued = finalize_release_item(model, tag, url, date, plain, source_name, conn,
                                    chinese_origin=chinese_origin)
    if queued:
        log.info(f"[{tag}] Queued: {title}")
    return queued


def finalize_release_item(
    model: genai.Client,
    tag: str,
    url: str,
    date: datetime,
    plain: str,
    source_name: str,
    conn: sqlite3.Connection,
    spokespersons: set[str] | None = None,
    raw_zh_text: str | None = None,
    chinese_origin: bool = False,
) -> bool:
    """Re-check whether this page is actually a genuine Q&A press
    conference before defaulting to release-style verbatim-paragraph
    extraction; generate the summary+anchor and queue the entry. Returns
    True if something was queued.

    `raw_zh_text`: the ORIGINAL, untranslated Chinese text, for
    Chinese-source callers. When set, the non-Q&A branch below uses
    select_relevant_chinese_paragraphs() (a narrow, explicit "does this
    paragraph name the US" keyword check — same one finalize_qa_item()'s
    own no-exchanges fallback already uses for fmprc/mofcom/mnd) instead
    of extract_key_paragraphs()'s LLM judgment call. Leave unset for
    English-original sources.

    `chinese_origin`: which direction the Q&A-reroute branch below should
    filter in — see filter_relevant_exchanges()'s own `require_china_
    mention` docstring for the real bug (a Rubio interview transcript
    almost entirely included because "does this mention the US" is
    trivially true throughout it) this exists to prevent. True for SCIO
    (a Chinese-government source that can still occasionally turn out to
    be Q&A-shaped); also treated as True whenever `raw_zh_text` is set
    (MFA leadership), since that's Chinese-origin by definition even if
    this particular call site didn't also set this flag explicitly.
    Leave False for state/treasury/ustr/whitehouse.
    """
    is_chinese_source = chinese_origin or raw_zh_text is not None
    paragraphs = [p.strip() for p in plain.split("\n") if p.strip()]
    paragraphs = [_unbracket_label(p) for p in paragraphs]
    paragraphs = _merge_orphan_speaker_labels(paragraphs)

    if content_type_from_paragraphs(paragraphs) == "qa":
        log.info(f"[{tag}] Content looks like a genuine Q&A, not a plain release — rerouting: {url}")
        exchanges = parse_qa_from_plaintext("\n\n".join(paragraphs), spokespersons or set(), model)
        exchanges = filter_relevant_exchanges(exchanges, require_china_mention=not is_chinese_source)
        if not exchanges:
            log.info(f"[{tag}] No relevant exchanges after rerouting — skipping: {url}")
            mark_seen(conn, url)
            return False
        summary, anchor = get_summary_and_anchor(model, plain, source_name, url)
        queue_entry("qa", date, summary, url, anchor, exchanges=exchanges)
        # mark_seen() happens in flush_pending_entries() — see its docstring.
        return True

    if raw_zh_text is not None:
        # Narrow, explicit "does this paragraph name the US" gate — see
        # this function's own docstring.
        zh_paras = select_relevant_chinese_paragraphs(raw_zh_text)
        if not zh_paras:
            log.info(f"[{tag}] No US-mentioning paragraphs found — skipping: {url}")
            mark_seen(conn, url)
            return False
        # One translation call over just the matched paragraphs, not the
        # whole page — same pattern as finalize_qa_item()'s equivalent
        # fallback.
        translated_block = translate_to_english(model, "\n\n".join(zh_paras))
        paras = [p.strip() for p in translated_block.split("\n\n") if p.strip()]
    else:
        paras = extract_key_paragraphs(model, plain)
    if not paras:
        log.info(f"[{tag}] No China-relevant paragraphs found — skipping: {url}")
        flag_for_review(url, plain[:80], "extract_key_paragraphs found no relevant paragraphs")
        mark_seen(conn, url)
        return False

    summary, anchor = get_summary_and_anchor(model, plain, source_name, url)
    queue_entry("release", date, summary, url, anchor, paragraphs=paras, source_label=source_name)
    # mark_seen() happens in flush_pending_entries() — see its docstring.
    return True


# ── RSS helpers ───────────────────────────────────────────────────────────────

def _rss_text(tag) -> str:
    return tag.get_text(strip=True) if tag else ""


def parse_rss(xml_text: str, limit: int = 50) -> list[dict]:
    soup = BeautifulSoup(xml_text, "xml")
    items = []
    for item in soup.find_all("item")[:limit]:
        content_tag = item.find("content:encoded") or item.find("description")
        link = _rss_text(item.find("link")) or _rss_text(item.find("guid"))
        items.append({
            "_type":       "rss",
            "title":       _rss_text(item.find("title")),
            "link":        link,
            "pubDate":     _rss_text(item.find("pubDate")),
            "content_raw": str(content_tag) if content_tag else "",
        })
    return items


def item_url(item: dict) -> str:
    """
    RSS items always carry a "link" string (parse_rss already resolves its
    own link-vs-guid fallback internally), so the `guid` fallback here only
    ever actually fires for a WP-API item missing "link" — and a WP-API
    item's raw `guid` field is `{"rendered": "https://..."}`, a dict, not a
    string. Every real source used by this pipeline includes "link" in its
    WP-API responses, so this path is untested-in-practice, but returning
    the raw dict unfixed here would silently poison every downstream use of
    the "URL" (sqlite storage/lookup, dedup, doc hyperlinks) with a dict
    instead of a string. Found by inspection, 2026-09-02 — hardened rather
    than left as a latent trap for the day some future source omits "link".
    """
    url = item.get("link") or item.get("guid", "")
    return url.get("rendered", "") if isinstance(url, dict) else url


def item_title(item: dict) -> str:
    if item.get("_type") == "rss":
        return item.get("title", "")
    raw = item.get("title", {})
    return BeautifulSoup(
        raw.get("rendered", "") if isinstance(raw, dict) else str(raw),
        "html.parser",
    ).get_text()


def item_content(item: dict) -> str:
    if item.get("_type") == "rss":
        return item.get("content_raw", "")
    raw = item.get("content", {})
    return raw.get("rendered", "") if isinstance(raw, dict) else str(raw)


def item_date(item: dict) -> datetime:
    if item.get("_type") == "rss":
        try:
            from email.utils import parsedate_to_datetime
            return parsedate_to_datetime(item["pubDate"]).replace(tzinfo=None)
        except Exception:
            return _utcnow()
    try:
        return datetime.fromisoformat(item.get("date", ""))
    except Exception:
        return _utcnow()


# ── Source: FMPRC ─────────────────────────────────────────────────────────────

def process_fmprc_item(
    url: str,
    title: str,
    label: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
) -> bool:
    """
    Per-item body for an FMPRC list entry. Split out from scrape_fmprc so
    backtest.py can call the exact same code path against a known URL
    instead of a re-implementation that could drift from the real pipeline.
    """
    tag = f"fmprc/{label}"
    try:
        resp = fetch(client, url)
        if not resp:
            return False

        raw_html = resp.text
        # CJK ratio must be measured on the isolated article body, not the
        # raw page: found live that FMPRC's EN-mirror pages carry a language
        # switcher menu ("简体中文", "Русский", ...) whose handful of CJK
        # characters was enough to tip a genuinely all-English article over
        # the cjk>100 threshold, routing it into the Chinese branch — where
        # the (correctly Chinese-keyword-empty, because there's no real
        # Chinese content) pre-filter then dropped it entirely instead of
        # ever reaching the English parse_qa() branch below. Confirmed via
        # backtest.py against 2026-07-28's Lin Jian/Lula entry.
        main_text = extract_main_text(raw_html)
        cjk = sum(1 for c in main_text if "一" <= c <= "鿿")

        raw_zh_text = None
        if cjk > 100:
            if not CHINESE_RELEVANCE_KEYWORDS.search(main_text):
                log.info(f"[{tag}] No relevant keywords in raw text — skipping translation: {url}")
                mark_seen(conn, url)
                return False
            log.info(f"[{tag}] Translating: {url}")
            translated = translate_to_english(model, main_text[:7000])
            exchanges  = parse_qa_from_plaintext(translated, FMPRC_SPOKESPERSONS, model)
            work_text  = translated
            raw_zh_text = main_text
        else:
            exchanges = parse_qa(raw_html, FMPRC_SPOKESPERSONS)
            work_text = main_text

        date_m = re.search(r"t(\d{8})_", url)
        date   = (
            datetime.strptime(date_m.group(1), "%Y%m%d") if date_m
            else _utcnow()
        )

        queued = finalize_qa_item(model, tag, url, date, exchanges,
                                   work_text, f"MFA {label}", conn,
                                   raw_zh_text=raw_zh_text)
        if queued:
            log.info(f"[{tag}] Queued: {title}")
        return queued

    except Exception as exc:
        log.error(f"[{tag}] Error on {url}: {exc}")
        return False


def scrape_fmprc(
    list_url: str,
    label: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
    client: httpx.Client,
) -> None:
    log.info(f"[fmprc/{label}] {list_url}")
    resp = fetch(client, list_url)
    if not resp:
        log.error(f"[fmprc/{label}] Failed to fetch list")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    raw_links = []
    for a in soup.find_all("a", href=re.compile(r"/\d{6}/t\d+_\d+\.html")):
        raw_links.append((urljoin(list_url, a["href"]), a.get_text(strip=True)))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[fmprc/{label}] {len(new_links)} new items")

    for url, title in new_links:
        try:
            process_fmprc_item(url, title, label, model, conn, client)
        except Exception as exc:
            log.error(f"[fmprc/{label}] Error on {url}: {exc}")


# ── Source: MOFCOM ────────────────────────────────────────────────────────────

def _fetch_mofcom_cms_list(
    tag: str,
    list_url: str,
    api_query: dict,
    link_pattern: re.Pattern,
    client: httpx.Client,
    skip_index_hrefs: bool = True,
) -> list[tuple[str, str]]:
    """
    MOFCOM's whole site (both the English mirror and the Chinese
    www.mofcom.gov.cn pages) runs on a JS CMS whose index pages are just a
    client-side-rendered shell — the article list is fetched by the page's
    own JS from this API-gateway endpoint, which we can call directly and
    skip the browser entirely. Confirmed live on 2026-08-04 for the English
    press-conference index; confirmed again 2026-08-05 that the Chinese
    daily-news-release index (`rcxwfb`) runs the exact same pattern with
    its own `webId`/`tplSetId`/`pageId`/`tagId` — `tagId` is the (literal,
    non-obfuscated) Chinese string the page's own script sends: "信息列表"
    ("information list") for the English press-conference index, "分页列表"
    ("paginated list") for the Chinese daily-news-release index. Every
    other MOFCOM index page (`ldrhd`, `bldhd`, `ztxwfbh`, `sjfzrfb`) likely
    follows the same pattern — check the page's `AuthorizedRead` script tag
    for its own `queryData` if adding one of those next.

    `skip_index_hrefs`: the "index" substring check below exists to filter
    out a link back to the section's own root nav page, which sneaks past
    `link_pattern` on some sections. Set False for a section whose REAL
    content pages all legitimately end in "index.html" themselves (found
    2026-09-03 adding `lxxwfbh`, MOFCOM's regular weekly press conference
    index — every real conference page there is a per-date directory whose
    canonical URL is that directory's own "index.html", e.g.
    ".../swbzklxxwfbh2026n8y27r/index.html" — the blanket substring check
    would otherwise discard every single real link on that page).
    """
    origin  = "/".join(list_url.split("/")[:3])  # e.g. "https://www.mofcom.gov.cn"
    api_url = f"{origin}/api-gateway/jpaas-publish-server/front/page/build/unit"
    log.info(f"[{tag}] {api_url}")

    try:
        time.sleep(REQUEST_SLEEP)
        api_resp = client.get(api_url, params=api_query, timeout=30.0)
        api_resp.raise_for_status()
        list_html = api_resp.json()["data"]["html"]
    except Exception as exc:
        log.error(f"[{tag}] Failed to fetch list via API gateway: {exc}")
        return []

    soup = BeautifulSoup(list_html, "html.parser")
    raw_links = []
    for a in soup.find_all("a", href=link_pattern):
        if skip_index_hrefs and "index" in a.get("href", ""):
            continue
        href  = urljoin(list_url, a["href"])
        title = a.get_text(strip=True)
        if title and len(title) > 10:
            raw_links.append((href, title))
    return raw_links


def scrape_mofcom(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    list_url = "https://english.mofcom.gov.cn/News/PressConference/index.html"
    api_query = {
        "parseType": "bulidstatic",
        "webId":     "6c75aaf37b20474fb2ff8831451b51a0",
        "tplSetId":  "coGxHOTsUcFSE5XydHHhN",
        "pageType":  "column",
        "tagId":     "信息列表",
        "editType":  "null",
        "pageId":    "cbf20136dcfc4e99b91dfd842528e132",
    }
    client = make_client(verify_ssl=False)  # MOFCOM cert untrusted by Python CA bundle

    raw_links = _fetch_mofcom_cms_list(
        "mofcom", list_url, api_query, re.compile(r"/News/PressConference/"), client
    )
    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[mofcom] {len(new_links)} new items")

    for url, title in new_links:
        try:
            process_mofcom_item(url, title, model, conn, client)
        except Exception as exc:
            log.error(f"[mofcom] Error on {url}: {exc}")


# All of MOFCOM's Chinese-language xwfb/* index pages run the identical JS
# CMS pattern (same webId/tplSetId/tagId — confirmed live 2026-08-05 by
# checking each page's own AuthorizedRead <script> tag) with just a
# different `pageId` and URL section slug. `ztxwfbh` (special press
# conferences) came back genuinely empty when checked (valid — MOFCOM
# hasn't published one recently — not a bug), wired up anyway since that
# can change day to day.
_MOFCOM_SECTION_PAGE_IDS = {
    "rcxwfb":  "95d89972d8aa4fcea511701cd0f212d9",  # daily news release
    "ldrhd":   "ee2cfa3f1f534d108378ece1327ee791",  # leadership activity
    "bldhd":   "df9488c25fd44ec58ed25c8fb2454e35",  # dept. leadership activity
    "sjfzrfb": "07abc5c54eee4b7181be2b5be6a1a57f",  # bureau/dept head announcements
    "ztxwfbh": "a4e4443a9fca438f86a447c403158f4e",  # special press conferences
}


def scrape_mofcom_section(
    tag: str,
    section: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    """
    Shared implementation for every MOFCOM xwfb/* Chinese index page — see
    _MOFCOM_SECTION_PAGE_IDS. Added 2026-08-05 after backtest.py showed a
    real past-tracker entry (the "China's Position on the Issue of
    So-called 'Overcapacity'" position paper) was sourced from `rcxwfb`,
    which wasn't being crawled at all before (process_mofcom_item already
    knows how to handle a Chinese www.mofcom.gov.cn article — see its
    docstring — it just never got fed one from here). Reuses
    process_mofcom_item entirely; this function is only the list-discovery
    half, generalized across all 5 sections instead of one-off per section.
    """
    list_url = f"https://www.mofcom.gov.cn/xwfb/{section}/index.html"
    api_query = {
        "parseType": "bulidstatic",
        "webId":     "8f43c7ad3afc411fb56f281724b73708",
        "tplSetId":  "52551ea0e2c14bca8c84792f7aa37ead",
        "pageType":  "column",
        "tagId":     "分页列表",
        "editType":  "null",
        "pageId":    _MOFCOM_SECTION_PAGE_IDS[section],
    }
    client = make_client(verify_ssl=False)

    raw_links = _fetch_mofcom_cms_list(
        tag, list_url, api_query, re.compile(rf"/xwfb/{section}/art/"), client
    )
    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[{tag}] {len(new_links)} new items")

    for url, title in new_links:
        try:
            process_mofcom_item(url, title, model, conn, client)
        except Exception as exc:
            log.error(f"[{tag}] Error on {url}: {exc}")


def scrape_mofcom_daily(model: genai.Client, conn: sqlite3.Connection, doc: Document) -> None:
    scrape_mofcom_section("mofcom_daily", "rcxwfb", model, conn, doc)


def scrape_mofcom_leadership(model: genai.Client, conn: sqlite3.Connection, doc: Document) -> None:
    scrape_mofcom_section("mofcom_leadership", "ldrhd", model, conn, doc)


def scrape_mofcom_dept_leadership(model: genai.Client, conn: sqlite3.Connection, doc: Document) -> None:
    scrape_mofcom_section("mofcom_dept_leadership", "bldhd", model, conn, doc)


def scrape_mofcom_bureau_heads(model: genai.Client, conn: sqlite3.Connection, doc: Document) -> None:
    scrape_mofcom_section("mofcom_bureau_heads", "sjfzrfb", model, conn, doc)


def scrape_mofcom_special_conf(model: genai.Client, conn: sqlite3.Connection, doc: Document) -> None:
    scrape_mofcom_section("mofcom_special_conf", "ztxwfbh", model, conn, doc)


# MOFCOM's regular WEEKLY press conference index — added 2026-09-03. Real
# ground-truth gap, found live: a real Aug 27, 2026 entry (Spokesperson
# Huang Ling on a proposed US tariff hike) turned out to live here, at
# "/xwfbzt/YYYY/swbzklxxwfbhYYYYnMyDr/index.html" — a URL namespace this
# project previously flagged in SOURCES.md as "not yet built" without
# realizing it was reachable through the exact same CMS API-gateway
# pattern as every other MOFCOM section, just with skip_index_hrefs=False
# (see _fetch_mofcom_cms_list's docstring — every real link here itself
# ends in "index.html", unlike the other sections). The date-slug format
# (`n8y27r` = month 8, day 27) is parseable up front, same trick as
# Treasury/USTR/SCIO, so the list walk can stop as soon as it's past this
# run's target start instead of relying on MAX_NEW_ITEMS_PER_RUN alone.
_MOFCOM_LXXWFBH_LINK_RE = re.compile(r"/xwfbzt/\d{4}/swbzklxxwfbh[^/]+/index\.html$")
_MOFCOM_LXXWFBH_DATE_RE = re.compile(r"swbzklxxwfbh(\d{4})n(\d{1,2})y(\d{1,2})r")


def scrape_mofcom_lxxwfbh(model: genai.Client, conn: sqlite3.Connection, doc: Document) -> None:
    tag = "mofcom_lxxwfbh"
    list_url = "https://www.mofcom.gov.cn/xwfb/lxxwfbh/index.html"
    api_query = {
        "parseType": "bulidstatic",
        "webId":     "8f43c7ad3afc411fb56f281724b73708",
        "tplSetId":  "52551ea0e2c14bca8c84792f7aa37ead",
        "pageType":  "column",
        "tagId":     "分页列表",
        "editType":  "null",
        "pageId":    "325f68c1c6a548b299f02ff013d577bd",
    }
    client = make_client(verify_ssl=False)

    raw_links = _fetch_mofcom_cms_list(
        tag, list_url, api_query, _MOFCOM_LXXWFBH_LINK_RE, client,
        skip_index_hrefs=False,
    )

    dated_items: list[tuple[date, str, str]] = []
    for href, title in raw_links:
        m = _MOFCOM_LXXWFBH_DATE_RE.search(href)
        item_date = (
            date(int(m.group(1)), int(m.group(2)), int(m.group(3))) if m
            else _utcnow().date()
        )
        dated_items.append((item_date, href, title))
    dated_items.sort(key=lambda t: t[0], reverse=True)  # newest first, defensively

    new_items: list[tuple[str, str, date]] = []
    for item_date, href, title in dated_items:
        if _RUN_TARGET_START is not None and item_date < _RUN_TARGET_START:
            log.info(f"[{tag}] Reached {item_date}, before this run's "
                     f"target start ({_RUN_TARGET_START}) — stopping.")
            break
        if not is_seen(conn, href):
            new_items.append((href, title, item_date))
    new_items = new_items[:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[{tag}] {len(new_items)} new items")

    for url, title, item_date in new_items:
        try:
            process_mofcom_item(url, title, model, conn, client, known_date=item_date)
        except Exception as exc:
            log.error(f"[{tag}] Error on {url}: {exc}")


def process_mofcom_item(
    url: str,
    title: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
    known_date: date | None = None,
) -> bool:
    """
    Per-item body for a MOFCOM list entry — see process_fmprc_item.

    Handles BOTH the English mirror (english.mofcom.gov.cn, what
    scrape_mofcom's own list-discovery targets) AND raw Chinese
    www.mofcom.gov.cn pages — found via backtest.py that a real past-tracker
    entry (a MOFCOM spokesperson Q&A on the DoD's Section 1286 sanctions
    list) was sourced directly from the Chinese page, not the English
    mirror; SOURCES.md previously assumed the Chinese MOFCOM pages were
    fully redundant with the English one, which this disproves — they're
    not, at least not always. mofcom.gov.cn's Chinese press-conference
    transcripts use generic 问/答 ("Question"/"Answer") labels rather than
    named spokespersons, unlike FMPRC, hence "Answer" in the spokesperson
    set below.

    `known_date`: the date already parsed from the URL by a caller with a
    date-encoding URL scheme (scrape_mofcom_lxxwfbh) — skips re-deriving
    it from the page body here. Left optional so scrape_mofcom's own
    English-mirror discovery (no date in its URLs) and backtest.py's
    direct-URL dispatch both keep working unchanged.
    """
    try:
        resp = fetch(client, url)
        if not resp:
            return False

        main_text = extract_text_from_response(resp)
        cjk = sum(1 for c in main_text if "一" <= c <= "鿿")
        if known_date is not None:
            item_date = datetime(known_date.year, known_date.month, known_date.day)
        else:
            date_m    = re.search(r"(\d{4}-\d{2}-\d{2})", resp.text)
            item_date = (
                datetime.strptime(date_m.group(1), "%Y-%m-%d") if date_m
                else _utcnow()
            )

        raw_zh_text = None
        if cjk > 100:
            if not CHINESE_RELEVANCE_KEYWORDS.search(main_text):
                log.info(f"[mofcom] No relevant keywords in raw text — skipping translation: {url}")
                mark_seen(conn, url)
                return False
            log.info(f"[mofcom] Translating: {url}")
            translated = translate_to_english(model, main_text[:7000])
            exchanges  = parse_qa_from_plaintext(
                translated, {"Spokesperson", "Minister", "Deputy", "Answer"}, model
            )
            work_text  = translated
            raw_zh_text = main_text
        else:
            exchanges = parse_qa(resp.text, {"Spokesperson", "Minister", "Deputy"})
            work_text = main_text
            # chinese_origin=True — MOFCOM is a Chinese-government source
            # same as SCIO; see classify_relevance's docstring for why
            # "does this involve China" is the wrong question here. Found
            # 2026-09-03 chasing the exact same bug on SCIO — this call
            # site had it too, just less visible since a wrong NO here
            # gets silently masked whenever exchanges is empty anyway
            # (finalize_qa_item's own no-exchanges fallback never runs
            # because this branch returns first).
            is_rel, reason = classify_relevance(model, f"{title}\n\n{work_text}", chinese_origin=True)  # full text — see process_release_common's comment
            if not is_rel:
                if not reason.startswith("Keyword pre-filter:"):
                    flag_for_review(url, title, reason)
                mark_seen(conn, url)
                return False

        queued = finalize_qa_item(model, "mofcom", url, item_date, exchanges,
                                   work_text, "MOFCOM press conference", conn,
                                   raw_zh_text=raw_zh_text)
        if queued:
            log.info(f"[mofcom] Queued: {title}")
        return queued

    except Exception as exc:
        log.error(f"[mofcom] Error on {url}: {exc}")
        return False


def process_state_item_by_url(
    url: str,
    title: str,
    date: datetime,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
) -> bool:
    """
    Fetch-by-URL variant of State's per-item processing, for backtest.py —
    which only has a historical URL + known date from a past tracker
    entry, not the original RSS/API item dict scrape_state() normally
    works from. Content quality is comparable (see NOTES.md, 2026-08-04):
    a direct fetch + extract_main_text() vs. the RSS feed's embedded
    content field produce similar-length, similarly clean text.
    """
    try:
        resp = fetch(client, url)
        if not resp:
            return False
        plain = extract_text_from_response(resp)
        return process_release_common("state", url, title, date, plain,
                                       "State Department", model, conn)
    except Exception as exc:
        log.error(f"[state] Error on {url}: {exc}")
        return False


# ── Source: State Dept ────────────────────────────────────────────────────────

def scrape_state(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    # Route renamed server-side at some point after this was first written:
    # the custom post type is now `state_press_release` (was
    # `press_releases`) — confirmed live 2026-08-05 by walking
    # state.gov/wp-json/'s own route discovery document. The RSS fallback
    # below was masking this (it's why the source kept working at all), but
    # fixing the real endpoint means results aren't solely dependent on the
    # RSS feed's own completeness/latency.
    api_url = (
        "https://www.state.gov/wp-json/wp/v2/state_press_release"
        "?per_page=50&orderby=date&_fields=id,date,title,link,content,excerpt"
    )
    rss_url = "https://www.state.gov/rss-feed/press-releases/feed/"
    client  = make_client()

    log.info("[state] Fetching WP API")
    items: list[dict] = []
    resp = fetch(client, api_url)
    if resp:
        try:
            items = resp.json()
        except Exception:
            items = []

    if not items:
        log.warning("[state] WP API failed — trying RSS")
        resp_rss = fetch(client, rss_url)
        if resp_rss:
            items = parse_rss(resp_rss.text)

    if not items:
        log.error("[state] Both endpoints failed")
        return

    new_items = [it for it in items if not is_seen(conn, item_url(it))]
    log.info(f"[state] {len(new_items)} new items")

    for it in new_items:
        url = item_url(it)
        try:
            title = item_title(it)
            plain = BeautifulSoup(item_content(it), "html.parser").get_text()
            process_release_common("state", url, title, item_date(it), plain,
                                    "State Department", model, conn)

        except Exception as exc:
            log.error(f"[state] Error on {url}: {exc}")


def _resolve_pdf_stub(
    client: httpx.Client, html: str, plain: str, base_url: str, min_len: int = 300,
) -> str:
    """
    Some whitehouse.gov "releases" pages are just a landing/stub page with
    a "Download" link to the actual PDF report — e.g. "The Great
    Transshipment Scam," a real past-tracker entry (a 25-page OTMP report)
    whose page text is only "Releases | The Great Transshipment Scam |
    The White House | ... | Download" (107 chars), not the report itself.
    That was silently dropped: `classify_relevance`/`extract_key_
    paragraphs` correctly found nothing relevant in text that was never
    the real content. Found live, 2026-09-01. Detected via a suspiciously
    short extracted-text length; if a same-page PDF link exists, follow
    and extract THAT instead. Reusable — wire in wherever another source
    turns out to use the same stub-page-plus-PDF pattern.
    """
    if len(plain) >= min_len:
        return plain
    soup = BeautifulSoup(html, "html.parser")
    pdf_link = soup.find("a", href=re.compile(r"\.pdf($|\?)", re.IGNORECASE))
    if not pdf_link:
        return plain
    pdf_url = urljoin(base_url, pdf_link["href"])
    log.info(f"[pdf-stub] Landing page too short ({len(plain)} chars) — following PDF link: {pdf_url}")
    pdf_resp = fetch(client, pdf_url)
    if not pdf_resp:
        return plain
    return extract_text_from_response(pdf_resp)


def process_whitehouse_item_by_url(
    url: str,
    title: str,
    date: datetime,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
) -> bool:
    """Fetch-by-URL variant for backtest.py — see
    process_state_item_by_url's docstring, same reasoning."""
    try:
        resp = fetch(client, url)
        if not resp:
            return False
        plain = extract_text_from_response(resp)
        plain = _resolve_pdf_stub(client, resp.text, plain, url)
        return process_release_common("whitehouse", url, title, date, plain,
                                       "White House", model, conn)
    except Exception as exc:
        log.error(f"[whitehouse] Error on {url}: {exc}")
        return False


# ── Source: White House ───────────────────────────────────────────────────────

def scrape_whitehouse(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    rss_url = "https://www.whitehouse.gov/news/feed/"
    client  = make_client()

    log.info(f"[whitehouse] {rss_url}")
    resp = fetch(client, rss_url)
    if not resp:
        log.error("[whitehouse] Failed to fetch RSS")
        return

    items     = parse_rss(resp.text)
    new_items = [it for it in items if not is_seen(conn, item_url(it))]
    log.info(f"[whitehouse] {len(new_items)} new items")

    for it in new_items:
        url = item_url(it)
        try:
            title = item_title(it)
            raw_content = item_content(it)
            plain = BeautifulSoup(raw_content, "html.parser").get_text()
            plain = _resolve_pdf_stub(client, raw_content, plain, url)
            process_release_common("whitehouse", url, title, item_date(it), plain,
                                    "White House", model, conn)

        except Exception as exc:
            log.error(f"[whitehouse] Error on {url}: {exc}")


# ── Source: Treasury ──────────────────────────────────────────────────────────

def scrape_treasury(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    list_url = "https://home.treasury.gov/news/press-releases"
    client   = make_client()

    log.info(f"[treasury] {list_url}")
    resp = fetch(client, list_url, retries=4)
    if not resp:
        log.error("[treasury] Failed after 4 retries — skipping")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    # Every real press-release row carries its own <time datetime="...">
    # right alongside the link (in the same small container div) — the
    # exact same structure confirmed on USTR, strictly newest-first.
    # Requiring the link to match a real content slug (two letters +
    # digits, e.g. "sb0620", "jy1234") rather than a broad "/news/press-
    # releases/" substring also filters out nav-menu links (a category
    # page like ".../statements-remarks" has no such slug) and — found
    # live, 2026-09-02 — pagination controls ("Page 2", "Next page"),
    # which otherwise inherited a bogus date from an unrelated ancestor
    # div and could have looked like a real, very-recent item.
    raw_items: list[tuple[date, str, str]] = []
    for a in soup.find_all("a", href=re.compile(r"/news/press-releases/[a-z]{2}\d+/?$")):
        href = urljoin(list_url, a["href"])
        container = a.find_parent("div")
        time_tag = container.find("time") if container else None
        if not time_tag or not time_tag.get("datetime"):
            continue
        item_date = datetime.fromisoformat(time_tag["datetime"].replace("Z", "+00:00")).date()
        title = a.get_text(strip=True)
        if title:
            raw_items.append((item_date, href, title))

    # De-dupe (a real release can appear in more than one listing widget
    # on the page) while preserving the newest-first order already confirmed.
    seen_hrefs: set[str] = set()
    new_links: list[tuple[str, str, date]] = []
    for item_date, href, title in raw_items:
        if href in seen_hrefs:
            continue
        seen_hrefs.add(href)
        if _RUN_TARGET_START is not None and item_date < _RUN_TARGET_START:
            log.info(f"[treasury] Reached {item_date}, before this run's "
                     f"target start ({_RUN_TARGET_START}) — stopping (list "
                     f"is newest-first, so everything after this is even "
                     f"older).")
            break
        if not is_seen(conn, href):
            new_links.append((href, title, item_date))
    new_links = new_links[:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[treasury] {len(new_links)} new items")

    for url, title, item_date in new_links:
        try:
            process_treasury_item(url, title, model, conn, client, known_date=item_date)
        except Exception as exc:
            log.error(f"[treasury] Error on {url}: {exc}")


def process_treasury_item(
    url: str,
    title: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
    known_date: date | None = None,
) -> bool:
    """Per-item body for a Treasury list entry — see process_fmprc_item.
    `known_date`: the date already read straight off the listing page
    (see scrape_treasury) — skips re-deriving it from a regex search over
    the fetched page text when already known."""
    try:
        resp = fetch(client, url, retries=4)
        if not resp:
            return False

        plain = extract_text_from_response(resp)
        if known_date is not None:
            item_dt = datetime(known_date.year, known_date.month, known_date.day)
        else:
            date_m = re.search(
                r"(January|February|March|April|May|June|July|August|"
                r"September|October|November|December)\s+\d{1,2},\s+\d{4}",
                plain,
            )
            item_dt = (
                datetime.strptime(date_m.group(0), "%B %d, %Y") if date_m
                else _utcnow()
            )
        return process_release_common("treasury", url, title, item_dt, plain,
                                       "Treasury Department", model, conn)

    except Exception as exc:
        log.error(f"[treasury] Error on {url}: {exc}")
        return False


# ── Source: USTR ──────────────────────────────────────────────────────────────

def scrape_ustr(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    list_url = "https://ustr.gov/about-us/policy-offices/press-office/press-releases"
    client   = make_client()

    log.info(f"[ustr] {list_url}")
    resp = fetch(client, list_url)
    if not resp:
        log.error("[ustr] Failed to fetch list")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    # Each listing row already carries its own publish date directly — a
    # Drupal "views-row" with a <time datetime="..."> right alongside the
    # link — confirmed live, 2026-09-02, strictly newest-first. This lets
    # the list-walk stop the moment it passes this run's target range,
    # with NO extra fetches (the date doesn't require visiting the
    # article at all) and no arbitrary item-count cap needed to bound a
    # date-scoped request — MAX_NEW_ITEMS_PER_RUN below is now only a
    # safety ceiling for the OTHER case (no target range set at all,
    # e.g. a plain `--source ustr` with no dates), not the thing deciding
    # how far back a dated request looks. See NOTES.md — this replaces
    # the earlier fix attempt (raising the cap to 150), which paid a
    # real ongoing cost for what's actually a one-time backlog problem;
    # this fix costs nothing extra for a normal week and correctly
    # reaches an old week's content regardless of backlog size.
    raw_items: list[tuple[date, str, str]] = []
    for row in soup.find_all("div", class_="views-row"):
        time_tag = row.find("time")
        link = row.find("a", href=re.compile(r"/about/policy-offices/press-office/press-releases/\d{4}/"))
        if not time_tag or not link or not time_tag.get("datetime"):
            continue
        item_date = datetime.fromisoformat(time_tag["datetime"].replace("Z", "+00:00")).date()
        raw_items.append((item_date, urljoin("https://ustr.gov", link["href"]), link.get_text(strip=True)))

    new_links: list[tuple[str, str, date]] = []
    for item_date, href, title in raw_items:
        if _RUN_TARGET_START is not None and item_date < _RUN_TARGET_START:
            log.info(f"[ustr] Reached {item_date}, before this run's target "
                     f"start ({_RUN_TARGET_START}) — stopping (list is "
                     f"newest-first, so everything after this is even older).")
            break
        if not is_seen(conn, href):
            new_links.append((href, title, item_date))
    new_links = new_links[:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[ustr] {len(new_links)} new items")

    for url, title, item_date in new_links:
        try:
            process_ustr_item(url, title, model, conn, client, known_date=item_date)
        except Exception as exc:
            log.error(f"[ustr] Error on {url}: {exc}")


def process_ustr_item(
    url: str,
    title: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
    known_date: date | None = None,
) -> bool:
    """Per-item body for a USTR list entry — see process_fmprc_item.
    `known_date`: the date already read straight off the listing page
    (see scrape_ustr) — skips re-deriving it from a regex search over
    the fetched page text when already known."""
    try:
        resp = fetch(client, url)
        if not resp:
            return False

        plain = extract_text_from_response(resp)
        if known_date is not None:
            item_dt = datetime(known_date.year, known_date.month, known_date.day)
        else:
            date_m = re.search(r"(\d{4}-\d{2}-\d{2})", resp.text)
            item_dt = (
                datetime.strptime(date_m.group(1), "%Y-%m-%d") if date_m
                else _utcnow()
            )
        return process_release_common("ustr", url, title, item_dt, plain, "USTR", model, conn)

    except Exception as exc:
        log.error(f"[ustr] Error on {url}: {exc}")
        return False


# ── Source: Department of Defense (war.gov) — DISABLED, 2026-09-02 ──────────
#
# Removed from active dispatch (not in SOURCES, not called from main()) —
# war.gov's actual article content is blocked at the infrastructure level
# for any non-browser client, confirmed after real, thorough investigation:
# a bare 403 from the homepage AND every article page, identical across
# plain httpx/curl, a real headless-Playwright browser (no JS challenge to
# solve — nothing to bypass), the OLD defense.gov domain, and multiple
# user-agent spoofs (including Googlebot/Bingbot); confirmed from BOTH this
# project's dev environment and the user's own home network independently.
# Checked for a substitute mirror too — DVIDS (the official DoD public
# affairs platform, itself unblocked) carries related but NOT the same
# content; the RSS feed's own description fields are too short to stand in
# for the real article; Wayback Machine's coverage of individual articles
# is incomplete; the department's own X accounts (DOWResponse, SecWar)
# are real and active but mostly retweets/general content, not a
# systematic feed of these specific releases. See NOTES.md for the full
# investigation trail. Left in the file, unreachable, rather than deleted,
# in case war.gov's policy ever changes — everything below still works
# exactly as built (the RSS-based discovery, the stop-after-first-403
# fix) if this is ever re-enabled.

def scrape_wardept(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    # The /News/Releases/ listing page is JS-rendered behind an Akamai bot
    # check that a plain fetch (and even a headless-Playwright session — bare
    # 403, no JS challenge to solve, so it's an IP/ASN-level block rather than
    # a fingerprinting one) can't get past from a datacenter IP. war.gov's own
    # RSS backend (DesktopModules/ArticleCS/RSS.ashx) is NOT behind that
    # block and returns clean XML, so we use it instead of scraping the page
    # or driving a browser at all. ContentType=9 is "News/Releases/Release/..."
    # (press releases); ContentType=11 is "News/Speeches/..." (also relevant —
    # e.g. Alvaro Smith's China-relations remarks are published there, not
    # under Releases). Confirmed live on 2026-08-04.
    rss_urls = {
        "release": "https://www.war.gov/DesktopModules/ArticleCS/RSS.ashx?ContentType=9&Site=945&max=50",
        "speech":  "https://www.war.gov/DesktopModules/ArticleCS/RSS.ashx?ContentType=11&Site=945&max=50",
    }
    client = make_client()

    raw_links: list[tuple[str, str]] = []
    for kind, rss_url in rss_urls.items():
        log.info(f"[wardept] {rss_url}")
        resp = fetch(client, rss_url)
        if not resp:
            log.warning(f"[wardept] Failed to fetch {kind} RSS")
            continue
        for it in parse_rss(resp.text):
            raw_links.append((item_url(it), item_title(it)))

    new_links = [(u, t) for u, t in raw_links if u and not is_seen(conn, u)][:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[wardept] {len(new_links)} new items")

    if not new_links:
        # Note for whoever reads the run log: if this is consistently 0 even
        # when the site clearly has new content, check whether war.gov is
        # itself reachable at all from this machine (see NOTES.md — article
        # pages were 403'd from the sandbox IP this was developed on; that's
        # an IP-reputation issue with war.gov's WAF, not a bug in this code).
        return
    for url, title in new_links:
        process_wardept_item(url, title, model, conn, client)
        # war.gov's Akamai block, when active, is site-wide and persistent
        # for the rest of THIS run — confirmed live, 2026-09-02, even the
        # homepage 403s for a non-browser client (see NOTES.md). Once one
        # item comes back 403, every remaining item will too, so stop here
        # instead of burning a request (+ REQUEST_SLEEP) on each of up to
        # 30 items just to watch them all fail the same way.
        if _LAST_FETCH_STATUS == 403:
            log.warning("[wardept] Got 403 — war.gov's block is active this "
                        "run, stopping here instead of repeating the same "
                        "failure on the remaining items.")
            break


def process_wardept_item(
    url: str,
    title: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
) -> bool:
    """Per-item body for a war.gov list entry — see process_fmprc_item."""
    try:
        resp = fetch(client, url)
        if not resp:
            return False

        plain  = extract_text_from_response(resp)
        date_m = re.search(
            r"(January|February|March|April|May|June|July|August|"
            r"September|October|November|December)\s+\d{1,2},\s+\d{4}",
            plain,
        )
        date  = (
            datetime.strptime(date_m.group(0), "%B %d, %Y") if date_m
            else _utcnow()
        )
        return process_release_common("wardept", url, title, date, plain,
                                       "Department of Defense", model, conn)

    except Exception as exc:
        log.error(f"[wardept] Error on {url}: {exc}")
        return False


# ── Source: MFA leadership speeches/activity (Chinese-only pages) ───────────

_MFA_LEADERSHIP_LINK_RE = re.compile(r"^\./\d{6}/t\d+_\d+\.shtml$")

def _mfa_leadership_should_skip_translate(plain_cn: str) -> bool:
    """
    Free pre-filter for process_mfa_leadership_item — plain keyword gate
    against CHINESE_RELEVANCE_KEYWORDS, same as FMPRC/MOFCOM/MND. Pulled
    out on its own so it's directly testable without a live fetch.

    Earlier today (2026-09-02) this was a length-OR-keyword hybrid
    instead, specifically to avoid dropping a real past-tracker entry (a
    2026-07-28 Wang Yi/Global Development Initiative item with no keyword
    hit) that had been cited as evidence a pure keyword filter was unsafe
    for this source. The user confirmed that entry was itself a human
    coding error in the original tracker — it shouldn't have been
    included at all, since it has no actual US-China relations content —
    which removes the evidence for treating this source any differently
    from the other Chinese sources. Reverted to a plain keyword gate.
    """
    return not CHINESE_RELEVANCE_KEYWORDS.search(plain_cn)


def scrape_mfa_leadership(
    list_url: str,
    label: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
    client: httpx.Client,
) -> None:
    """
    MFA leadership speeches / leadership activity — Chinese-only, no EN
    mirror. Usually one-way statements/readouts rather than Q&A transcripts
    (finalize_release_item() re-checks and reroutes if a given page turns
    out to actually be a Q&A). Added 2026-08-04; live-tested against both
    list pages.
    """
    log.info(f"[mfa_leadership/{label}] {list_url}")
    resp = fetch(client, list_url)
    if not resp:
        log.error(f"[mfa_leadership/{label}] Failed to fetch list")
        return

    soup = BeautifulSoup(resp.text, "html.parser")
    raw_links = []
    for a in soup.find_all("a", href=_MFA_LEADERSHIP_LINK_RE):
        href  = urljoin(list_url, a["href"])
        title = a.get_text(strip=True)
        if title:
            raw_links.append((href, title))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[mfa_leadership/{label}] {len(new_links)} new items")

    for url, title in new_links:
        try:
            process_mfa_leadership_item(url, title, label, model, conn, client)
        except Exception as exc:
            log.error(f"[mfa_leadership/{label}] Error on {url}: {exc}")


def process_mfa_leadership_item(
    url: str,
    title: str,
    label: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
) -> bool:
    """Per-item body for an MFA leadership list entry — see
    process_fmprc_item."""
    tag = f"mfa_leadership/{label}"
    try:
        # Date is parsed straight from the URL, before any fetch at all —
        # this is the one thing we can know about an item for free. Added
        # 2026-09-02: if this run has a target range (see queue_entry()'s
        # comment) and this item's date falls outside it, skip the
        # network fetch AND the expensive translate call entirely, not
        # just the eventual doc-write (queue_entry() would have dropped
        # it anyway, but only after paying for the fetch+translate first).
        # Left unseen (not marked) so a future run whose target actually
        # covers this date still finds and correctly includes it.
        date_m = re.search(r"t(\d{8})_", url)
        item_date = (
            datetime.strptime(date_m.group(1), "%Y%m%d") if date_m
            else _utcnow()
        )
        if _RUN_TARGET_START is not None:
            d = item_date.date()
            if not (_RUN_TARGET_START <= d <= _RUN_TARGET_END):
                log.info(f"[{tag}] {d} is outside this run's target range "
                         f"({_RUN_TARGET_START}..{_RUN_TARGET_END}) — "
                         f"skipping without fetching: {url}")
                return False

        resp = fetch(client, url)
        if not resp:
            return False

        plain_cn = extract_text_from_response(resp)
        cjk = sum(1 for c in plain_cn if "一" <= c <= "鿿")
        if cjk < 50:
            log.info(f"[{tag}] Not enough Chinese text — skipping: {url}")
            mark_seen(conn, url)
            return False

        # Free keyword pre-filter, same as FMPRC/MOFCOM/MND — added
        # 2026-09-02 per user request, since this was previously the one
        # Chinese source that translated EVERY item unconditionally (the
        # slowest, most expensive call), on the theory that this source's
        # real editorial bar ("substantive top-leadership activity") was
        # broader than a topic keyword. That theory rested on one real
        # past-tracker entry (2026-07-28, Wang Yi/Global Development
        # Initiative) that matched no keyword — the user has since
        # confirmed that entry was itself a human coding error in the
        # original tracker, not a genuine editorial exception, so the
        # keyword gate applies here exactly like every other Chinese
        # source. See _mfa_leadership_should_skip_translate()'s own
        # docstring for the fuller history.
        if _mfa_leadership_should_skip_translate(plain_cn):
            log.info(f"[{tag}] No topic keyword — skipping without translating: {url}")
            mark_seen(conn, url)
            return False

        log.info(f"[{tag}] Translating: {url}")
        translated = translate_to_english(model, plain_cn[:7000])

        # raw_zh_text=plain_cn — see extract_key_paragraphs()'s docstring
        # for why the old general=True bar ("substantive Chinese
        # leadership activity", no US mention required) was retired
        # 2026-09-03: it let a Wang Yi/India-border readout with zero US
        # mentions through, purely because the free keyword pre-filter
        # above matched "Tibet" as a shared topic. raw_zh_text routes this
        # through the same narrow "does this paragraph name the US" gate
        # fmprc/mofcom/mnd already use for their own release-shaped content.
        queued = finalize_release_item(model, tag, url, item_date, translated,
                                        f"MFA {label}", conn, raw_zh_text=plain_cn)
        if queued:
            log.info(f"[{tag}] Queued: {title}")
        return queued

    except Exception as exc:
        log.error(f"[{tag}] Error on {url}: {exc}")
        return False


# ── Source: State Council Information Office (SCIO), English pressroom ─────

_SCIO_LINK_RE = re.compile(r"/pressroom/\d{4}-\d{2}/\d{2}/content_\d+\.html")


_SCIO_PAGE_LIMIT = 10  # safety cap on how many pages deep to paginate per
                        # list_url — well beyond what a normal week (even a
                        # delayed catch-up run) should ever need; just a
                        # backstop against an infinite loop if a list_url's
                        # pagination ever behaves unexpectedly.


def scrape_scio(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    """
    SCIO English pressroom — already in English, so this follows the same
    shape as scrape_state/scrape_treasury/scrape_ustr rather than needing
    translation. Two listing pages per the links doc.

    Paginates each list_url (node_8020819_2.html, _3.html, ...) with the
    same date-aware early-stop as scrape_treasury/scrape_ustr, added
    2026-09-03 after a real live run missed two real ground-truth entries
    (Aug 26 and Aug 28) that were sitting in plain sight on PAGE 2 of
    node_8020819 — this function previously only ever fetched page 1 of
    each list_url. node_8020819 in particular is a fast-moving general
    feed (30 items can cover as little as ~5 days), so page 1 alone isn't
    enough once a run is even a few days behind. Each item's date comes
    straight from its own URL (`/pressroom/YYYY-MM/DD/content_...html`) —
    no extra fetch needed to know it, same trick as MFA leadership's
    URL-embedded date.
    """
    list_urls = [
        "http://english.scio.gov.cn/pressroom/node_8020819.html",
        "http://english.scio.gov.cn/pressroom/node_8020805.html",
    ]
    client = make_client()

    raw_items: list[tuple[date, str, str]] = []
    seen_hrefs: set[str] = set()
    for base_url in list_urls:
        stem = base_url.rsplit(".html", 1)[0]
        for page in range(1, _SCIO_PAGE_LIMIT + 1):
            list_url = base_url if page == 1 else f"{stem}_{page}.html"
            log.info(f"[scio] {list_url}")
            resp = fetch(client, list_url)
            if not resp:
                log.error(f"[scio] Failed to fetch {list_url}")
                break

            soup = BeautifulSoup(resp.text, "html.parser")
            page_items: list[tuple[date, str, str]] = []
            for a in soup.find_all("a", href=_SCIO_LINK_RE):
                href = urljoin(list_url, a["href"])
                title = a.get_text(strip=True)
                if not title or len(title) <= 10 or href in seen_hrefs:
                    continue
                seen_hrefs.add(href)
                date_m = re.search(r"/pressroom/(\d{4})-(\d{2})/(\d{2})/", href)
                item_date = (
                    date(int(date_m.group(1)), int(date_m.group(2)), int(date_m.group(3)))
                    if date_m else _utcnow().date()
                )
                page_items.append((item_date, href, title))

            if not page_items:
                # Either an empty/duplicate page or we've exhausted this
                # list_url's real pagination (some sites keep serving the
                # last page's content past its actual final page number).
                break
            raw_items.extend(page_items)

            oldest_on_page = min(d for d, _, _ in page_items)
            if _RUN_TARGET_START is not None and oldest_on_page < _RUN_TARGET_START:
                log.info(f"[scio] Reached {oldest_on_page} on page {page} of "
                         f"{base_url}, before this run's target start "
                         f"({_RUN_TARGET_START}) — stopping this list_url "
                         f"(list is newest-first, so everything after this "
                         f"is even older).")
                break

    new_items: list[tuple[str, str, date]] = []
    for item_date, href, title in raw_items:
        if _RUN_TARGET_START is not None and item_date < _RUN_TARGET_START:
            continue
        if not is_seen(conn, href):
            new_items.append((href, title, item_date))
    new_items = new_items[:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[scio] {len(new_items)} new items")

    for url, title, item_date in new_items:
        try:
            process_scio_item(url, title, model, conn, client, known_date=item_date)
        except Exception as exc:
            log.error(f"[scio] Error on {url}: {exc}")


def process_scio_item(
    url: str,
    title: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
    known_date: date | None = None,
) -> bool:
    """Per-item body for an SCIO list entry — see process_fmprc_item.

    `known_date`: the date already parsed from the URL by scrape_scio's
    list walk — skips re-deriving it here. Left optional (re-derives from
    `url` if omitted) so backtest.py can keep calling this directly with
    just a URL, same as every other process_*_item's `known_date` param."""
    try:
        resp = fetch(client, url)
        if not resp:
            return False

        plain = extract_text_from_response(resp)
        if known_date is not None:
            item_date = datetime(known_date.year, known_date.month, known_date.day)
        else:
            date_m = re.search(r"/pressroom/(\d{4})-(\d{2})/(\d{2})/", url)
            item_date = (
                datetime(int(date_m.group(1)), int(date_m.group(2)), int(date_m.group(3)))
                if date_m else _utcnow()
            )
        return process_release_common("scio", url, title, item_date, plain,
                                       "State Council Information Office", model, conn,
                                       chinese_origin=True)

    except Exception as exc:
        log.error(f"[scio] Error on {url}: {exc}")
        return False


# ── Source: Ministry of National Defense (mod.gov.cn) ───────────────────────

_MND_LINK_RE = re.compile(r"/gfbw/xwfyr/(yzxwfb|lxjzh_246940)/\d+\.html")
# \w, not A-Za-z — same accented-name fix as bug #13 (found on a State
# Dept transcript with "PRESIDENT ARÉVALO"), applied here preventively:
# this normalizes "Label:text" spacing to "Label: text" BEFORE _QA_RE ever
# sees it, and _QA_RE requires `:\s+` (whitespace after the colon) to
# match at all — if an accented foreign name (a defense minister quoted
# in an MND transcript, say) broke THIS regex, the added space would
# never happen, and the downstream label match would silently fail too:
# the same "invisible, not misclassified" cascade bug #13 already caused
# once, just one step earlier in the pipeline. Not yet observed live for
# this specific regex — applied on the same reasoning that already proved
# out 4 times elsewhere in this file today, not waiting for a live
# failure to confirm it.
_ENGLISH_COLON_RE = re.compile(r"^([A-Z][\w .]{1,30}):(?=\S)")


def _looks_english(text: str) -> bool:
    """True if `text` has essentially no CJK — i.e. it's the English half of
    mod.gov.cn's bilingual press-conference pages (see scrape_mnd)."""
    if len(text.strip()) < 4:
        return False
    cjk = sum(1 for c in text if "一" <= c <= "鿿")
    return cjk < 3


def scrape_mnd(
    model: genai.Client,
    conn: sqlite3.Connection,
    doc: Document,
) -> None:
    """
    Ministry of National Defense weekly press conferences. mod.gov.cn
    publishes these bilingually IN THE SAME PAGE — each Chinese Q/A
    paragraph is immediately followed by its official English translation
    (e.g. "记者：...” then “Question:..."; "蒋斌：..." then "Jiang Bin:...").
    So unlike FMPRC/MOFCOM we skip our own Gemini translation and use MND's
    own English paragraphs directly when they're present — more accurate
    (official wording) and cheaper. Falls back to translating the Chinese
    if a page turns out not to have an English half. Added 2026-08-04, live
    verified against http://www.mod.gov.cn/gfbw/xwfyr/yzxwfb/16474241.html.
    """
    list_urls = [
        "http://www.mod.gov.cn/gfbw/xwfyr/yzxwfb/index.html",
        "http://www.mod.gov.cn/gfbw/xwfyr/lxjzh_246940/index.html",
    ]
    client = make_client(verify_ssl=False)  # mod.gov.cn cert untrusted by Python CA bundle

    raw_links: list[tuple[str, str]] = []
    for list_url in list_urls:
        log.info(f"[mnd] {list_url}")
        resp = fetch(client, list_url)
        if not resp:
            log.error(f"[mnd] Failed to fetch {list_url}")
            continue
        soup = BeautifulSoup(resp.text, "html.parser")
        for a in soup.find_all("a", href=_MND_LINK_RE):
            href  = urljoin(list_url, a["href"])
            title = a.get_text(strip=True)
            if title:
                raw_links.append((href, title))

    new_links = [(u, t) for u, t in raw_links if not is_seen(conn, u)][:MAX_NEW_ITEMS_PER_RUN]
    log.info(f"[mnd] {len(new_links)} new items")

    for url, title in new_links:
        try:
            process_mnd_item(url, title, model, conn, client)
        except Exception as exc:
            log.error(f"[mnd] Error on {url}: {exc}")


def process_mnd_item(
    url: str,
    title: str,
    model: genai.Client,
    conn: sqlite3.Connection,
    client: httpx.Client,
) -> bool:
    """Per-item body for an MND list entry — see process_fmprc_item."""
    try:
        resp = fetch(client, url)
        if not resp:
            return False

        soup = BeautifulSoup(resp.text, "html.parser")
        paragraphs = [p.get_text(separator=" ", strip=True) for p in soup.find_all("p")]
        paragraphs = [
            p for p in paragraphs
            if p and "版权声明" not in p and "站点地图" not in p
        ]

        english_paras = [_ENGLISH_COLON_RE.sub(r"\1: ", p) for p in paragraphs if _looks_english(p)]

        raw_zh_text = None
        if len(english_paras) >= 2:
            exchanges = _build_exchanges(english_paras, MND_SPOKESPERSONS)
            work_text = "\n\n".join(english_paras)
        else:
            chinese_paras = [p for p in paragraphs if not _looks_english(p)]
            plaintext = "\n".join(chinese_paras)
            if not CHINESE_RELEVANCE_KEYWORDS.search(plaintext):
                log.info(f"[mnd] No relevant keywords in raw text — skipping translation: {url}")
                mark_seen(conn, url)
                return False
            log.info(f"[mnd] No English half found, translating: {url}")
            translated = translate_to_english(model, plaintext[:7000])
            exchanges  = parse_qa_from_plaintext(translated, MND_SPOKESPERSONS, model)
            work_text  = translated
            raw_zh_text = plaintext

        date_m = re.search(r"(\d{4})-(\d{2})-(\d{2})", title) or re.search(
            r"(\d{4})-(\d{2})-(\d{2})", "\n".join(paragraphs)
        )
        date = datetime.strptime(date_m.group(0), "%Y-%m-%d") if date_m else _utcnow()

        queued = finalize_qa_item(model, "mnd", url, date, exchanges,
                                   work_text, "Ministry of National Defense", conn,
                                   raw_zh_text=raw_zh_text)
        if queued:
            log.info(f"[mnd] Queued: {title}")
        return queued

    except Exception as exc:
        log.error(f"[mnd] Error on {url}: {exc}")
        return False


# ── Source: X (Twitter) ──────────────────────────────────────────────────────
#
# Added 2026-09-01: previously "blocked" (see SOURCES.md) — unauthenticated
# scraping of x.com is actively blocked and against ToS, and the official
# API's old subscription tiers (Free = no read access at all; Basic ~$200/
# month) made this a non-starter for a hobby-scale weekly tracker. X moved
# to pay-per-use billing in Feb 2026 (~$0.005/post read — confirmed via
# docs.x.com, re-checked 2026-09-02), which is cheap enough to actually use
# here. User provided X_API_KEY (an App-only OAuth 2.0 Bearer token) and
# pay-per-use credit.
#
# REDESIGNED 2026-09-02 around X's real search-based cost/matching model
# (see NOTES.md for the full story): the original per-account design called
# GET /2/users/:id/tweets once per account, paying for and downloading
# EVERY tweet from every account, then discarding most as irrelevant only
# AFTER paying for it — confirmed live, a first real run cost $1.735 (94%
# of that run's total spend) fetching 329 tweets this way. Research into
# X's actual billing model found the fix: "a search that returns 20 posts
# is billed as 20 post reads" — the SAME per-post price as a timeline
# fetch, but a search query can combine a keyword filter AND multiple
# accounts into ONE call, so X only returns (and bills for) tweets that
# already match. This reuses the exact keyword-recall tradeoff every OTHER
# source in this pipeline already accepts for cost reasons — just applied
# one step earlier (before paying for the tweet, not after).
#
# Two group queries per run instead of one call per account: non-PRC
# accounts filtered by CHINA-relevant keywords (their bar: "does this
# mention China"), PRC accounts filtered by US-mention keywords (their
# bar: "does this mention the US" — see _PRC_X_ACCOUNTS's own comment for
# why that's the opposite test). A useful side effect: X's search "from:"
# operator takes a plain USERNAME, not a numeric user ID, so the old
# username->user-ID lookup step (a separate billed "user read" per
# account) is gone entirely — one less cost, one less thing to cache.
#
# Cost discipline that remains: (1) `since_id` (X's own recommended
# incremental-poll parameter) — now tracked per SEARCH GROUP, not per
# account, since a combined query has one shared result stream — so an
# already-seen tweet is never re-billed on the next run, on top of our own
# is_seen() dedup; (2) only the SOURCES.md "normal" tier is scraped by
# default (11 accounts) — the 7 "less important" accounts are listed but
# not polled, to control cost; flip X_INCLUDE_LESS_IMPORTANT to add them.

X_STATE_PATH = os.path.join(DATA_DIR, "x_accounts_state.json")

# Baked directly into the X search query — see the module comment above.
# Deliberately mirrors the *shape* of RELEVANCE_KEYWORDS/CHINESE_RELEVANCE_
# KEYWORDS (same terms this pipeline already trusts as a relevance signal
# elsewhere), not a separately-invented list, so this doesn't introduce a
# new, differently-tuned notion of "relevant" just for X. Quoted phrases
# use X search's own exact-phrase syntax.
_X_CHINA_SEARCH_TERMS = (
    'china OR chinese OR beijing OR "xi jinping" OR taiwan OR "hong kong" '
    'OR tariff OR tariffs OR "trade war" OR sanctions OR "export control" '
    'OR "export controls" OR huawei OR tiktok OR "rare earth"'
)
_X_US_SEARCH_TERMS = (
    'US OR "united states" OR america OR washington OR trump OR rubio '
    'OR bessent OR vance OR lutnick'
)

# search-group key ("china"/"prc") -> prospective new since_id, staged by
# scrape_x() and only actually persisted by flush_pending_entries() once
# the doc write it corresponds to has succeeded — see scrape_x()'s
# docstring for why. Keyed by GROUP now, not by username — a combined
# search query has one shared result stream per group, not one per
# account.
_PENDING_X_SINCE_IDS: dict[str, str] = {}

X_ACCOUNTS_NORMAL = [
    "ChineseEmbinUS", "SpoxCHN_LinJian", "SpoxCHN_MaoNing",
    "RapidResponse47", "JDVance", "StevenCheung47", "WhiteHouse",
    "PressSec", "SecScottBessent", "realDonaldTrump", "SecRubio",
]
X_ACCOUNTS_LESS_IMPORTANT = [
    "USTradeRep", "DOWResponse", "CIADirector", "SecWar",
    "StephenM", "DNIGabbard", "howardlutnick",
]

# @username -> the bold speaker label a tweet-sourced entry's body
# paragraph gets, matching the real past trackers exactly (e.g. "Chinese
# Embassy: American teacher...", "Rapid Response 47: .@SecScottBessent:
# ...", "President Trump: FBI Shuts Down..."). Added 2026-09-02 per user
# request. Falls back to the bare username for any account not listed
# here (e.g. a future addition) rather than failing — see
# process_x_tweet()'s use of this.
_X_ACCOUNT_DISPLAY_NAMES = {
    "ChineseEmbinUS":  "Chinese Embassy",
    "SpoxCHN_LinJian": "Lin Jian",
    "SpoxCHN_MaoNing": "Mao Ning",
    "RapidResponse47": "Rapid Response 47",
    "JDVance":         "Vice President Vance",
    "StevenCheung47":  "Steven Cheung",
    "WhiteHouse":      "White House",
    "PressSec":        "Press Secretary",
    "SecScottBessent": "Secretary Bessent",
    "realDonaldTrump": "President Trump",
    "SecRubio":        "Secretary Rubio",
    "USTradeRep":      "USTR",
    "DOWResponse":     "Department of War",
    "CIADirector":     "CIA Director",
    "SecWar":          "Secretary of War",
    "StephenM":        "Stephen Miller",
    "DNIGabbard":      "DNI Gabbard",
    "howardlutnick":   "Secretary Lutnick",
}
X_INCLUDE_LESS_IMPORTANT = False

# These three accounts are PRC-government-run — their own posts are
# trivially "about China" almost by definition (a Chinese embassy tweeting
# about Chinese industry, agriculture, disaster relief, etc.), so
# classify_relevance()'s bar ("does this explicitly and substantively
# involve China") is the wrong test here: it's designed for US-origin
# sources (state.gov, whitehouse.gov, ...) where "mentions China" is a
# meaningfully rare, relevant signal — applied to a PRC account's own
# feed, it's satisfied by nearly everything they post. Confirmed live,
# 2026-09-01: classify_relevance said YES to a Chinese Embassy tweet about
# "China's digital publishing industry" growth, with ZERO US mention
# anywhere, reasoning only that it "directly referenc[es] China." Handled
# in process_x_tweet() with the narrower _EXPLICIT_US_MENTION_RE check
# instead (same fix already applied to filter_relevant_exchanges — see
# its docstring), matching what these tweets actually need to clear the
# bar for a US-CHINA relations tracker.
_PRC_X_ACCOUNTS = {"ChineseEmbinUS", "SpoxCHN_LinJian", "SpoxCHN_MaoNing"}

# Pay-per-use rate confirmed live via docs.x.com, re-checked 2026-09-02 —
# see the module comment above. Update here if X's pricing changes. (No
# per-user-lookup rate anymore — the search-based redesign never resolves
# a numeric user ID at all; "from:" takes a plain username.)
_X_USD_PER_POST_READ = 0.005


def _log_x_cost(label: str, units: int, unit_price: float) -> None:
    """
    Same USAGE_LOG_PATH/summarize_usage_log() plumbing as _log_usage(), but
    for X's own pay-per-use API charges rather than an LLM call — these are
    real, precisely-known costs (X bills a flat rate per read, not tokens),
    so `total_tokens` here holds a unit count (users or posts read) rather
    than an actual token count. summarize_usage_log() only sums numbers and
    reports a USD total, so this reuses it correctly without changes.
    """
    usd = units * unit_price
    log.info(f"[usage] {label} via X API: units={units} usd=${usd:.6f}")
    try:
        with open(USAGE_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps({
                "ts": _utcnow().isoformat(),
                "label": label,
                "provider": "X API",
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "reasoning_tokens": 0,
                "total_tokens": units,
                "usd": usd,
            }) + "\n")
    except Exception as exc:
        log.warning(f"Failed to append to {USAGE_LOG_PATH}: {exc}")


def _x_api_key() -> str:
    api_key = os.environ.get("X_API_KEY")
    if not api_key:
        raise RuntimeError("X_API_KEY is not set")
    return api_key


def _load_x_state() -> dict:
    if Path(X_STATE_PATH).exists():
        return json.loads(Path(X_STATE_PATH).read_text())
    return {}


def _save_x_state(state: dict) -> None:
    Path(X_STATE_PATH).write_text(json.dumps(state, indent=2))


def build_x_search_query(accounts: list[str], terms: str) -> str:
    """
    '(term1 OR term2 ...) (from:user1 OR from:user2 ...)' — X search syntax:
    parenthesized OR-groups, ANDed together by simple juxtaposition. Pulled
    out as its own pure function so the actual query string is testable
    without a live call. Comfortably under "recent search"'s 512-char query
    limit for both real groups this pipeline builds (checked: ~300 chars
    for the 8-account China-terms group, ~150 for the 3-account US-terms
    group — see NOTES.md for the exact counts).
    """
    from_clause = " OR ".join(f"from:{u}" for u in accounts)
    return f"({terms}) ({from_clause})"


def _x_search_recent(query: str, since_id: str | None, max_results: int = 100) -> list[dict]:
    """
    GET /2/tweets/search/recent — ONE call covering multiple accounts AND
    a keyword filter, replacing the old per-account GET /2/users/:id/tweets
    polling (2026-09-02). X bills a search result exactly like a timeline
    result (same $/post-read rate), but a search only RETURNS (and bills
    for) posts that already match the query — see the module comment for
    the real cost data that motivated this. `expansions=author_id` +
    `user.fields=username` gets each result's author's username back
    without a separate lookup call (X's search "from:" operator takes a
    plain username directly — no numeric user-ID resolution needed at
    all with this design, unlike the old per-account timeline approach).
    `since_id` here is a GROUP-level cursor (see _PENDING_X_SINCE_IDS's
    comment), not per-account.
    """
    params = {
        "query": query,
        "max_results": max_results,
        "tweet.fields": "created_at,author_id",
        "expansions": "author_id",
        "user.fields": "username",
    }
    if since_id:
        params["since_id"] = since_id

    resp = httpx.get(
        "https://api.x.com/2/tweets/search/recent",
        headers={"Authorization": f"Bearer {_x_api_key()}"},
        params=params,
        timeout=30.0,
    )
    if resp.status_code != 200:
        log.error(f"[x] Search failed: {resp.status_code} {resp.text[:200]}")
        return []
    data = resp.json()
    tweets = data.get("data", [])
    if tweets:
        _log_x_cost("x_post_read", len(tweets), _X_USD_PER_POST_READ)
    users_by_id = {u["id"]: u.get("username", "?") for u in data.get("includes", {}).get("users", [])}
    for t in tweets:
        t["username"] = users_by_id.get(t.get("author_id"), "?")
    return tweets


def process_x_tweet(username: str, tweet: dict, model: genai.Client, conn: sqlite3.Connection) -> bool:
    """
    Per-tweet body — a tweet is already the whole "document" (short, no
    Q&A structure, nothing to extract), so this skips the heavier release-
    entry machinery (content_type_from_paragraphs, extract_key_paragraphs)
    that exists for multi-paragraph pages and just classifies the tweet
    text directly.
    """
    tweet_id = tweet.get("id")
    text = (tweet.get("text") or "").strip()
    url = f"https://x.com/{username}/status/{tweet_id}"
    if not text or is_seen(conn, url):
        return False

    # Date-only parsing (matches every other source — entries group by day,
    # not time-of-day) so this doesn't need to match X's exact timestamp
    # format precisely.
    created_at = tweet.get("created_at", "")
    date = datetime.strptime(created_at[:10], "%Y-%m-%d") if len(created_at) >= 10 else _utcnow()

    if username in _PRC_X_ACCOUNTS:
        # Free keyword check, not classify_relevance() — see
        # _PRC_X_ACCOUNTS's docstring: for a PRC-government account,
        # "does this mention China" is trivially true almost always, so
        # the actual bar has to be "does this mention the US," and a
        # short tweet doesn't need LLM nuance to check that.
        if not _EXPLICIT_US_MENTION_RE.search(text):
            mark_seen(conn, url)
            return False
    else:
        is_rel, reason = classify_relevance(model, text)
        if not is_rel:
            if not reason.startswith("Keyword pre-filter:"):
                flag_for_review(url, text[:80], reason)
            mark_seen(conn, url)
            return False

    summary, anchor = get_summary_and_anchor(model, text, f"@{username} on X", url)
    source_label = _X_ACCOUNT_DISPLAY_NAMES.get(username, username)
    queue_entry("release", date, summary, url, anchor, paragraphs=[text], source_label=source_label)
    # mark_seen() happens in flush_pending_entries() — see its docstring.
    log.info(f"[x] Queued: @{username} — {text[:80]}")
    return True


def scrape_x(model: genai.Client, conn: sqlite3.Connection, doc: Document) -> None:
    """
    Two combined search queries per run — one for the non-PRC accounts
    (China-keyword-filtered), one for the PRC accounts (US-keyword-
    filtered) — instead of one GET .../tweets call per account. See the
    module comment above for why (real cost data from the old per-account
    design) and build_x_search_query()/_x_search_recent() for how.

    `since_id` is deliberately NOT saved to disk here, even though a given
    group's tweets are already fully processed by the time this function
    moves to the next group — the same mark_seen()/flush ordering bug
    (see flush_pending_entries()'s docstring, 2026-09-01) applies to this
    state file too: X's `output/x_accounts_state.json` is the "have we
    already read this" bookkeeping X itself is billed against, exactly
    parallel to tracker.db's seen_urls for the doc. `scrape_x` covers ALL
    accounts as one "source," flushed once after this whole function
    returns — if `since_id` advanced to disk per-group as we went and the
    process then crashed before that flush, the next run's `since_id`
    would already be past those tweets, so X's API would never return them
    again: permanent, silent loss of content ALREADY PAID FOR, not just a
    wasted retry. Instead, each group's prospective new `since_id` is
    staged in `_PENDING_X_SINCE_IDS` and only actually written by
    flush_pending_entries(), after doc.save() confirms the corresponding
    entries are durably on disk. Worst case on a crash before that: the
    next run re-reads (re-bills) a few already-seen tweets, which
    `process_x_tweet`'s `is_seen()` check then correctly skips re-queuing
    — a small wasted cost, never lost content.
    """
    accounts = list(X_ACCOUNTS_NORMAL)
    if X_INCLUDE_LESS_IMPORTANT:
        accounts += X_ACCOUNTS_LESS_IMPORTANT

    prc_accounts = [a for a in accounts if a in _PRC_X_ACCOUNTS]
    other_accounts = [a for a in accounts if a not in _PRC_X_ACCOUNTS]

    groups = []
    if other_accounts:
        groups.append(("china", other_accounts, _X_CHINA_SEARCH_TERMS))
    if prc_accounts:
        groups.append(("prc", prc_accounts, _X_US_SEARCH_TERMS))

    state = _load_x_state()
    search_groups_state = state.get("_search_groups", {})

    for group_key, group_accounts, terms in groups:
        query = build_x_search_query(group_accounts, terms)
        since_id = search_groups_state.get(group_key, {}).get("since_id")
        # 100, not MAX_NEW_ITEMS_PER_RUN (30) — that constant is calibrated
        # for "max items from ONE list page," but one search call here
        # covers MULTIPLE accounts' combined matches, so reusing it would
        # be tighter than intended and risk missing real content during a
        # busy week. 100 is "recent search"'s own documented per-call max.
        tweets = _x_search_recent(query, since_id, max_results=100)
        log.info(f"[x/{group_key}] {len(tweets)} new matching tweets since last run")

        for tweet in tweets:
            try:
                process_x_tweet(tweet.get("username", "?"), tweet, model, conn)
            except Exception as exc:
                log.error(f"[x/{group_key}] Error on tweet {tweet.get('id', '?')}: {exc}")

        # Explicit max(), not tweets[0] — search results are typically
        # newest-first but that's not a documented guarantee the way it
        # is for the old timeline endpoint, and getting this wrong would
        # silently re-bill (though never lose) content. Staged, not
        # saved — see this function's docstring.
        if tweets:
            _PENDING_X_SINCE_IDS[group_key] = str(max(int(t["id"]) for t in tweets))

        # No REQUEST_SLEEP here — that constant exists for politeness
        # toward government websites being scraped via fetch()/httpx
        # directly against .gov servers; X's v2 API is a paid, purpose-
        # built API with its own generous published rate limits (far
        # beyond what 1-2 search calls per run gets remotely close to),
        # and _x_search_recent() calls httpx directly, never through
        # fetch() — so this sleep was never protecting against an actual
        # X-side limit.


# ── Main ──────────────────────────────────────────────────────────────────────

SOURCES = {
    "fmprc_conf":              "MFA press conferences",
    "fmprc_remarks":           "MFA spokesperson remarks",
    "mfa_leadership_speeches": "MFA leadership speeches (Chinese)",
    "mfa_leadership_activity": "MFA leadership activity (Chinese)",
    "mofcom":                  "MOFCOM press conferences",
    "mofcom_daily":            "MOFCOM daily news release (Chinese)",
    "mofcom_leadership":       "MOFCOM leadership activity (Chinese)",
    "mofcom_dept_leadership":  "MOFCOM dept. leadership activity (Chinese)",
    "mofcom_bureau_heads":     "MOFCOM bureau/dept head announcements (Chinese)",
    "mofcom_special_conf":     "MOFCOM special press conferences (Chinese)",
    "mofcom_lxxwfbh":          "MOFCOM regular weekly press conferences (Chinese)",
    "scio":                    "State Council Information Office",
    "mnd":                     "Ministry of National Defense (mod.gov.cn)",
    "state":                   "State Dept",
    "whitehouse":              "White House",
    "treasury":                "Treasury",
    "ustr":                    "USTR",
    # "wardept" (Dept of War / war.gov) intentionally removed from active
    # dispatch, 2026-09-02 — see the DISABLED note on scrape_wardept()
    # itself and NOTES.md for the full investigation. Its article pages
    # are blocked at the infrastructure level (Akamai) for any non-
    # browser client, confirmed even against a real headless-browser
    # test and from multiple independent networks — not something a
    # code fix can work around. The function is left in the file,
    # unreachable, rather than deleted, in case this policy ever changes.
    "x":                       "X (Twitter) accounts",
}


def _parse_user_date(s: str) -> date:
    """
    Accepts either 'YYYY-MM-DD' or plain 'YYYYMMDD' (no dashes) — the
    latter added 2026-09-02 per user request, typing dashes for every
    date felt unnecessarily fussy for something typed into a plain-
    language terminal prompt (see "Run Weekly Tracker (Mac).command"). Raises
    ValueError on anything else, same as date.fromisoformat() alone did
    — main()'s caller already turns that into a friendly message instead
    of a raw traceback.
    """
    if re.fullmatch(r"\d{8}", s):
        return datetime.strptime(s, "%Y%m%d").date()
    return date.fromisoformat(s)


def default_week_range(today: date | None = None) -> tuple[date, date]:
    """
    The tracker's natural unit has always been a Tuesday-through-Monday
    week (see NOTES.md and run_scheduled.sh's Monday cron) — so "the most
    recent week" means the most recently COMPLETED such week, not
    whatever's still in progress today.

    end   = the most recent Monday on or before `today` (today itself,
            if today IS Monday)
    start = the Tuesday 6 days before that

    This one formula covers every day of the week without a special case:
    run it on Monday and you get last Tuesday through today; run it on
    Tuesday (or any other day) and you get the exact same week, since
    yesterday's Monday is still the most recent completed one — the
    in-progress days of the CURRENT week are deliberately left for next
    week's run, matching how the weekly deliverable has always worked.
    """
    today = today or date.today()
    end = today - timedelta(days=today.weekday())  # Monday == weekday() 0
    start = end - timedelta(days=6)
    return start, end


def _format_week_for_filename(start: date, end: date) -> str:
    """'Aug 25-31, 2026' / 'Aug 25-Sep 1, 2026' style label for the dated
    doc copy — spans months correctly instead of assuming both dates fall
    in the same one."""
    if start.month == end.month:
        return f"{start.strftime('%b')} {start.day}-{end.day}, {end.year}"
    return f"{start.strftime('%b')} {start.day}-{end.strftime('%b')} {end.day}, {end.year}"


def main() -> None:
    global MAX_NEW_ITEMS_PER_RUN
    parser = argparse.ArgumentParser(description="US-China tracker scraper")
    parser.add_argument(
        "--source",
        choices=list(SOURCES.keys()),
        help="Run a single source only (default: all)",
    )
    parser.add_argument(
        "--start", type=str, default=None,
        help="Week start date, YYYYMMDD or YYYY-MM-DD (default: last complete Tuesday)",
    )
    parser.add_argument(
        "--end", type=str, default=None,
        help="Week end date, YYYYMMDD or YYYY-MM-DD (default: last complete Monday)",
    )
    parser.add_argument(
        "-v", "--verbose", action="store_true",
        help="Show full detail on screen instead of just the progress bar "
             "and final summary (the full detail is always written to "
             "logs/ either way)",
    )
    parser.add_argument(
        "--no-open", action="store_true",
        help="Don't automatically open the finished doc when the run ends "
             "(run_scheduled.sh passes this — nobody's watching an "
             "unattended cron/launchd run to see a doc pop open)",
    )
    parser.add_argument(
        "--max-items", type=int, default=None,
        help=f"How many not-yet-seen items per source to consider (default: "
             f"{MAX_NEW_ITEMS_PER_RUN}). Only worth raising for a deliberate "
             f"catch-up run against a source with a real backlog (e.g. "
             f"--max-items 200) — the default stays low because a normal "
             f"week's new content per source is small; raising it here "
             f"doesn't change what a future normal run costs.",
    )
    args = parser.parse_args()
    if args.verbose:
        _console_handler.setLevel(logging.INFO)
    if args.max_items is not None:
        MAX_NEW_ITEMS_PER_RUN = args.max_items

    default_start, default_end = default_week_range()
    try:
        week_start = _parse_user_date(args.start) if args.start else default_start
        week_end   = _parse_user_date(args.end) if args.end else default_end
    except ValueError:
        # A friendly message instead of a raw traceback — added 2026-09-02
        # once a plain-language date prompt (see "Run Weekly Tracker (Mac).command")
        # started inviting a non-technical user to type a date directly,
        # where a typo (a slash instead of a dash, a 2-digit year, ...) is
        # a real, expected failure mode rather than a rare programmer error.
        parser.error(
            f"Couldn't understand that date. Use YYYYMMDD (like 20260804) or "
            f"YYYY-MM-DD (like 2026-08-04) — got --start={args.start!r} "
            f"--end={args.end!r}."
        )
    if week_start > week_end:
        parser.error(f"--start ({week_start}) is after --end ({week_end})")

    # A REAL filter, not just a label — see queue_entry()'s own comment.
    # An item whose actual date falls outside [week_start, week_end] never
    # gets queued into this run's output at all, regardless of what a
    # source happens to find. Set as module globals because queue_entry()
    # (the single choke point every source funnels through) has no other
    # way to see this run's target range.
    global _RUN_TARGET_START, _RUN_TARGET_END
    _RUN_TARGET_START, _RUN_TARGET_END = week_start, week_end
    _QUEUED_URL_SLUGS_THIS_RUN.clear()

    week_label = _format_week_for_filename(week_start, week_end)
    print(f"US-China Relations Tracker — covering {week_label}\n")

    run_start_monotonic = time.monotonic()
    run_start_ts = _utcnow().isoformat()

    model  = init_llm()
    conn   = init_db()
    doc    = get_or_create_doc(DOC_PATH)
    client = make_client()

    run_all = args.source is None
    s = args.source
    total_new = 0

    sources_to_run = list(SOURCES.keys()) if run_all else [s]
    # logging_redirect_tqdm(): a plain logging.StreamHandler writes
    # straight to the terminal with no idea a progress bar is live there
    # — a WARNING/ERROR firing mid-run (a real one, e.g. "[mofcom] Failed
    # to fetch list...") used to land in the middle of tqdm's own cursor-
    # control sequences, corrupting the bar into multiple broken/
    # duplicate-looking lines. This temporarily reroutes every logging
    # handler's output through tqdm.write() instead, which knows to clear
    # the bar, print the line cleanly, then redraw the bar — for the
    # whole rest of this function, not just around one log call, since
    # any source's fetch can log a warning/error at any moment. Confirmed
    # this was the actual mechanism (not something more exotic) by
    # reproducing the exact garbled-bar shape from a real MOFCOM SSL
    # error live, 2026-09-03.
    with logging_redirect_tqdm():
        pbar = tqdm(total=len(sources_to_run), desc="Starting...", unit="source")
        source_failures: dict[str, list[str]] = {}

        def run(key: str, fn, *fn_args):
            nonlocal total_new
            if run_all or s == key:
                pbar.set_description(SOURCES.get(key, key))
                capture = _SourceErrorCapture()
                log.addHandler(capture)
                try:
                    fn(*fn_args)
                except Exception as exc:
                    log.error(f"[{key}] Unhandled error: {exc}")
                finally:
                    log.removeHandler(capture)
                if capture.messages:
                    source_failures[key] = capture.messages
                # Flush after each source rather than only once at the very end:
                # bounds how much work is lost if the whole process gets killed
                # mid-run, while still collapsing repeated date headings within
                # (and, via the module-level "last date written", across)
                # sources — see the PENDING_ENTRIES comment block above.
                n = flush_pending_entries(doc, conn)
                if n:
                    log.info(f"[{key}] Wrote {n} entries to {DOC_PATH}")
                    total_new += n
                pbar.update(1)

        run("fmprc_conf",    scrape_fmprc,
            "https://www.fmprc.gov.cn/eng/xw/fyrbt/lxjzh/",
            "press conference", model, conn, doc, client)

        run("fmprc_remarks", scrape_fmprc,
            "https://www.fmprc.gov.cn/eng/xw/fyrbt/fyrbt/",
            "spokesperson remarks", model, conn, doc, client)

        run("mfa_leadership_speeches", scrape_mfa_leadership,
            "https://www.mfa.gov.cn/web/ziliao_674904/zyjh_674906/",
            "leadership speeches", model, conn, doc, client)

        run("mfa_leadership_activity", scrape_mfa_leadership,
            "https://www.mfa.gov.cn/web/wjdt_674879/wjbxw_674885/",
            "leadership activity", model, conn, doc, client)

        run("mofcom",        scrape_mofcom,     model, conn, doc)
        run("mofcom_daily",  scrape_mofcom_daily, model, conn, doc)
        run("mofcom_leadership",      scrape_mofcom_leadership,      model, conn, doc)
        run("mofcom_dept_leadership", scrape_mofcom_dept_leadership, model, conn, doc)
        run("mofcom_bureau_heads",    scrape_mofcom_bureau_heads,    model, conn, doc)
        run("mofcom_special_conf",    scrape_mofcom_special_conf,    model, conn, doc)
        run("mofcom_lxxwfbh",         scrape_mofcom_lxxwfbh,         model, conn, doc)
        run("scio",          scrape_scio,       model, conn, doc)
        run("mnd",           scrape_mnd,        model, conn, doc)
        run("state",         scrape_state,      model, conn, doc)
        run("whitehouse",    scrape_whitehouse, model, conn, doc)
        run("treasury",      scrape_treasury,   model, conn, doc)
        run("ustr",          scrape_ustr,       model, conn, doc)
        # "wardept" intentionally not run — see the SOURCES dict comment and
        # scrape_wardept()'s own DISABLED note.
        run("x",             scrape_x,          model, conn, doc)

        pbar.close()

    # The actual per-week document — rendered fresh from the `entries`
    # table (see render_doc_for_range()'s docstring), not a copy of the
    # ever-growing master doc. Contains exactly week_start..week_end,
    # nothing else — whether those entries were queued in THIS run or a
    # previous one, so re-running an already-covered week is both fast
    # (nothing new to queue) and still produces the complete, correct
    # document for it. The master (DOC_PATH / tracker_output.docx) is
    # untouched by this — it keeps accumulating every entry ever written,
    # which is what dedup/is_seen() needs; this is purely a separate,
    # filtered VIEW for the human-facing weekly deliverable.
    week_doc = render_doc_for_range(conn, week_start, week_end)
    dated_name = f"US-China Tracker {week_label}.docx"
    dated_path = os.path.join(DATA_DIR, dated_name)
    week_doc.save(dated_path)

    elapsed = _format_duration(time.monotonic() - run_start_monotonic)
    llm_tokens, x_reads, run_usd = _summarize_run_usage(run_start_ts)
    usage_bits = [f"{llm_tokens:,} tokens"]
    if x_reads:
        usage_bits.append(f"{x_reads} X reads")
    usage_line = " + ".join(usage_bits)

    print(f"\nDone — {total_new} new entr{'y' if total_new == 1 else 'ies'} added for {week_label}.")
    print(f"Saved to: {dated_path}")
    print(f"Took {elapsed} — {usage_line}, est. cost ${run_usd:.4f}.")

    # Explicit pass/fail-per-source report, added 2026-09-03 per user
    # request — a real fetch failure (a site down, a list page that
    # stopped loading) was previously only visible in logs/*.log, which
    # nobody's watching mid-run; the console only shows WARNING+ by
    # default (see _console_handler) so an ERROR-level failure like
    # "[scio] Failed to fetch ..." could scroll by without ever being
    # seen. This makes a real failure impossible to miss without needing
    # -v or a log file at all.
    if source_failures:
        print(f"\n⚠ {len(source_failures)} source(s) had errors this run:")
        for key, messages in source_failures.items():
            label = SOURCES.get(key, key)
            unique = list(dict.fromkeys(messages))  # de-dupe, keep order
            print(f"  - {label} ({key}): {len(messages)} error(s)")
            for msg in unique[:3]:
                print(f"      {msg}")
            if len(unique) > 3:
                print(f"      ...and {len(unique) - 3} more (see logs/ for the full detail)")
    else:
        print("\nNo source errors this run.")

    # Standing reminder, printed every run regardless of errors above —
    # these three sources have no scraping code at all (not a failure,
    # a deliberate scope decision; see input/notes/SOURCES.md), so
    # nothing above would ever flag them as missing. Easy to forget
    # they're not covered since most other sources ARE.
    print(
        "\nReminder: Truth Social, YouTube, and Dept of War (war.gov) are "
        "NOT scraped by this tool (see input/notes/SOURCES.md for why) — "
        "check those manually if this week might include something from them."
    )

    if not args.no_open and sys.platform == "darwin":
        # Opens the finished doc in whatever app is set to handle .docx
        # (Word, by default) — added 2026-09-02 per user request, so the
        # deliverable is right there instead of needing to go find it in
        # Finder. `check=False`: if this fails for any reason (no doc app
        # installed, sandboxed environment, ...), that's not a reason to
        # make the run itself look like it failed — the file is already
        # saved either way, this is just a convenience on top.
        subprocess.run(["open", dated_path], check=False)

    log.info("All done.")


if __name__ == "__main__":
    main()

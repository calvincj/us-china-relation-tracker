#!/usr/bin/env python3
from __future__ import annotations
"""
Offline regression tests for scraper.py's pure logic — the classification/
label-parsing/keyword functions that don't need a network call or an LLM
to test. No API keys required, no cost, runs in under a second.

Every test here is anchored to a REAL bug found via live testing during
development (see NOTES.md for the full story on each) — this file exists
specifically so none of them can silently regress as the code changes.
Where a test's docstring says "bug #N," that's the NOTES.md section with
the original live evidence.

This deliberately does NOT test anything that needs a live fetch, an LLM
call, or a real API key (classify_relevance's actual judgment,
translate_to_english, the process_*_item()/scrape_*() functions
end-to-end) — that's what backtest.py is for, against real past-tracker
weeks. This file only covers the parts that are pure functions of their
input: label/speaker parsing, keyword matching, content-type
classification.

Usage (run from the project root, not from inside code/):
    python3 code/test_scraper.py
    python3 -m unittest code.test_scraper -v   # (needs code/__init__.py — not set up; use the line above)
"""

import json
import os
import sqlite3
import sys
import tempfile
import unittest
from datetime import date, datetime
from unittest import mock

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import scraper as S
import httpx


class TestFetchRetryLogic(unittest.TestCase):
    """fetch()'s `retries` parameter used to be silently ignored for ANY
    HTTP error status (403/404/500/521/...) — `except httpx.HTTPStatusError`
    returned None on the very first attempt regardless of `retries`, even
    though Treasury's call site explicitly passes retries=4. Found
    reviewing the function directly, 2026-09-01 (not via live testing —
    the specific failures actually observed this session, war.gov's 403
    and scio.gov.cn's 521, both happened to be cases where retrying
    wouldn't have helped anyway, so this gap was invisible in practice
    until read closely). Fixed: a 5xx now gets retried with backoff, same
    as a network-level failure; a 4xx is still an immediate bail (those
    are essentially always persistent, retrying wastes a request)."""

    class _FakeResponse:
        def __init__(self, status_code):
            self.status_code = status_code
            self.headers = {}

        def raise_for_status(self):
            if self.status_code >= 400:
                request = httpx.Request("GET", "https://example.com/test")
                raise httpx.HTTPStatusError("error", request=request, response=self)

    class _FakeClient:
        """Returns FakeResponses from a scripted list, one per call.get()."""
        def __init__(self, status_codes):
            self.status_codes = list(status_codes)
            self.call_count = 0

        def get(self, url):
            code = self.status_codes[min(self.call_count, len(self.status_codes) - 1)]
            self.call_count += 1
            return TestFetchRetryLogic._FakeResponse(code)

    def test_5xx_is_retried_until_success(self):
        client = self._FakeClient([503, 503, 200])
        with mock.patch.object(S.time, "sleep"):  # skip real waits in the test
            resp = S.fetch(client, "https://example.com/test", retries=3)
        self.assertIsNotNone(resp)
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(client.call_count, 3, "should have retried through both 503s")

    def test_5xx_gives_up_after_exhausting_retries(self):
        client = self._FakeClient([503, 503, 503])
        with mock.patch.object(S.time, "sleep"):
            resp = S.fetch(client, "https://example.com/test", retries=3)
        self.assertIsNone(resp)
        self.assertEqual(client.call_count, 3)

    def test_4xx_is_not_retried(self):
        # A 404/403 should bail on the FIRST attempt — retrying a
        # persistent client error just wastes a request.
        client = self._FakeClient([404, 200])
        with mock.patch.object(S.time, "sleep"):
            resp = S.fetch(client, "https://example.com/test", retries=3)
        self.assertIsNone(resp)
        self.assertEqual(client.call_count, 1, "must not retry a 4xx even though a later call would succeed")

    def test_success_on_first_try(self):
        client = self._FakeClient([200])
        with mock.patch.object(S.time, "sleep"):
            resp = S.fetch(client, "https://example.com/test", retries=3)
        self.assertIsNotNone(resp)
        self.assertEqual(client.call_count, 1)

    def test_last_fetch_status_records_the_real_code(self):
        # A caller (e.g. scrape_wardept(), for war.gov's persistent 403
        # block) needs to know WHICH status code a None return meant,
        # without fetch()'s Optional[Response] contract changing for
        # every other caller — see _LAST_FETCH_STATUS's own comment.
        client = self._FakeClient([403])
        with mock.patch.object(S.time, "sleep"):
            S.fetch(client, "https://example.com/test", retries=3)
        self.assertEqual(S._LAST_FETCH_STATUS, 403)

    def test_last_fetch_status_resets_to_none_on_success(self):
        client = self._FakeClient([403])
        with mock.patch.object(S.time, "sleep"):
            S.fetch(client, "https://example.com/test", retries=3)
        self.assertEqual(S._LAST_FETCH_STATUS, 403)
        client2 = self._FakeClient([200])
        with mock.patch.object(S.time, "sleep"):
            S.fetch(client2, "https://example.com/test", retries=3)
        self.assertIsNone(S._LAST_FETCH_STATUS, "a later successful fetch must clear the stale status")


class TestRelevanceKeywordBoundaries(unittest.TestCase):
    """Bug #18: a Latin-only \\b(?:...) with no CLOSING \\b let a short
    alternative match as a mere prefix of an unrelated word — "AI" matched
    inside "UnmAnned AIrcraft Systems" because the leading boundary
    (space→"A") was satisfied and nothing checked the trailing one."""

    def test_ai_does_not_match_inside_aircraft(self):
        self.assertIsNone(S.RELEVANCE_KEYWORDS.search("unmanned aircraft systems"))

    def test_ai_still_matches_as_a_real_word(self):
        self.assertIsNotNone(S.RELEVANCE_KEYWORDS.search("the AI race"))
        self.assertIsNotNone(S.RELEVANCE_KEYWORDS.search("artificial intelligence competition"))

    def test_other_real_keywords_still_match(self):
        for text in ["a 25% tariff on Chinese goods", "export control on semiconductor chips",
                     "U.S. trade deficit", "sanctions on Chinese entities"]:
            with self.subTest(text=text):
                self.assertIsNotNone(S.RELEVANCE_KEYWORDS.search(text))

    def test_plural_forms_of_singular_keywords_still_match(self):
        # Adding the trailing \b above (to fix the AI/aircraft bug) also
        # silently broke ordinary plurals of the affected single-word
        # keywords — "tariff\b" doesn't match inside "tariffs". Caught by
        # this test suite, not by the live testing that found the
        # original bug.
        for text in ["new tariffs on Chinese goods", "sanctions were announced", "chips export",
                     "several semiconductors", "export controls expanded", "import duties rose",
                     "trade deficits widened", "trade surpluses narrowed", "forced transfers of technology"]:
            with self.subTest(text=text):
                self.assertIsNotNone(S.RELEVANCE_KEYWORDS.search(text))

    def test_us_source_keywords_include_bare_china_mention(self):
        # US_SOURCE_RELEVANCE_KEYWORDS extends RELEVANCE_KEYWORDS with a
        # bare China/Beijing/Xi Jinping mention, for screening US-origin
        # releases (which need "does this mention China at all," unlike
        # RELEVANCE_KEYWORDS itself, built for finding the US side WITHIN
        # inherently-China-focused content).
        self.assertIsNotNone(S.US_SOURCE_RELEVANCE_KEYWORDS.search("a deal with China"))
        self.assertIsNotNone(S.US_SOURCE_RELEVANCE_KEYWORDS.search("Beijing announced"))

    def test_does_not_match_latin_south_central_america(self):
        # Real false positive, 2026-09-02 — same fix as
        # TestExplicitUsMention's version, applied here too since
        # RELEVANCE_KEYWORDS shares the same "America[n]" alternative.
        for text in [
            "cooperation with Latin America on trade",
            "trade ties with South America",
            "engagement across Central America",
        ]:
            with self.subTest(text=text):
                self.assertIsNone(S.RELEVANCE_KEYWORDS.search(text))

    def test_does_not_match_us_dollars_as_a_currency_unit(self):
        # Real false positive, 2026-09-02 — a Chinese economic release
        # quoting a USD-equivalent figure (原文 "8100亿美元") translates to
        # English as "810 billion U.S. dollars," which used to trip this
        # regex on "U.S." alone even though the article never actually
        # discusses the United States as a country/actor. Mirrors the
        # Latin America fix above: same shape of bug, same fix (negative
        # lookahead instead of lookbehind, since "dollar(s)" trails here).
        for text in [
            "China exported 810 billion U.S. dollars worth of goods last year.",
            "The fund is valued at 3 trillion U.S. dollars.",
            "The United States dollar weakened against the yuan this week.",
        ]:
            with self.subTest(text=text):
                self.assertIsNone(S.RELEVANCE_KEYWORDS.search(text))

    def test_still_matches_us_as_a_real_actor_near_dollar_amounts(self):
        # The dollar exclusion must be narrow — a real US mention that
        # merely happens to be near a dollar figure elsewhere in the
        # sentence must still match.
        self.assertIsNotNone(
            S.RELEVANCE_KEYWORDS.search(
                "The United States announced a $2 billion aid package."
            )
        )

    def test_matches_bare_us_without_periods(self):
        # Real, high-impact bug, live 2026-09-03: a real LLM translation
        # of a MOFCOM Q&A rendered the abbreviation as "US government"
        # (no periods) rather than "U.S. government" — this pattern's
        # only US-abbreviation alternative was the strictly-dotted
        # "U\.S\b", so a real reporter's question about a proposed US
        # tariff hike silently failed the relevance check purely because
        # of which of two extremely common, equally valid spellings the
        # translation happened to use.
        for text in ["the US government said", "The US Treasury announced new sanctions"]:
            with self.subTest(text=text):
                self.assertIsNotNone(S.RELEVANCE_KEYWORDS.search(text))

    def test_does_not_match_lowercase_us_pronoun(self):
        # The bare-"US" fix must be case-SENSITIVE (uppercase only) — the
        # rest of this pattern is case-insensitive, which would otherwise
        # make a plain lowercase "us" (the pronoun) match too.
        for text in ["let us know if you have questions", "join us for the event",
                     "please contact us for more information"]:
            with self.subTest(text=text):
                self.assertIsNone(S.RELEVANCE_KEYWORDS.search(text))

    def test_bare_us_dollar_exclusion_still_works(self):
        # The dollar-amount exclusion must apply to the new bare-"US"
        # alternative too, not just the dotted "U.S." one.
        self.assertIsNone(
            S.RELEVANCE_KEYWORDS.search(
                "China exported 810 billion US dollars worth of goods."
            )
        )


class TestExplicitUsMention(unittest.TestCase):
    """The narrow "does this actually name the US" pattern — used by
    filter_relevant_exchanges (bug #19) and process_x_tweet's PRC-account
    branch (bug #23). Deliberately narrower than RELEVANCE_KEYWORDS: no
    topic words (Taiwan, AI, tariff), only direct US/named-official
    mentions."""

    def test_matches_direct_us_mentions(self):
        for text in ["the U.S. imposed tariffs", "United States officials said",
                     "Washington announced", "President Trump stated", "Secretary Bessent said"]:
            with self.subTest(text=text):
                self.assertIsNotNone(S._EXPLICIT_US_MENTION_RE.search(text))

    def test_does_not_match_topic_only_mentions(self):
        # These all discuss China-flashpoint topics with ZERO US
        # involvement — exactly the false-positive shape from bug #19.
        for text in [
            "erroneous remarks and actions on Taiwan made by the Japanese government",
            "cooperation on artificial intelligence with all other countries",
            "China's digital publishing industry continues to grow",
        ]:
            with self.subTest(text=text):
                self.assertIsNone(S._EXPLICIT_US_MENTION_RE.search(text))

    def test_does_not_match_latin_south_central_america(self):
        # Real false positive, 2026-09-02: "America" alone matched inside
        # "Latin America"/"South America"/"Central America" (a region, not
        # the US) — and this regex has no LLM judgment downstream to catch
        # it (filter_relevant_exchanges/PRC-account tweets decide directly
        # off this match).
        for text in [
            "China deepened cooperation with Latin America on trade.",
            "China deepened cooperation with South America on trade.",
            "Discussions with Central America continued this week.",
        ]:
            with self.subTest(text=text):
                self.assertIsNone(S._EXPLICIT_US_MENTION_RE.search(text))

    def test_still_matches_american_as_a_real_word(self):
        # The lookbehind must only block "Latin/South/Central America",
        # not "America[n]" in general.
        self.assertIsNotNone(S._EXPLICIT_US_MENTION_RE.search("The American ambassador met with officials."))

    def test_does_not_match_us_dollars_as_a_currency_unit(self):
        # Same fix and reasoning as RELEVANCE_KEYWORDS's version above,
        # applied here too since this regex shares the same "U.S"/"United
        # States" alternatives — and has no LLM judgment downstream to
        # catch a miss.
        for text in [
            "China's central bank holds 3 trillion U.S. dollars in reserves.",
            "The deal was worth 500 million U.S. dollars.",
            "The United States dollar weakened against the yuan this week.",
        ]:
            with self.subTest(text=text):
                self.assertIsNone(S._EXPLICIT_US_MENTION_RE.search(text))

    def test_still_matches_us_as_a_real_actor_near_dollar_amounts(self):
        self.assertIsNotNone(
            S._EXPLICIT_US_MENTION_RE.search(
                "The United States announced a $2 billion aid package."
            )
        )

    def test_matches_bare_us_without_periods(self):
        # Same real bug and fix as RELEVANCE_KEYWORDS's version above —
        # this regex in particular has NO downstream LLM check, so a miss
        # here silently drops the whole Q&A block with nothing to catch it.
        for text in [
            "Media reports indicate that the US government is considering a tariff hike.",
            "The US Treasury announced new sanctions today.",
        ]:
            with self.subTest(text=text):
                self.assertIsNotNone(S._EXPLICIT_US_MENTION_RE.search(text))

    def test_does_not_match_lowercase_us_pronoun(self):
        for text in ["let us know if you have questions", "join us for the event"]:
            with self.subTest(text=text):
                self.assertIsNone(S._EXPLICIT_US_MENTION_RE.search(text))


class TestClassifyByLabels(unittest.TestCase):
    """Bug #16: content_type_from_paragraphs/_exchanges used to say "qa"
    if ANY label looked asker-shaped — true but not sufficient. A media-
    coverage roundup citing several outlets (each appearing once) isn't a
    real interactive Q&A; a real press conference has ONE (rarely two)
    person actually answering, no matter how many different outlets ask.
    Also covers the "CGTN Reporter"/"CNBC Reporter" compound-label gap
    found fixing this (role word only recognized as a bare exact match,
    not as a suffix on a longer label) and the original .search()-vs-
    .match() outlet-suffix bug from 2026-08-04."""

    def test_media_roundup_is_release_not_qa(self):
        # Real case: a USTR release citing 5 outlets + 8 individually-
        # quoted officials, none of them a repeat answerer.
        labels = [
            "Here are the top headlines", "Inside U.S. Trade", "Bloomberg", "CNBC", "Axios", "The Hill",
            "Senator Pete Ricketts", "Chairman Jason Smith", "Representative Adrian Smith",
            "Representative Claudia Tenney", "Representative Dan Newhouse", "Representative Rosa DeLauro",
            "Representative Chris Deluzio", "Representative Linda Sanchez",
        ]
        self.assertEqual(S._classify_by_labels(labels), "release")

    def test_real_press_conference_is_qa(self):
        # Real case: MOFCOM's He Yadong answering 5 different reporters,
        # some labeled "<Outlet> Reporter" (the compound-label gap).
        labels = [
            "He Yadong", "Xinhua News Agency Journalist", "He Yadong", "CGTN Reporter", "He Yadong",
            "Libération Daily Reporter", "He Yadong", "Kyodo News Reporter", "He Yadong", "CNBC Reporter",
        ]
        self.assertEqual(S._classify_by_labels(labels), "qa")

    def test_single_exchange_with_moderator_is_qa(self):
        # Real case: Rubio/Arévalo joint press availability. Arévalo's
        # whole turn is one long CONT-continued block, so his label never
        # repeats — the bar is "how many DISTINCT non-asker names," not
        # "does one repeat," and this must still pass with exactly 1.
        self.assertEqual(S._classify_by_labels(["MODERATOR", "PRESIDENT ARÉVALO"]), "qa")

    def test_wire_service_labels_are_qa(self):
        labels = ["Reuters", "Lin Jian", "AFP", "Lin Jian", "CCTV", "Lin Jian"]
        self.assertEqual(S._classify_by_labels(labels), "qa")

    def test_outlet_name_containing_the_matches_via_search_not_match(self):
        # 2026-08-04 bug: .match() anchors at position 0 regardless of the
        # pattern's own ^ markers, so "The New York Times" (doesn't START
        # with "times") never matched under the old .match() call.
        labels = ["A Tarde", "The New York Times", "Antara", "Lin Jian", "Lin Jian", "Lin Jian"]
        self.assertEqual(S._classify_by_labels(labels), "qa")

    def test_fewer_than_two_labels_is_release(self):
        self.assertEqual(S._classify_by_labels(["Lin Jian"]), "release")
        self.assertEqual(S._classify_by_labels([]), "release")

    def test_no_asker_shaped_label_at_all_is_release(self):
        # Several named officials each quoted once, no one asking anything.
        labels = ["Yan Dong", "Lin Weilong", "Han Yong", "He Shaojun"]
        self.assertEqual(S._classify_by_labels(labels), "release")


class TestBuildExchanges(unittest.TestCase):
    """_build_exchanges() turns a paragraph list into typed Q/A/CONT
    exchanges via regex alone — no LLM, no reproduction, just slicing the
    original text at detected label boundaries."""

    def test_bare_single_letter_labels(self):
        # Bug #14: {1,40} required at least 1 char between the letter and
        # the colon, so bare "Q:"/"A:" (MOFCOM's translated 问：/答：) never
        # matched at all.
        paragraphs = [
            "Q: On May 2, 2026, the Ministry of Commerce issued a prohibition order.",
            "A: Since 2025, the United States has imposed sanctions on Chinese firms.",
        ]
        exchanges = S._build_exchanges(paragraphs, {"Spokesperson", "Answer"})
        self.assertEqual(len(exchanges), 2)
        self.assertEqual(exchanges[0]["type"], "Q")
        self.assertEqual(exchanges[1]["type"], "A")

    def test_accented_names_in_labels(self):
        # Bug #13: [A-Za-z0-9...] silently failed to match "É" in
        # "PRESIDENT ARÉVALO" — not misclassified, invisible.
        paragraphs = ["MODERATOR: Please begin.", "PRESIDENT ARÉVALO: Thank you."]
        exchanges = S._build_exchanges(paragraphs, set())
        self.assertEqual(len(exchanges), 2)
        self.assertEqual(exchanges[1]["speaker"], "PRESIDENT ARÉVALO")

    def test_metadata_fields_are_not_speaker_labels(self):
        # Bug found live 2026-09-01: a MOFCOM page's translated boilerplate
        # ("Category: News", "Source: Xinhua News Agency", "Type: Reprint")
        # matched the "Label: text" shape just as well as a real speaker,
        # tricking the parser into thinking a plain document had Q&A
        # structure.
        paragraphs = [
            "Category: News", "Source: Xinhua News Agency", "Type: Reprint",
            "China's Position on So-Called Overcapacity",
            "The Ministry of Commerce released this position paper today.",
        ]
        exchanges = S._build_exchanges(paragraphs, set())
        self.assertEqual(exchanges, [])  # no real speaker labels at all

    def test_repeated_unnamed_spokesperson_promoted_to_answerer(self):
        # A named spokesperson not in the role-word list (MOFCOM's
        # rotating regular-conference officials, unlike FMPRC's maintained
        # name list) — repeating 2+ times is the signal that promotes
        # them from the default "Q" typing to "A".
        paragraphs = [
            "He Yadong: Welcome everyone.",
            "CGTN Reporter: What about tariffs?",
            "He Yadong: We oppose unilateral tariffs.",
            "Kyodo News Reporter: What about rare earths?",
            "He Yadong: We have addressed this before.",
        ]
        exchanges = S._build_exchanges(paragraphs, {"Spokesperson"})
        he_yadong_types = {e["type"] for e in exchanges if e["speaker"] == "He Yadong"}
        self.assertEqual(he_yadong_types, {"A"})
        reporter_types = {e["type"] for e in exchanges if e["speaker"] and "Reporter" in e["speaker"]}
        self.assertEqual(reporter_types, {"Q"})

    def test_asterisk_divider_between_topics_is_dropped(self):
        # Real bug, live 2026-09-03: FMPRC's own pages use a bare line of
        # repeated asterisks as a visual divider between unrelated topics
        # within the same press conference. It doesn't match "Label: text"
        # so it fell through to a plain CONT paragraph and got written
        # into the tracker as if it were more of the PRECEDING answer.
        paragraphs = [
            "Spokesperson: China opposes unilateral tariff measures in all forms.",
            "**************************************************",
            "Reuters: Nepal once asked China for shared monitoring data.",
        ]
        exchanges = S._build_exchanges(paragraphs, {"Spokesperson"})
        self.assertEqual(len(exchanges), 2)
        self.assertTrue(all("*" not in e["text"] for e in exchanges))

    def test_other_dashed_or_underscored_dividers_are_also_dropped(self):
        for divider in ["----------", "__________", "=========="]:
            with self.subTest(divider=divider):
                paragraphs = [
                    "Spokesperson: Some real content here.",
                    divider,
                    "Reporter: A real follow-up question.",
                ]
                exchanges = S._build_exchanges(paragraphs, {"Spokesperson"})
                self.assertEqual(len(exchanges), 2)

    def test_short_real_dash_joined_phrase_is_not_treated_as_a_divider(self):
        # _SEPARATOR_LINE_RE requires the SAME character repeated 4+ times
        # — must not eat a real short hyphenated aside.
        paragraphs = ["Spokesperson: This is a well-known, long-standing policy."]
        exchanges = S._build_exchanges(paragraphs, {"Spokesperson"})
        self.assertEqual(len(exchanges), 1)

    def test_long_parenthetical_outlet_label_is_recognized(self):
        # Real bug, live 2026-09-03: a reporter's label translated as
        # "International Market News Agency (IMNA) Reporter:" (47 chars,
        # with parentheses) failed to match the old {0,40}, \w-only label
        # regex on TWO independent counts — too long, and \w doesn't cover
        # "(" / ")". The label fell through as an unrecognized CONT line,
        # which shifted filter_relevant_exchanges' block boundaries enough
        # that a real, US-mentioning exchange (about a proposed US tariff
        # hike) got silently folded into an unrelated block and dropped.
        paragraphs = [
            "Spokesperson: Thank you for the question.",
            "International Market News Agency (IMNA) Reporter: Media reports "
            "say the US government is considering a new tariff.",
        ]
        exchanges = S._build_exchanges(paragraphs, {"Spokesperson"})
        self.assertEqual(len(exchanges), 2)
        self.assertEqual(exchanges[1]["type"], "Q")
        self.assertEqual(exchanges[1]["speaker"], "International Market News Agency (IMNA) Reporter")


class TestMergeOrphanSpeakerLabels(unittest.TestCase):
    """Bugs #13/#14: a "SPEAKER:" label alone on its own line/paragraph,
    with the actual remarks starting on the NEXT line, is invisible to
    _QA_RE/_LABEL_RE (both require "Label: text" on ONE line) unless
    merged first."""

    def test_merges_orphan_label_with_following_paragraph(self):
        paragraphs = ["MODERATOR:", "Please welcome our guest.", "PRESIDENT ARÉVALO:", "Thank you all."]
        merged = S._merge_orphan_speaker_labels(paragraphs)
        self.assertEqual(merged, ["MODERATOR: Please welcome our guest.", "PRESIDENT ARÉVALO: Thank you all."])

    def test_bare_qa_shorthand_orphan_labels(self):
        # {0,40}, not {1,40}: a bare "Q:"/"A:" orphan label has ZERO
        # characters between the letter and colon.
        paragraphs = ["Q:", "What is your comment?", "A:", "We firmly oppose this."]
        merged = S._merge_orphan_speaker_labels(paragraphs)
        self.assertEqual(merged, ["Q: What is your comment?", "A: We firmly oppose this."])

    def test_non_orphan_paragraphs_untouched(self):
        paragraphs = ["Lin Jian: This is a complete statement on one line."]
        self.assertEqual(S._merge_orphan_speaker_labels(paragraphs), paragraphs)

    def test_orphan_label_with_parenthetical_outlet_name(self):
        # Real bug, live 2026-09-03 — see TestBuildExchanges's version of
        # this same fix. An orphan label this long, with a parenthetical
        # abbreviation, previously failed _ORPHAN_LABEL_RE on both counts
        # (length and character class) and never got merged at all.
        paragraphs = [
            "International Market News Agency (IMNA) Reporter:",
            "Media reports say the US government is considering a new tariff.",
        ]
        merged = S._merge_orphan_speaker_labels(paragraphs)
        self.assertEqual(merged, [
            "International Market News Agency (IMNA) Reporter: "
            "Media reports say the US government is considering a new tariff."
        ])


class TestUnbracketLabel(unittest.TestCase):
    """Bug #15: MOFCOM's *regular* press conferences wrap speaker names in
    brackets ("【何亚东】："), which translate_to_english carries through as
    English brackets ("[He Yadong]:") — a shape none of the "Label: text"
    regexes recognized (they all require the label to start with a plain
    letter, not "[")."""

    def test_strips_english_brackets_with_content(self):
        self.assertEqual(S._unbracket_label("[He Yadong]: Thank you."), "He Yadong: Thank you.")

    def test_strips_fullwidth_brackets(self):
        self.assertEqual(S._unbracket_label("【何亚东】：Thank you."), "何亚东: Thank you.")

    def test_orphan_bracketed_label_keeps_clean_colon(self):
        # Must NOT add a trailing space when there's no content on the
        # same line — that would break _ORPHAN_LABEL_RE's `:$` anchor
        # downstream (a real bug in the first draft of this fix).
        self.assertEqual(S._unbracket_label("[He Yadong]:"), "He Yadong:")

    def test_unbracketed_paragraph_untouched(self):
        text = "This paragraph has no leading bracketed label at all."
        self.assertEqual(S._unbracket_label(text), text)


class TestChineseUsMention(unittest.TestCase):
    """select_relevant_chinese_paragraphs()'s keyword filter — per user
    request 2026-09-01, replacing an LLM judgment call for Chinese-source
    release-type paragraph selection. Deliberately excludes bare 美元 (US
    dollar) the same way CHINESE_RELEVANCE_KEYWORDS does — every real term
    is a 2+ character compound whose second character isn't 元, so no
    separate exclusion regex is needed."""

    def test_selects_paragraphs_mentioning_the_us(self):
        text = "美国对华加征关税。\n中国坚决反对单边主义。\n中非合作持续深化。"
        selected = S.select_relevant_chinese_paragraphs(text)
        self.assertEqual(selected, ["美国对华加征关税。"])

    def test_bare_currency_mention_does_not_match(self):
        self.assertIsNone(S._CHINESE_US_MENTION_RE.search("5.5万亿元（约8100亿美元）"))

    def test_real_us_reference_terms_match(self):
        for term in ["美国", "中美", "美方", "华盛顿", "白宫"]:
            with self.subTest(term=term):
                self.assertIsNotNone(S._CHINESE_US_MENTION_RE.search(f"关于{term}的声明"))


class TestMarkSeenFlushOrdering(unittest.TestCase):
    """Bug #21: mark_seen() used to commit to tracker.db immediately when
    an item was queued, before its content was durably written to the
    doc (which only happens once per SOURCE, in flush_pending_entries()).
    A crash in between permanently marked the item seen — never retried —
    while its content vanished. Fixed: mark_seen() now happens INSIDE
    flush_pending_entries(), after doc.save() succeeds.

    IMPORTANT: flush_pending_entries() always calls doc.save(S.DOC_PATH)
    — the module-level constant, NOT a path derived from the `doc` object
    passed in — so passing a throwaway in-memory Document() here is NOT
    enough isolation on its own; every test below must also patch
    S.DOC_PATH to a temp file. Found 2026-09-02: neither test originally
    did this, so every test run was silently overwriting the REAL
    output/tracker_output.docx with this test's synthetic one-entry
    fixture — a real, ongoing data-loss bug in the test suite itself,
    not just a theoretical risk. (test_x_since_id_staged_not_saved_
    until_flush already patched S.X_STATE_PATH the same way — this was
    the same lesson, just not applied to S.DOC_PATH too.)"""

    def setUp(self):
        import sqlite3
        self.conn = sqlite3.connect(":memory:")
        self.conn.execute("CREATE TABLE IF NOT EXISTS seen_urls (url TEXT PRIMARY KEY, date_seen TEXT)")
        self.conn.execute(
            "CREATE TABLE IF NOT EXISTS entries (id INTEGER PRIMARY KEY AUTOINCREMENT, "
            "url TEXT, date TEXT NOT NULL, kind TEXT NOT NULL, "
            "summary TEXT, anchor TEXT, exchanges_json TEXT, paragraphs_json TEXT, "
            "source_label TEXT)"
        )
        self.conn.commit()
        S.PENDING_ENTRIES.clear()
        S._PENDING_X_SINCE_IDS.clear()
        S._QUEUED_URL_SLUGS_THIS_RUN.clear()
        self._tmpdir = tempfile.TemporaryDirectory()
        self._orig_doc_path = S.DOC_PATH
        S.DOC_PATH = os.path.join(self._tmpdir.name, "test_tracker_output.docx")

    def tearDown(self):
        S.PENDING_ENTRIES.clear()
        S._PENDING_X_SINCE_IDS.clear()
        S._QUEUED_URL_SLUGS_THIS_RUN.clear()
        S.DOC_PATH = self._orig_doc_path
        self._tmpdir.cleanup()

    def test_not_marked_seen_until_flush(self):
        url = "https://example.com/test-item"
        S.queue_entry("release", datetime(2026, 9, 1), "Test summary", url, "Test", paragraphs=["body"])
        self.assertFalse(S.is_seen(self.conn, url), "must NOT be marked seen before flush — a crash here must be safely retryable")

        doc = S.Document()
        S._set_doc_defaults(doc)
        n = S.flush_pending_entries(doc, self.conn)
        self.assertEqual(n, 1)
        self.assertTrue(S.is_seen(self.conn, url), "must be marked seen once the write is durably confirmed")

    def test_x_since_id_staged_not_saved_until_flush(self):
        # Bug #22 — same ordering mistake, different subsystem. Staged
        # under a search-GROUP key ("china"/"prc"), not a username, since
        # the 2026-09-02 search-based redesign — see scrape_x()'s comment.
        state_path = os.path.join(self._tmpdir.name, "x_state.json")
        orig_path = S.X_STATE_PATH
        S.X_STATE_PATH = state_path
        try:
            S._PENDING_X_SINCE_IDS["china"] = "999999"
            self.assertEqual(S._load_x_state(), {}, "must NOT be on disk before flush")

            doc = S.Document()
            S._set_doc_defaults(doc)
            S.flush_pending_entries(doc, self.conn)  # 0 doc entries, since_id still must persist
            state = S._load_x_state()
            self.assertEqual(state.get("_search_groups", {}).get("china", {}).get("since_id"), "999999")
        finally:
            S.X_STATE_PATH = orig_path


class TestContentTypeFromParagraphs(unittest.TestCase):
    """content_type_from_paragraphs() end-to-end on realistic paragraph
    lists — exercises _LABEL_RE + _classify_by_labels together."""

    def test_position_paper_with_metadata_is_release(self):
        paragraphs = [
            "Category: News", "Source: Xinhua News Agency",
            "China's Position on So-Called Overcapacity",
            "The following is the full text of the position paper.",
        ]
        self.assertEqual(S.content_type_from_paragraphs(paragraphs), "release")

    def test_real_transcript_is_qa(self):
        paragraphs = ["Reuters: What is your comment?", "Lin Jian: We firmly oppose this.",
                      "AFP: Any further remarks?", "Lin Jian: That is all for today."]
        self.assertEqual(S.content_type_from_paragraphs(paragraphs), "qa")


class TestEnglishColonSpacing(unittest.TestCase):
    """MND's _ENGLISH_COLON_RE normalizes "Label:text" (no space) to
    "Label: text" before _QA_RE ever sees it — and _QA_RE requires a
    space after the colon to match at all, so a name this regex can't
    recognize breaks BOTH regexes in sequence, invisibly. Fixed
    preventively (not from an observed live failure) using the same \\w
    fix already proven correct 4 times today for the identical Latin-only-
    character-class shape."""

    def test_accented_name_gets_spaced(self):
        self.assertEqual(
            S._ENGLISH_COLON_RE.sub(r"\1: ", "Minister García:We discussed trade issues."),
            "Minister García: We discussed trade issues.",
        )

    def test_plain_ascii_name_still_works(self):
        self.assertEqual(
            S._ENGLISH_COLON_RE.sub(r"\1: ", "Question:What is your comment?"),
            "Question: What is your comment?",
        )


class TestItemUrl(unittest.TestCase):
    """item_url()'s guid fallback only ever fires in practice for a WP-API
    item missing "link" — and WP-API's raw `guid` field is a dict
    ({"rendered": "..."}), not a string. Untested-in-practice (every real
    source's WP-API responses include "link"), but hardened 2026-09-02 so
    a future source that omits it doesn't silently poison dedup/doc
    hyperlinks with a dict instead of a URL string."""

    def test_uses_link_when_present(self):
        self.assertEqual(S.item_url({"link": "https://example.com/a"}), "https://example.com/a")

    def test_falls_back_to_string_guid(self):
        self.assertEqual(S.item_url({"guid": "https://example.com/b"}), "https://example.com/b")

    def test_falls_back_to_dict_guid_rendered(self):
        self.assertEqual(
            S.item_url({"guid": {"rendered": "https://example.com/c"}}),
            "https://example.com/c",
        )

    def test_empty_link_falls_back_to_guid(self):
        self.assertEqual(S.item_url({"link": "", "guid": "https://example.com/d"}), "https://example.com/d")

    def test_neither_present_returns_empty_string(self):
        self.assertEqual(S.item_url({}), "")


class TestDefaultWeekRange(unittest.TestCase):
    """The tracker's week is Tuesday-through-Monday. Run on Monday or on
    the Tuesday right after it, and both must resolve to the exact same
    completed week — that's the whole point of the "most recent complete
    week" default (see default_week_range()'s docstring), not just a
    same-answer-most-of-the-time coincidence."""

    def test_monday_gives_last_tuesday_through_today(self):
        start, end = S.default_week_range(date(2026, 8, 31))  # a Monday
        self.assertEqual((start, end), (date(2026, 8, 25), date(2026, 8, 31)))

    def test_tuesday_gives_the_same_week_as_the_monday_before_it(self):
        start, end = S.default_week_range(date(2026, 9, 1))  # the next day, Tuesday
        self.assertEqual((start, end), (date(2026, 8, 25), date(2026, 8, 31)))

    def test_every_other_weekday_also_reports_last_weeks_completed_range(self):
        for d in [date(2026, 9, 2), date(2026, 9, 3), date(2026, 9, 4),
                  date(2026, 9, 5), date(2026, 9, 6)]:  # Wed through Sun
            with self.subTest(day=d):
                self.assertEqual(
                    S.default_week_range(d),
                    (date(2026, 8, 25), date(2026, 8, 31)),
                )

    def test_week_label_same_month(self):
        self.assertEqual(
            S._format_week_for_filename(date(2026, 8, 25), date(2026, 8, 31)),
            "Aug 25-31, 2026",
        )

    def test_week_label_spanning_months(self):
        self.assertEqual(
            S._format_week_for_filename(date(2026, 8, 25), date(2026, 9, 1)),
            "Aug 25-Sep 1, 2026",
        )


class TestFormatDuration(unittest.TestCase):
    """The final run summary's "Took ..." line — whichever unit reads
    naturally for the elapsed time, no leading zero units."""

    def test_seconds_only(self):
        self.assertEqual(S._format_duration(47), "47s")

    def test_minutes_and_seconds(self):
        self.assertEqual(S._format_duration(967), "16m 7s")

    def test_hours_and_minutes(self):
        self.assertEqual(S._format_duration(3900), "1h 05m")


class TestHallucinatedOfficials(unittest.TestCase):
    """get_summary_and_anchor()'s programmatic backstop against a real
    failure: generate_summary() invented "Treasury Secretary Janet
    Yellen" for a real G20 statement that names no individual official
    anywhere — Yellen hasn't been Treasury Secretary since Jan 2025."""

    def test_flags_a_name_not_in_the_source(self):
        summary = "Treasury Secretary Janet Yellen announced new measures."
        source = "We, the G20 Finance Ministers, met today to discuss trade."
        self.assertEqual(S._hallucinated_officials(summary, source), ["Yellen"])

    def test_does_not_flag_a_name_that_really_is_in_the_source(self):
        summary = "Treasury Secretary Scott Bessent testified before Congress."
        source = "Secretary Bessent said the economy remains strong."
        self.assertEqual(S._hallucinated_officials(summary, source), [])

    def test_case_insensitive_match_against_an_all_caps_transcript_label(self):
        # Real bug, live 2026-09-03: a real State Department interview
        # transcript labels its speaker "SECRETARY RUBIO:" (all caps, a
        # common transcript convention) — a case-sensitive check missed
        # this entirely and falsely flagged a correctly-named summary as
        # a hallucination, triggering a needless (and less accurate)
        # regeneration.
        summary = "Secretary of State Marco Rubio discussed the Venezuela oil deal."
        source = "SECRETARY RUBIO: This arrangement benefits US energy security."
        self.assertEqual(S._hallucinated_officials(summary, source), [])

    def test_no_names_in_summary_is_clean(self):
        summary = "The Treasury Department released a new report on trade."
        source = "This report covers trade statistics for the past year."
        self.assertEqual(S._hallucinated_officials(summary, source), [])

    def test_trump_is_not_checked(self):
        # Deliberately excluded — "Trump administration"/"Trump tariffs"
        # show up constantly as an adjective, not a claim that he
        # personally did the specific act being summarized.
        summary = "The Trump administration announced new tariffs on China."
        source = "A White House official announced new tariffs today."
        self.assertEqual(S._hallucinated_officials(summary, source), [])


class TestSummarizeRunUsage(unittest.TestCase):
    """Auto-shown cost/token summary at the end of a run — added so a
    non-technical user sees this without a separate manual command.
    USAGE_LOG_PATH accumulates across the pipeline's entire history, so
    this must only sum records at or after the calling run's own start
    time, not every run ever made."""

    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self._orig_path = S.USAGE_LOG_PATH
        S.USAGE_LOG_PATH = os.path.join(self._tmpdir.name, "usage_log.jsonl")

    def tearDown(self):
        S.USAGE_LOG_PATH = self._orig_path
        self._tmpdir.cleanup()

    def _write(self, ts, provider, total_tokens, usd):
        with open(S.USAGE_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps({
                "ts": ts, "label": "test", "provider": provider,
                "prompt_tokens": 0, "completion_tokens": 0,
                "reasoning_tokens": 0, "total_tokens": total_tokens, "usd": usd,
            }) + "\n")

    def test_ignores_records_before_run_start(self):
        self._write("2026-01-01T00:00:00", "Gemini", 500, 0.001)  # a prior run
        self._write("2026-01-02T00:00:00", "Gemini", 300, 0.002)  # this run
        llm_tokens, x_reads, usd = S._summarize_run_usage("2026-01-01T12:00:00")
        self.assertEqual(llm_tokens, 300)
        self.assertAlmostEqual(usd, 0.002)

    def test_x_reads_counted_separately_from_llm_tokens(self):
        self._write("2026-01-02T00:00:00", "Gemini", 300, 0.002)
        self._write("2026-01-02T00:00:01", "X API", 5, 0.025)
        llm_tokens, x_reads, usd = S._summarize_run_usage("2026-01-01T00:00:00")
        self.assertEqual(llm_tokens, 300)
        self.assertEqual(x_reads, 5)
        self.assertAlmostEqual(usd, 0.027)

    def test_no_log_file_returns_zeros(self):
        self.assertEqual(S._summarize_run_usage("2026-01-01T00:00:00"), (0, 0, 0.0))


class TestSourceErrorCapture(unittest.TestCase):
    """The end-of-run "which sources failed" report (added 2026-09-03 per
    user request) hooks this handler onto `log` for the duration of each
    source's run() call — it must only collect ERROR+, not the routine
    INFO-level "0 new items"/"skipping" messages every healthy source
    logs constantly."""

    def test_captures_error_level_messages(self):
        capture = S._SourceErrorCapture()
        S.log.addHandler(capture)
        try:
            S.log.error("[scio] Failed to fetch list")
            S.log.error("[scio] Error on http://example.com: timeout")
        finally:
            S.log.removeHandler(capture)
        self.assertEqual(capture.messages, [
            "[scio] Failed to fetch list",
            "[scio] Error on http://example.com: timeout",
        ])

    def test_does_not_capture_info_level_messages(self):
        # A healthy source logs plenty of INFO — "0 new items", "No
        # relevant exchanges — skipping" — none of that is a failure.
        capture = S._SourceErrorCapture()
        S.log.addHandler(capture)
        try:
            S.log.info("[scio] 0 new items")
            S.log.warning("[scio] Something worth a second look, not an error")
        finally:
            S.log.removeHandler(capture)
        self.assertEqual(capture.messages, [])


class TestQueueEntryHardDateFilter(unittest.TestCase):
    """queue_entry() now HARD-filters by _RUN_TARGET_START/_RUN_TARGET_END
    — found live, 2026-09-02: a run targeting "Aug 25-31" was still
    queuing a genuinely-found Sept 1 MFA leadership item into that week's
    output, because the range was previously just a label, never a
    filter. An out-of-range item must be dropped WITHOUT being marked
    seen, so a future run whose target covers it still finds it."""

    def setUp(self):
        S.PENDING_ENTRIES.clear()
        S._QUEUED_URL_SLUGS_THIS_RUN.clear()
        self._orig_start, self._orig_end = S._RUN_TARGET_START, S._RUN_TARGET_END

    def tearDown(self):
        S.PENDING_ENTRIES.clear()
        S._QUEUED_URL_SLUGS_THIS_RUN.clear()
        S._RUN_TARGET_START, S._RUN_TARGET_END = self._orig_start, self._orig_end

    def test_in_range_entry_is_queued(self):
        S._RUN_TARGET_START, S._RUN_TARGET_END = date(2026, 8, 25), date(2026, 8, 31)
        S.queue_entry("release", datetime(2026, 8, 28), "s", "https://example.com/a", "s")
        self.assertEqual(len(S.PENDING_ENTRIES), 1)

    def test_out_of_range_entry_is_dropped_not_queued(self):
        S._RUN_TARGET_START, S._RUN_TARGET_END = date(2026, 8, 25), date(2026, 8, 31)
        S.queue_entry("release", datetime(2026, 9, 1), "s", "https://example.com/b", "s")
        self.assertEqual(len(S.PENDING_ENTRIES), 0)

    def test_no_target_range_set_queues_everything(self):
        # backtest.py / format_entry.py never set a target range — must
        # not filter at all in that case.
        S._RUN_TARGET_START, S._RUN_TARGET_END = None, None
        S.queue_entry("release", datetime(2026, 9, 1), "s", "https://example.com/c", "s")
        self.assertEqual(len(S.PENDING_ENTRIES), 1)


class TestUrlDedupSlug(unittest.TestCase):
    """Real bug, live 2026-09-03: state.gov's own WP-API listed the SAME
    "Foundry School" headline under two different post IDs two days
    apart, with URL slugs differing by exactly one missing hyphen
    ("workforce-behind" vs "workforcebehind") — a source-side republish,
    not something our normal URL-based is_seen() dedup can catch since
    the URLs are genuinely different strings."""

    def test_collapses_a_missing_hyphen(self):
        a = "https://www.state.gov/.../launches-foundry-school-to-build-the-workforce-behind-americas-comeback/"
        b = "https://www.state.gov/.../launches-foundry-school-to-build-the-workforcebehind-americas-comeback/"
        self.assertEqual(S._url_dedup_slug(a), S._url_dedup_slug(b))

    def test_genuinely_different_slugs_stay_different(self):
        a = "https://www.state.gov/releases/some-real-article-about-taiwan/"
        b = "https://www.state.gov/releases/a-totally-different-story-about-iran/"
        self.assertNotEqual(S._url_dedup_slug(a), S._url_dedup_slug(b))


class TestQueueEntryRepublishDuplicateGuard(unittest.TestCase):
    """queue_entry()'s second, narrower dedup net on top of is_seen() —
    see _QUEUED_URL_SLUGS_THIS_RUN's docstring for the real case."""

    def setUp(self):
        S.PENDING_ENTRIES.clear()
        S._QUEUED_URL_SLUGS_THIS_RUN.clear()
        self._orig_start, self._orig_end = S._RUN_TARGET_START, S._RUN_TARGET_END
        S._RUN_TARGET_START, S._RUN_TARGET_END = None, None

    def tearDown(self):
        S.PENDING_ENTRIES.clear()
        S._QUEUED_URL_SLUGS_THIS_RUN.clear()
        S._RUN_TARGET_START, S._RUN_TARGET_END = self._orig_start, self._orig_end

    def test_second_near_identical_slug_is_dropped(self):
        S.queue_entry("release", datetime(2026, 9, 1), "s1",
                       "https://www.state.gov/x/workforce-behind-comeback/", "s1")
        S.queue_entry("release", datetime(2026, 9, 3), "s2",
                       "https://www.state.gov/x/workforcebehind-comeback/", "s2")
        self.assertEqual(len(S.PENDING_ENTRIES), 1)

    def test_genuinely_different_urls_both_queued(self):
        S.queue_entry("release", datetime(2026, 9, 1), "s1", "https://www.state.gov/x/story-one/", "s1")
        S.queue_entry("release", datetime(2026, 9, 1), "s2", "https://www.state.gov/x/story-two/", "s2")
        self.assertEqual(len(S.PENDING_ENTRIES), 2)


class TestRenderDocForRange(unittest.TestCase):
    """The per-week output doc is now rendered FRESH from the `entries`
    table, not a copy of the ever-growing master doc — found live,
    2026-09-02, that the "dated" file was actually the entire cumulative
    history under a misleadingly narrow filename."""

    def setUp(self):
        self._tmpdir = tempfile.TemporaryDirectory()
        self.conn = sqlite3.connect(":memory:")
        self.conn.execute(
            "CREATE TABLE entries (id INTEGER PRIMARY KEY AUTOINCREMENT, "
            "url TEXT, date TEXT NOT NULL, kind TEXT NOT NULL, "
            "summary TEXT, anchor TEXT, exchanges_json TEXT, paragraphs_json TEXT, "
            "source_label TEXT)"
        )

    def tearDown(self):
        self.conn.close()
        self._tmpdir.cleanup()

    def _insert(self, date_str, url, summary="Something happened.", source_label=None):
        self.conn.execute(
            "INSERT INTO entries (url, date, kind, summary, anchor, exchanges_json, paragraphs_json, source_label) "
            "VALUES (?, ?, 'release', ?, ?, NULL, ?, ?)",
            (url, date_str, summary, summary, json.dumps([summary]), source_label),
        )
        self.conn.commit()

    def test_only_entries_in_range_are_included(self):
        self._insert("2026-08-28", "https://example.com/a", summary="In-range item happened.")
        self._insert("2026-09-01", "https://example.com/b", summary="Out-of-range item happened.")
        doc = S.render_doc_for_range(self.conn, date(2026, 8, 25), date(2026, 8, 31))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("In-range item", full_text)
        self.assertNotIn("Out-of-range item", full_text)

    def test_already_scraped_week_still_renders_without_a_new_run(self):
        # Simulates "if you already got those dates, just give me the
        # doc" — entries inserted by an EARLIER run still render
        # correctly for a later call with no new queuing at all.
        self._insert("2026-08-26", "https://example.com/c", summary="Already scraped item.")
        doc = S.render_doc_for_range(self.conn, date(2026, 8, 25), date(2026, 8, 31))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Already scraped item", full_text)

    def test_empty_range_still_gets_every_days_heading(self):
        # A day with nothing found still gets its own date heading, no
        # body under it — confirmed against the real past trackers (a
        # quiet "Sunday, August 2, 2026" heading immediately followed by
        # the next day's heading, nothing in between). "Nothing found"
        # is real information, not the same as "we didn't check."
        doc = S.render_doc_for_range(self.conn, date(2026, 8, 25), date(2026, 8, 27))
        headings = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        self.assertEqual(headings, [
            "Tuesday, August 25, 2026",
            "Wednesday, August 26, 2026",
            "Thursday, August 27, 2026",
        ])


class TestMfaLeadershipShouldSkipTranslate(unittest.TestCase):
    """MFA leadership's free pre-filter — a plain keyword gate against
    CHINESE_RELEVANCE_KEYWORDS, same as FMPRC/MOFCOM/MND. Was briefly a
    length-OR-keyword hybrid (2026-09-02) to avoid dropping a real
    past-tracker entry with no keyword hit (a Wang Yi/Global Development
    Initiative item) — the user confirmed that entry was itself a human
    coding error in the original tracker, not a genuine editorial
    exception, so this source now matches every other Chinese source."""

    def test_no_keyword_is_skipped(self):
        text = "外交部发言人将于明日举行例行记者会。"  # a routine notice, no keyword
        self.assertTrue(S._mfa_leadership_should_skip_translate(text))

    def test_keyword_hit_is_not_skipped(self):
        text = "王毅同美国国务卿通话，讨论关税问题。"  # hits 美国/关税
        self.assertFalse(S._mfa_leadership_should_skip_translate(text))

    def test_long_text_with_no_keyword_is_still_skipped(self):
        # Length alone is no longer a safety net — even a long document
        # gets skipped if it never hits a real topic keyword, matching
        # every other Chinese source's behavior.
        text = "全球发展倡议合作理念与实践。" * 60
        self.assertTrue(S._mfa_leadership_should_skip_translate(text))


class TestBuildXSearchQuery(unittest.TestCase):
    """The search-based X redesign (2026-09-02) — combines multiple
    accounts + a keyword filter into one query, so X only returns (and
    bills for) tweets that already match, instead of every tweet from
    every account. See scrape_x()'s module comment for the real cost
    data (94% of one run's spend) that motivated this."""

    def test_combines_terms_and_accounts(self):
        query = S.build_x_search_query(["alice", "bob"], "china OR taiwan")
        self.assertEqual(query, "(china OR taiwan) (from:alice OR from:bob)")

    def test_single_account(self):
        query = S.build_x_search_query(["alice"], "china")
        self.assertEqual(query, "(china) (from:alice)")

    def test_real_groups_stay_under_the_512_char_query_limit(self):
        # The two actual groups this pipeline builds — a regression test
        # against ever silently exceeding "recent search"'s query cap as
        # accounts/terms are added over time.
        prc_accounts = [a for a in S.X_ACCOUNTS_NORMAL if a in S._PRC_X_ACCOUNTS]
        other_accounts = [a for a in S.X_ACCOUNTS_NORMAL if a not in S._PRC_X_ACCOUNTS]
        china_query = S.build_x_search_query(other_accounts, S._X_CHINA_SEARCH_TERMS)
        prc_query = S.build_x_search_query(prc_accounts, S._X_US_SEARCH_TERMS)
        self.assertLess(len(china_query), 512)
        self.assertLess(len(prc_query), 512)


class TestParseUserDate(unittest.TestCase):
    """--start/--end (and the .command prompt) accept plain YYYYMMDD, not
    just YYYY-MM-DD — dashes felt unnecessarily fussy to type for a
    plain-language terminal prompt, per user request 2026-09-02."""

    def test_plain_digits(self):
        self.assertEqual(S._parse_user_date("20260804"), date(2026, 8, 4))

    def test_dashed_iso_format_still_works(self):
        self.assertEqual(S._parse_user_date("2026-08-04"), date(2026, 8, 4))

    def test_garbage_raises_value_error(self):
        with self.assertRaises(ValueError):
            S._parse_user_date("08/04/2026")


class TestXmlSafe(unittest.TestCase):
    """A stray control character in text passed to a Word run makes
    lxml raise a hard ValueError and crash the WHOLE run, not just skip
    that one entry — found live, 2026-09-02, a real Aug 18-24 run
    crashed mid-flush on exactly this inside generate_summary()'s LLM
    output, the first real exercise of that path since ENABLE_LLM_
    SUMMARY was turned back on."""

    def test_strips_null_byte(self):
        self.assertEqual(S._xml_safe("before\x00after"), "beforeafter")

    def test_strips_other_control_chars(self):
        self.assertEqual(S._xml_safe("a\x0bb\x0cc\x1fd"), "abcd")

    def test_preserves_tab_newline_cr(self):
        self.assertEqual(S._xml_safe("a\tb\nc\rd"), "a\tb\nc\rd")

    def test_preserves_normal_and_cjk_text(self):
        text = "Lin Jian answered questions on 中美关系."
        self.assertEqual(S._xml_safe(text), text)

    def test_run_does_not_crash_on_bad_text(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        S._run(p, "bad\x00text")  # must not raise
        self.assertEqual(p.text, "badtext")

    def test_add_hyperlink_does_not_crash_on_bad_text(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        S.add_hyperlink(p, "bad\x00anchor", "https://example.com")  # must not raise
        self.assertNotIn("\x00", p.text)


class TestStripStrayMarkup(unittest.TestCase):
    """A real White House/Ford entry came out with literal markdown bold
    and an HTML anchor tag in the finished doc — confirmed by re-fetching
    the actual source page that the raw scraped text has none of this
    markup, so extract_key_paragraphs()'s LLM call added it despite
    being told to preserve exact text. Found live, 2026-09-02."""

    def test_strips_real_bad_text(self):
        bad = ('**Ford Motor Company** <a href="https://example.com">'
               '**announced**</a> **it will reshore production**.')
        cleaned = S._strip_stray_markup(bad)
        self.assertNotIn("**", cleaned)
        self.assertNotIn("<a", cleaned)
        self.assertIn("Ford Motor Company announced it will reshore production", cleaned)

    def test_does_not_eat_a_real_less_than_comparison(self):
        # A blanket "<[^>]+>" strip would mistake this for a tag — the
        # regex must only match a recognized HTML tag name.
        text = "GDP growth stayed <5% target this quarter."
        self.assertEqual(S._strip_stray_markup(text), text)

    def test_run_strips_markup_from_written_text(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        S._run(p, "**bold claim** with <b>a tag</b>")
        self.assertEqual(p.text, "bold claim with a tag")


class TestSourceLabelInReleaseBody(unittest.TestCase):
    """Every body paragraph in the real past trackers is bold-prefixed
    with who's speaking — a named official (Q&A, already handled), an X
    account ("Chinese Embassy:", "Rapid Response 47:"), or a plain
    institutional label for a release with no named speaker at all
    ("State Council Press Release: A senior official..."). Confirmed
    directly against the real tracker docx before implementing. Added
    2026-09-02 per user request."""

    def _first_para_text_and_bold(self, doc):
        # paragraph 0 is the summary line; paragraph 1 is the first body paragraph
        p = doc.paragraphs[1]
        first_run_bold = p.runs[0].bold if p.runs else None
        return p.text, first_run_bold

    def test_source_label_bolded_on_first_paragraph_only(self):
        from docx import Document
        doc = Document()
        S._set_doc_defaults(doc)
        S.add_release_entry_body(
            doc, "Summary sentence.",
            ["A senior official said something.", "The official continued speaking."],
            source_label="State Council Press Release",
        )
        text, bold = self._first_para_text_and_bold(doc)
        self.assertTrue(text.startswith("State Council Press Release:"))
        self.assertTrue(bold)
        # second paragraph must NOT repeat the label
        second_para = doc.paragraphs[2]
        self.assertFalse(second_para.text.startswith("State Council Press Release"))

    def test_existing_speaker_label_wins_over_source_label(self):
        # A paragraph that already has its own "Name: text" shape (e.g. a
        # Q&A-shaped release fallback) must not get a SECOND label
        # stacked on top of it.
        from docx import Document
        doc = Document()
        S._set_doc_defaults(doc)
        S.add_release_entry_body(
            doc, "Summary sentence.",
            ["Zhang Xiaogang: China's position is clear."],
            source_label="Ministry of National Defense",
        )
        text, bold = self._first_para_text_and_bold(doc)
        self.assertTrue(text.startswith("Zhang Xiaogang:"))
        self.assertNotIn("Ministry of National Defense", text)

    def test_x_account_display_names_cover_all_configured_accounts(self):
        # Every account this pipeline actually polls must have a real
        # display name, not silently fall back to the raw username.
        all_accounts = S.X_ACCOUNTS_NORMAL + S.X_ACCOUNTS_LESS_IMPORTANT
        for account in all_accounts:
            with self.subTest(account=account):
                self.assertIn(account, S._X_ACCOUNT_DISPLAY_NAMES)

    def test_unknown_x_account_falls_back_to_username(self):
        self.assertEqual(
            S._X_ACCOUNT_DISPLAY_NAMES.get("SomeFutureAccount", "SomeFutureAccount"),
            "SomeFutureAccount",
        )


class TestIsolateItemErrors(unittest.TestCase):
    """
    _isolate_item_errors() — the shared per-item error boundary added
    2026-09-04 (see NOTES.md) after finding source loops with no
    isolation at all could let one item's exception abort every
    remaining item in that source for the rest of the run.
    """

    def test_swallows_an_exception_and_logs_it(self):
        with self.assertLogs(S.log, level="ERROR") as cm:
            with S._isolate_item_errors("testsrc", "http://example.com/bad"):
                raise RuntimeError("simulated failure")
        self.assertIn("[testsrc] Error on http://example.com/bad", cm.output[0])
        self.assertIn("simulated failure", cm.output[0])

    def test_does_not_swallow_success_or_return_value(self):
        ran = []
        with S._isolate_item_errors("testsrc", "http://example.com/ok"):
            ran.append("did work")
        self.assertEqual(ran, ["did work"])

    def test_a_loop_keeps_going_after_one_item_fails(self):
        # The actual regression this exists to prevent: item #2 raising
        # must not stop items #3/#4 from ever being attempted.
        attempted = []
        items = ["ok1", "bad", "ok2", "ok3"]
        for item in items:
            with S._isolate_item_errors("testsrc", item):
                attempted.append(item)
                if item == "bad":
                    raise RuntimeError("simulated failure")
        self.assertEqual(attempted, items)


class TestAtomicDocSave(unittest.TestCase):
    """
    _atomic_doc_save() — added 2026-09-04 per user request ("save
    progress along the way... so if crash then progress still saved").
    A plain doc.save(path) writes straight to the target and can leave a
    corrupted file if the process dies mid-write; this must never do that.
    """

    def test_success_leaves_no_temp_file_behind(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as d:
            path = os.path.join(d, "out.docx")
            doc = Document()
            doc.add_paragraph("hello")
            S._atomic_doc_save(doc, path)
            self.assertTrue(os.path.exists(path))
            leftovers = [f for f in os.listdir(d) if f != "out.docx"]
            self.assertEqual(leftovers, [])

    def test_failure_leaves_original_untouched_and_no_temp_file(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as d:
            path = os.path.join(d, "out.docx")
            Document().save(path)
            original_bytes = open(path, "rb").read()

            class ExplodingDoc:
                def save(self, p):
                    raise RuntimeError("simulated crash mid-save")

            with self.assertRaises(RuntimeError):
                S._atomic_doc_save(ExplodingDoc(), path)

            self.assertEqual(open(path, "rb").read(), original_bytes)
            leftovers = [f for f in os.listdir(d) if f != "out.docx"]
            self.assertEqual(leftovers, [])


class TestParseStateResultDate(unittest.TestCase):
    """
    _parse_state_result_date() — found live, 2026-09-04, the same day
    state.gov's real pagination shipped: some `.collection-result-meta`
    containers have an extra leading span naming an official (e.g.
    "Marco Rubio") BEFORE the real date span, which a naive "just take
    the first span" selector silently misread as the date, dropping
    those specific entries.
    """

    def _meta(self, html: str):
        from bs4 import BeautifulSoup
        return BeautifulSoup(html, "html.parser").find("div")

    def test_leading_name_span_is_skipped_for_the_real_date_span(self):
        meta = self._meta(
            '<div class="collection-result-meta">'
            '<span>Marco Rubio</span><span>August 12, 2026</span></div>'
        )
        self.assertEqual(S._parse_state_result_date(meta), S.date(2026, 8, 12))

    def test_single_date_span_with_no_name(self):
        meta = self._meta(
            '<div class="collection-result-meta"><span>September 1, 2026</span></div>'
        )
        self.assertEqual(S._parse_state_result_date(meta), S.date(2026, 9, 1))

    def test_no_parseable_date_returns_none(self):
        meta = self._meta(
            '<div class="collection-result-meta"><span>Someone Else</span></div>'
        )
        self.assertIsNone(S._parse_state_result_date(meta))


class TestParseMofcomListDate(unittest.TestCase):
    """
    _parse_mofcom_list_date() — the list widget shows dates in two
    different formats depending on which MOFCOM index it's rendering:
    bracketed "[2026-09-03]" on the Chinese xwfb/* sections, bare
    "08/04/2026" on the English press-conference index.
    """

    def test_bracketed_chinese_section_format(self):
        self.assertEqual(S._parse_mofcom_list_date("[2026-09-03]"), S.date(2026, 9, 3))

    def test_unbracketed_chinese_section_format(self):
        self.assertEqual(S._parse_mofcom_list_date("2026-09-03"), S.date(2026, 9, 3))

    def test_english_press_conference_format(self):
        self.assertEqual(S._parse_mofcom_list_date("08/04/2026"), S.date(2026, 8, 4))

    def test_none_input_returns_none(self):
        self.assertIsNone(S._parse_mofcom_list_date(None))

    def test_unparseable_text_returns_none(self):
        self.assertIsNone(S._parse_mofcom_list_date("not a date"))


if __name__ == "__main__":
    unittest.main(verbosity=2)
